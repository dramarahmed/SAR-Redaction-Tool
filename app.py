"""
SAR Redaction Tool
NHS Subject Access Request · Multi-format document bundle processor
UK GDPR / DPA 2018 / ICO / BMA / NHS England guidance
Human-in-the-loop redaction review
"""

import streamlit as st
import streamlit.components.v1 as components
import ollama
import fitz  # PyMuPDF
import base64
import datetime
import time
import io
import os
import re
import json
import zipfile
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from pathlib import Path

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text as parse_rtf
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    from PIL import Image as PILImage, ImageDraw as PILImageDraw
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pytesseract
    if not PIL_AVAILABLE:
        from PIL import Image as PILImage  # ensure PILImage defined even if above block failed
    # Confirm the Tesseract binary is actually present before claiming it's available
    pytesseract.get_tesseract_version()
    TESSERACT_AVAILABLE = True
except Exception:
    TESSERACT_AVAILABLE = False


# =============================================================================
# Timing helpers
# =============================================================================

def _fmt_elapsed(seconds: float) -> str:
    """Format a number of seconds as '1m 23s' or '45s'."""
    s = int(seconds)
    if s < 60:
        return f"{s}s"
    return f"{s // 60}m {s % 60:02d}s"


def _timing_suffix(start: float, done: int, total: int) -> str:
    """Return '  |  elapsed 1m 23s  ·  ~2m 10s remaining' string."""
    elapsed = time.time() - start
    parts = [f"elapsed {_fmt_elapsed(elapsed)}"]
    if done > 0 and done < total:
        remaining = elapsed / done * (total - done)
        parts.append(f"~{_fmt_elapsed(remaining)} remaining")
    return "  |  " + "  ·  ".join(parts)


# =============================================================================
# NHS SAR Redaction Ontology
# =============================================================================

REDACTION_TAGS = {
    # ── Auto-redact ───────────────────────────────────────────────────────────
    "THIRD_PARTY_IDENTIFIER": {
        "label":  "Third-party identifier",
        "desc":   "Name or identifying detail of a private individual (family, carer, neighbour, friend)",
        "action": "redact",
    },
    "CONFIDENTIAL_DISCLOSURE": {
        "label":  "Confidential third-party disclosure",
        "desc":   "Information given in confidence by a third party; anonymous or pseudonymous reports",
        "action": "redact",
    },
    "OTHER_PATIENT_DATA": {
        "label":  "Other patient's data",
        "desc":   "Data belonging to a different patient (misfiled notes, clinic list error, wrong results)",
        "action": "redact",
    },
    "AGENCY_CONFIDENTIAL_INFO": {
        "label":  "Agency / social care report",
        "desc":   "Social worker, school, police or probation report that identifies a third party",
        "action": "redact",
    },
    "INDIRECT_IDENTIFIER": {
        "label":  "Indirect identifier",
        "desc":   "Text that would identify a third party without naming them explicitly",
        "action": "redact",
    },
    # ── Escalate for qualified human decision ─────────────────────────────────
    "CLINICIAN_CONTEXT_AMBIGUOUS": {
        "label":  "Clinician — context ambiguous",
        "desc":   (
            "A clinician name that appears in a non-professional context: named as a patient, "
            "as a complainant, as the subject of an internal complaint or investigation, or "
            "where their role is unclear (e.g. locum/agency staff). "
            "Clinicians named in their ordinary professional capacity are NOT redacted."
        ),
        "action": "escalate",
    },
    "SAFEGUARDING_RISK": {
        "label":  "Safeguarding concern",
        "desc":   "Safeguarding referral, MARAC, CP concern, LAC, MASH referral — requires clinical/IG review",
        "action": "escalate",
    },
    "DOMESTIC_ABUSE_CONTEXT": {
        "label":  "Domestic abuse disclosure",
        "desc":   "Domestic abuse, coercive control, DASH assessment, MARAC referral — escalate",
        "action": "escalate",
    },
    "CHILD_PROTECTION": {
        "label":  "Child protection information",
        "desc":   "CP plan, S47/S17 enquiry, CP conference, LADO — escalate",
        "action": "escalate",
    },
    "SERIOUS_HARM_RISK": {
        "label":  "Serious harm risk",
        "desc":   (
            "Information whose disclosure could cause serious physical or mental harm to the "
            "patient or a third party (DPA 2018 Sch.3 para.5 / s.15). "
            "Includes acute active suicide/self-harm risk, credible violence risk, acute psychotic "
            "risk. Routine historical mental health notes are NOT covered by this exemption."
        ),
        "action": "escalate",
    },
    "SENSITIVE_CLINICAL_OPINION": {
        "label":  "Harmful clinical opinion",
        "desc":   (
            "Clinical opinion that, if disclosed, could cause serious harm or engages a specific "
            "exemption — NOT routine clinical opinion, which is the patient's own data and must "
            "be disclosed. Covers: explicit notes on symptom fabrication / factitious disorder, "
            "notes recording a credible and current risk of violence by the patient, or opinion "
            "that would directly identify and potentially harm a named third party."
        ),
        "action": "escalate",
    },
    "LEGAL_PRIVILEGE": {
        "label":  "Legal / investigation material",
        "desc":   (
            "Material that may attract an exemption under DPA 2018 Sch.3: legal advice, court "
            "reports, expert witness reports, internal disciplinary or complaints investigations "
            "(Sch.3 para.19), management forecasting / planning information (Sch.3 para.6), or "
            "formal negotiation records (Sch.3 para.7). Requires IG / legal review."
        ),
        "action": "escalate",
    },
    "DPA_SCHEDULE3_EXEMPTION": {
        "label":  "DPA 2018 Sch.3 — other exemption",
        "desc":   (
            "Content that may engage a Schedule 3 DPA 2018 exemption not captured elsewhere: "
            "research, statistics or history data (Sch.3 para.8); exam scripts before publication "
            "(Sch.3 para.9); regulatory / supervisory body material (Sch.3 para.10); or data "
            "from a separate data controller whose provenance is unclear in a shared-care or "
            "ICB-held record. Requires IG review to identify the applicable head of exemption."
        ),
        "action": "escalate",
    },
}

SECTION_ORDER = [
    "Clinical Records",
    "Referral Letters",
    "Correspondence",
    "Results and Investigations",
    "Miscellaneous",
]


def _extract_document_date(text: str) -> datetime.date:
    """
    Extract the most relevant date from document text using common NHS/UK formats.
    Searches the first 2 000 characters (where letterhead dates normally appear).
    Returns datetime.date.min if no plausible date is found (sorts to end of section).
    """
    sample = text[:2000]

    _MONTH_MAP = {
        "jan": 1, "january": 1,
        "feb": 2, "february": 2,
        "mar": 3, "march": 3,
        "apr": 4, "april": 4,
        "may": 5,
        "jun": 6, "june": 6,
        "jul": 7, "july": 7,
        "aug": 8, "august": 8,
        "sep": 9, "september": 9,
        "oct": 10, "october": 10,
        "nov": 11, "november": 11,
        "dec": 12, "december": 12,
    }

    _MIN = datetime.date(1990, 1, 1)
    _MAX = datetime.date.today()
    candidates = []

    def _add(y, m, d):
        try:
            dt = datetime.date(int(y), int(m), int(d))
            if _MIN <= dt <= _MAX:
                candidates.append(dt)
        except ValueError:
            pass

    # DD/MM/YYYY  DD-MM-YYYY  DD.MM.YYYY
    for m in re.finditer(r'\b(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})\b', sample):
        _add(m.group(3), m.group(2), m.group(1))

    # YYYY-MM-DD
    for m in re.finditer(r'\b(\d{4})-(\d{2})-(\d{2})\b', sample):
        _add(m.group(1), m.group(2), m.group(3))

    # DD Month YYYY  (e.g. "15 January 2024" or "15 Jan 2024")
    for m in re.finditer(r'\b(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})\b', sample):
        mn = _MONTH_MAP.get(m.group(2).lower())
        if mn:
            _add(m.group(3), mn, m.group(1))

    return max(candidates) if candidates else datetime.date.min

ACCEPTED_FORMATS = ["pdf", "docx", "doc", "tiff", "tif", "rtf", "txt", "zip"]


class _FileWrapper:
    """Wraps raw bytes + filename to behave like a Streamlit UploadedFile."""
    def __init__(self, name: str, data: bytes):
        self.name = name
        self._data = data

    def read(self) -> bytes:
        return self._data


def _detect_patient_name(filename: str, text: str = "") -> str:
    """
    Try to detect the patient's full name from:
      1. NHS EPR filename convention: '…SURNAME, Firstname (Title) NHSnum date.ext'
      2. Common document header patterns: 'Patient: Ms Firstname SURNAME'

    Returns 'Firstname Surname' (title-cased) or empty string if not found.
    Used as a fallback when the operator has not typed the patient name in the sidebar.
    """
    # ── 1. Filename pattern ──────────────────────────────────────────────────
    # Typical: '2022-09-14_hash_Description SURNAME, Firstname (Ms) 1000 …'
    m = re.search(
        r'\b([A-Z]{2,}),\s+([A-Za-z][a-z]+)\s+\((?:Mr|Mrs|Ms|Miss|Dr|Prof)',
        filename,
    )
    if m:
        return f"{m.group(2)} {m.group(1).title()}"

    # ── 2. Document text header ──────────────────────────────────────────────
    if text:
        sample = text[:2000]
        for pat in (
            r'Patient:\s+(?:Mr|Mrs|Ms|Miss|Dr|Prof)\.?\s+([A-Za-z][a-z]+)\s+([A-Z][A-Za-z]+)',
            r'Patient:\s+([A-Za-z][a-z]+)\s+([A-Z]{2,})',
            r'Name:\s+(?:Mr|Mrs|Ms|Miss|Dr|Prof)\.?\s+([A-Za-z][a-z]+)\s+([A-Z][A-Za-z]+)',
        ):
            m = re.search(pat, sample)
            if m:
                return f"{m.group(1)} {m.group(2).title()}"

    return ""


def _detect_guardian_name(text: str) -> str:
    """
    Extract the named parent/guardian from a paediatric record header.
    Returns the full name string (including title such as 'Mrs') as it
    appears before any parenthetical annotation, or '' if not found.
    """
    sample = text[:1500]
    for pat in (
        r'(?:Parent/Guardian|Registered Parent):\s+'
        r'((?:Mr|Mrs|Ms|Miss|Dr|Prof)\.?\s+[A-Za-z][a-z]+\s+[A-Z][A-Za-z]+)',
        r'(?:Parent/Guardian|Registered Parent):\s+'
        r'([A-Za-z][a-z]+\s+[A-Z][A-Za-z]+)',
    ):
        m = re.search(pat, sample)
        if m:
            return m.group(1).strip()
    return ""


def _detect_patient_dob(text: str) -> str:
    """
    Extract the patient's own DOB from the record header (first 1500 chars).
    Returns the date string as it appears (e.g. '27/06/1978') or '' if not found.
    Used to prevent the LLM from accidentally flagging the patient's own DOB
    as third-party data.
    """
    sample = text[:1500]
    for pat in (
        r'DOB:\s+(\d{1,2}/\d{1,2}/\d{4})',
        r'Date of Birth:\s+(\d{1,2}/\d{1,2}/\d{4})',
        r'DOB:\s+(\d{2}\.\d{2}\.\d{4})',
        r'D\.O\.B\.?:\s+(\d{1,2}/\d{1,2}/\d{4})',
    ):
        m = re.search(pat, sample)
        if m:
            return m.group(1).strip()
    return ""


_SUPPORTED_EXTS = {"pdf", "docx", "doc", "tiff", "tif", "rtf", "txt"}


def _expand_zip(name: str, data: bytes) -> list:
    """Return a list of _FileWrapper for every supported file inside a ZIP."""
    result = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for entry in zf.infolist():
                if entry.is_dir():
                    continue
                inner_name = Path(entry.filename).name
                ext = inner_name.rsplit(".", 1)[-1].lower() if "." in inner_name else ""
                if ext in _SUPPORTED_EXTS:
                    result.append(_FileWrapper(inner_name, zf.read(entry.filename)))
    except Exception as exc:
        # Return a sentinel so the caller can surface the error
        result.append(_FileWrapper(f"[ZIP ERROR] {name}", b""))
        result[-1]._zip_error = str(exc)
    return result


def _collect_all_files(uploaded_files, folder_path: str) -> list:
    """
    Combine Streamlit-uploaded files with any files from a folder path.
    Expands ZIPs from both sources.
    Returns a flat list of _FileWrapper / UploadedFile objects.
    """
    raw = list(uploaded_files or [])

    # Load files from folder
    if folder_path:
        fp = Path(folder_path)
        if fp.is_dir():
            for f in sorted(fp.iterdir()):
                if f.is_file():
                    ext = f.suffix.lower().lstrip(".")
                    if ext in _SUPPORTED_EXTS or ext == "zip":
                        raw.append(_FileWrapper(f.name, f.read_bytes()))

    # Expand ZIPs
    flat = []
    for uf in raw:
        name = uf.name
        if name.lower().endswith(".zip"):
            data = uf.read()
            flat.extend(_expand_zip(name, data))
        else:
            flat.append(uf)

    return flat

NHS_BLUE = (0.0,  0.478, 0.784)
WHITE    = (1.0,  1.0,   1.0)
BLACK    = (0.0,  0.0,   0.0)
GREY     = (0.45, 0.45,  0.45)
LT_GREY  = (0.92, 0.92,  0.92)


# ── Logo (resized to 60 px tall for the header) ───────────────────────────────
_LOGO_B64 = ""
_LOGO_PATH = Path(__file__).parent / "logo.jpg"
try:
    if _LOGO_PATH.exists() and TESSERACT_AVAILABLE:   # PIL already imported
        _img = PILImage.open(_LOGO_PATH)
        _ratio = 60 / _img.height
        _img = _img.resize((int(_img.width * _ratio), 60), PILImage.LANCZOS)
        _buf = io.BytesIO()
        _img.save(_buf, format="JPEG", quality=85)
        _LOGO_B64 = "data:image/jpeg;base64," + base64.b64encode(_buf.getvalue()).decode()
    elif _LOGO_PATH.exists():
        # PIL available but tesseract flag is False — PIL still imported if Pillow is installed
        from PIL import Image as _PImg
        _img = _PImg.open(_LOGO_PATH)
        _ratio = 60 / _img.height
        _img = _img.resize((int(_img.width * _ratio), 60), _PImg.LANCZOS)
        _buf = io.BytesIO()
        _img.save(_buf, format="JPEG", quality=85)
        _LOGO_B64 = "data:image/jpeg;base64," + base64.b64encode(_buf.getvalue()).decode()
except Exception:
    _LOGO_B64 = ""


# ── Light-theme CSS ────────────────────────────────────────────────────────────
def _inject_css():
    st.markdown("""
<style>
/* ═══ App background — white/light-grey ═══ */
.stApp { background: #f5f7fa !important; }
.stApp::before { display: none !important; }
.main .block-container { background: transparent; padding-top: 1rem; }
/* ═══ Sidebar ═══ */
[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 2px solid #dde5f0 !important;
    box-shadow: 2px 0 10px rgba(0,94,184,.07) !important;
}
[data-testid="stSidebar"] * {
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}

/* ═══ Headings ═══ */
h1 { color: #0a2040 !important; font-weight: 700 !important; }
h2, h3 { color: #1a3060 !important; font-weight: 600 !important; }
p, li { color: #2c3e5a !important; }
.stCaption, [data-testid="stCaptionContainer"] p { color: #5a7090 !important; }

/* ═══ Buttons — NHS blue ═══ */
/* Must set -webkit-text-fill-color as well as color; Streamlit injects
   -webkit-text-fill-color from textColor in config.toml which otherwise
   wins over `color` in Chrome/WebKit regardless of specificity.          */
.stButton > button,
.stButton > button * {
    background: #ffffff !important;
    border: 1.5px solid #005EB8 !important;
    color: #005EB8 !important;
    -webkit-text-fill-color: #005EB8 !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    transition: all .18s ease !important;
    box-shadow: 0 1px 4px rgba(0,94,184,.10) !important;
}
.stButton > button:hover,
.stButton > button:hover * {
    background: #005EB8 !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    box-shadow: 0 3px 12px rgba(0,94,184,.22) !important;
}
.stButton > button[kind="primary"],
.stButton > button[kind="primary"] * {
    background: #005EB8 !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    border-color: #005EB8 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[kind="primary"]:hover * {
    background: #004a9a !important;
    -webkit-text-fill-color: #ffffff !important;
}
.stDownloadButton > button,
.stDownloadButton > button * {
    background: #1a7a36 !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    border-color: #1a7a36 !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
}
.stDownloadButton > button:hover,
.stDownloadButton > button:hover * {
    background: #155f2b !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ═══ Metrics ═══ */
[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid #dde5f0 !important;
    border-radius: 12px !important;
    padding: 16px 20px !important;
    box-shadow: 0 2px 8px rgba(0,94,184,.07) !important;
}
[data-testid="stMetricValue"] { color: #0a2040 !important; font-weight: 700 !important; }
[data-testid="stMetricLabel"] { color: #4a6080 !important; }

/* ═══ Expanders ═══ */
[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid #dde5f0 !important;
    border-radius: 10px !important;
    margin-bottom: 8px !important;
}
[data-testid="stExpander"]:hover { border-color: #005EB8 !important; }
[data-testid="stExpanderHeader"] { color: #1a3060 !important; font-weight: 500 !important; }
[data-testid="stExpanderDetails"] {
    background: #f9fafc !important;
    border-top: 1px solid #e8edf5 !important;
}

/* ═══ Inputs ═══ */
input, textarea,
[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea,
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stNumberInput > div > div > input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    border: 1px solid #b8cce0 !important;
    border-radius: 7px !important;
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}
input::placeholder, textarea::placeholder {
    color: #8095b0 !important;
    -webkit-text-fill-color: #8095b0 !important;
}
input:focus, textarea:focus,
[data-baseweb="input"]:focus-within,
[data-baseweb="textarea"]:focus-within {
    border-color: #005EB8 !important;
    box-shadow: 0 0 0 2px rgba(0,94,184,.15) !important;
}

/* ═══ Selectbox ═══ */
.stSelectbox [data-baseweb="select"],
.stSelectbox > div > div,
[data-baseweb="select"],
[data-baseweb="select"] > div {
    background: #ffffff !important;
    border-color: #b8cce0 !important;
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}
[data-baseweb="select"] [data-baseweb="single-value"],
[data-baseweb="select"] [data-baseweb="placeholder"],
[data-baseweb="select"] span,
[data-baseweb="select"] p {
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}
[data-baseweb="popover"], [data-baseweb="menu"],
[role="listbox"], ul[data-baseweb="menu"] {
    background: #ffffff !important;
    border: 1px solid #dde5f0 !important;
    border-radius: 8px !important;
    box-shadow: 0 4px 20px rgba(0,0,0,.10) !important;
}
[role="option"], [data-baseweb="menu-item"], li[role="option"] {
    background: #ffffff !important;
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}
[role="option"]:hover, [data-baseweb="menu-item"]:hover {
    background: #e8f0fc !important;
    color: #005EB8 !important;
    -webkit-text-fill-color: #005EB8 !important;
}
[aria-selected="true"][role="option"] {
    background: #005EB8 !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ═══ Alert boxes — let Streamlit's light theme handle colour tinting ═══ */
[data-testid="stAlert"] p,
[data-testid="stAlert"] li,
[data-testid="stAlert"] span,
[data-testid="stAlert"] div,
[data-testid="stAlertContentBlock"] p,
[data-testid="stAlertContentBlock"] li,
[data-testid="stAlertContentBlock"] span {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

/* ═══ Progress bar — NHS blue shimmer ═══ */
[data-testid="stProgressBar"] > div > div {
    background: linear-gradient(90deg, #005EB8, #0081d3, #005EB8) !important;
    background-size: 200% 100% !important;
    animation: sar-shimmer 1.8s linear infinite !important;
    border-radius: 4px !important;
}
@keyframes sar-shimmer { 0%{background-position:200% 0} 100%{background-position:-200% 0} }

/* ═══ Data tables ═══ */
[data-testid="stDataEditor"], [data-testid="stDataFrame"] {
    background: #ffffff !important;
    border: 1px solid #dde5f0 !important;
    border-radius: 8px !important;
}

/* ═══ Code blocks ═══ */
.stCode, code, pre {
    background: #f0f4f8 !important;
    border: 1px solid #dde5f0 !important;
    border-radius: 7px !important;
    color: #1a3060 !important;
}

/* ═══ Containers with border ═══ */
[data-testid="stVerticalBlockBorderWrapper"] {
    background: #ffffff !important;
    border-color: #dde5f0 !important;
    border-radius: 10px !important;
}

/* ═══ Divider ═══ */
hr { border-color: #dde5f0 !important; }

/* ═══ Toggle ═══ */
.stToggle > label > div { background: #c8d8ec !important; }

/* ═══ File uploader ═══ */
[data-testid="stFileUploader"] *,
[data-testid="stFileUploaderDropzone"] *,
[data-testid="stFileUploaderFileList"] *,
[data-testid="stFileUploaderFile"] * {
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: #f5f8ff !important;
    border: 1.5px dashed #005EB8 !important;
    border-radius: 8px !important;
}

/* ═══ Radio / Checkbox ═══ */
.stRadio label, .stRadio span,
.stRadio > div > label span,
[data-testid="stRadio"] label,
[data-testid="stRadio"] span,
.stCheckbox label span,
[data-testid="stCheckbox"] span {
    color: #1a2a4a !important;
    -webkit-text-fill-color: #1a2a4a !important;
}

/* ═══ Scrollbar — NHS blue, clearly visible ═══ */
::-webkit-scrollbar { width: 14px; height: 14px; }
::-webkit-scrollbar-track { background: #e8edf5; border-radius: 7px; }
::-webkit-scrollbar-thumb {
    background: #005EB8;
    border-radius: 7px;
    border: 3px solid #e8edf5;
    background-clip: padding-box;
}
::-webkit-scrollbar-thumb:hover {
    background: #004a9a;
    border: 3px solid #e8edf5;
    background-clip: padding-box;
}
* { scrollbar-width: auto; scrollbar-color: #005EB8 #e8edf5; }

/* ═══ Header card — solid NHS blue ═══ */
.sar-header {
    display: flex; align-items: center; gap: 18px;
    padding: 18px 24px;
    background: #005EB8;
    border-radius: 14px; margin-bottom: 20px;
    box-shadow: 0 4px 20px rgba(0,94,184,.25);
}
.sar-header img { height: 54px; width: auto; border-radius: 8px; }
.sar-header-text { flex: 1; }
.sar-header-text h1 { margin: 0 !important; font-size: 1.55rem !important; font-weight: 700 !important; color: #ffffff !important; line-height: 1.2 !important; }
.sar-header-text p  { margin: 5px 0 0; font-size: .82rem; color: rgba(210,230,255,.85); }

/* ═══ Badges ═══ */
.badge-local {
    display: inline-flex; align-items: center; gap: 5px;
    background: #e6f9ec; border: 1px solid #52c271;
    border-radius: 20px; padding: 3px 11px;
    font-size: .72rem; color: #1a7a36; font-weight: 600; letter-spacing: .3px;
}
.badge-test {
    display: inline-flex; align-items: center; gap: 5px;
    background: #fff8e6; border: 1px solid #f0a030;
    border-radius: 20px; padding: 3px 11px;
    font-size: .72rem; color: #8a5a00; font-weight: 600; letter-spacing: .3px;
}

/* ═══ Disclaimer ═══ */
.sar-disclaimer {
    background: #fff8e6;
    border: 1px solid #f0a030;
    border-radius: 8px; padding: 10px 14px; margin: 10px 0;
    font-size: .76rem; color: #7a4a00; line-height: 1.55;
}
</style>""", unsafe_allow_html=True)


# ── Sound effects (Web Audio API via components.html) ────────────────────────
def _play_sound(sound: str):
    """Inject a zero-height iframe that plays a Web Audio tone sequence."""
    _SOUNDS = {
        "chime": "[[523,.0,.18],[659,.16,.18],[784,.32,.35]]",     # C5-E5-G5
        "fanfare": "[[523,.0,.15],[659,.14,.15],[784,.28,.15],[1047,.42,.45]]",  # C5-E5-G5-C6
        "click": "[[880,.0,.06]]",
    }
    notes = _SOUNDS.get(sound, _SOUNDS["chime"])
    components.html(f"""
<script>
(function(){{
  try {{
    var ctx = new (window.AudioContext||window.webkitAudioContext)();
    var notes = {notes};
    notes.forEach(function(n){{
      var o=ctx.createOscillator(), g=ctx.createGain();
      o.connect(g); g.connect(ctx.destination);
      o.frequency.value=n[0]; o.type='sine';
      g.gain.setValueAtTime(0, ctx.currentTime+n[1]);
      g.gain.linearRampToValueAtTime(0.18, ctx.currentTime+n[1]+0.04);
      g.gain.exponentialRampToValueAtTime(0.001, ctx.currentTime+n[1]+n[2]);
      o.start(ctx.currentTime+n[1]); o.stop(ctx.currentTime+n[1]+n[2]+0.05);
    }});
  }} catch(e) {{}}
}})();
</script>""", height=0)


# =============================================================================
# JSON extraction — robust, handles code fences, preamble, minor errors
# =============================================================================

def _extract_json(raw: str):
    if not raw:
        return None

    # Strategy 1: JSON inside a ```json ... ``` fence (greedy to capture full nested object)
    m = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass

    # Strategy 2: first { ... last }
    if "{" in raw and "}" in raw:
        start = raw.index("{")
        end   = raw.rindex("}") + 1
        candidate = raw[start:end]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            # Strategy 3: auto-fix common LLM JSON mistakes then retry
            fixed = candidate
            fixed = re.sub(r",\s*([}\]])",    r"\1",      fixed)  # trailing commas
            fixed = re.sub(r'(?<!")None(?!")',  '"null"',  fixed)  # Python None
            fixed = re.sub(r'(?<!")True(?!")',  '"true"',  fixed)  # Python True
            fixed = re.sub(r'(?<!")False(?!")', '"false"', fixed)  # Python False
            try:
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass

    return None


# =============================================================================
# Ollama helpers
# =============================================================================

# Models tried in order when sorting the sidebar dropdown.
# Add or change entries to suit the hardware on each deployment.
PREFERRED_MODELS = [
    "qwen2.5:14b",
    "qwen3.5:9b",
    "qwen2.5:7b",
    "qwen2.5:32b",
    "qwen2.5",
    "llama3.1:8b",
    "llama3.1",
    "llama3",
]


def _rank_model(name: str) -> int:
    for i, pref in enumerate(PREFERRED_MODELS):
        if name.startswith(pref):
            return i
    return len(PREFERRED_MODELS)


def check_ollama_connection():
    try:
        resp   = ollama.list()
        models = resp.models if hasattr(resp, "models") else []
        names  = []
        for m in models:
            if hasattr(m, "model"):   names.append(m.model)
            elif isinstance(m, dict): names.append(m.get("name", m.get("model", "")))
        names = sorted([n for n in names if n], key=_rank_model)
        return True, names
    except Exception:
        return False, []


_SAR_SYSTEM = (
    "You are an NHS Information Governance SAR redaction specialist. "
    "You respond with valid JSON only. No preamble, no explanation, no markdown."
)

_SAR_PROMPT_TMPL = """\
You are an NHS Information Governance officer processing a Subject Access Request (SAR).
Analyse ONLY the text between the --- markers below.
Apply UK GDPR / DPA 2018 / ICO guidance and the BMA guidance on access to health records.

━━━ DO NOT FLAG FOR REDACTION ━━━
{patient_line}\
• The patient's own name, DOB, NHS number, address, clinical findings, diagnoses,
  medications and test results — this is their own personal data and MUST be disclosed.
  NOTE: only the patient's OWN DOB is protected. Any date of birth that differs from
  the patient's DOB and belongs to a third party (e.g. a perpetrator, next of kin,
  or misfiled patient) MUST be flagged as THIRD_PARTY_IDENTIFIER or OTHER_PATIENT_DATA.
• Clinical process words used alone — these describe clinical activities and are NOT
  personal identifiers. NEVER flag words such as: prescription, referral, consultation,
  appointment, medication, dosage, treatment, examination, assessment, admission,
  discharge, procedure, operation, review, follow-up, blood test, scan, result.
  Only flag the PERSONAL DATA around them (e.g. a third party's name, their DOB).
• Routine clinical opinion — clinical opinions, assessments and judgements recorded about
  the patient are the patient's own data. Do NOT escalate them unless they meet the
  specific "SENSITIVE_CLINICAL_OPINION" criteria below.
• Clinician names (GP, nurse, consultant, pharmacist, AHP) appearing in their ORDINARY
  PROFESSIONAL CAPACITY — e.g. signing a letter, recording a consultation, ordering a test.
  Exception: escalate as CLINICIAN_CONTEXT_AMBIGUOUS if the clinician is named as a
  patient, as the complainant/subject of a complaint, or in a context where their personal
  data (not their professional act) is being recorded.
• NHS Trust, hospital, GP practice, clinic or department names, and their postal
  addresses — a clinic's building address in a letterhead is NOT personal data.
• Clinician and healthcare staff contact details at their place of work that are
  publicly listed (e.g. a GP surgery's main phone number, a ward switchboard
  number, a generic admin email). DO redact a clinician's personal direct-dial
  number or personal email address if it is clearly individual rather than shared.
• Standard appointment dates, referral acknowledgements, administrative notices.
• Job titles and role descriptions alone (e.g. "SEN coordinator", "class teacher",
  "key worker", "social worker", "named nurse", "care coordinator") — these are NOT
  personal data. Only redact the individual's personal name, not their job title.
• In a paediatric record (patient described as "child"), the named parent or guardian
  listed in the record header (e.g. "Parent/Guardian: Mrs Chloe Green") is the SAR
  requestor acting on the child's behalf — do NOT redact their name ANYWHERE in the
  document, even when it appears again in the body or in correspondence.
• Clinician and healthcare professional names (including abbreviated forms such as
  "Dr M. Robertson", "Dr J. Cole", "Nurse Ward") when appearing in their professional
  capacity — do NOT redact. This applies to ALL registered health professionals:
  GPs, hospital consultants, nurses, pharmacists, physiotherapists, occupational
  therapists, optometrists, dentists, radiographers, and any other AHP or clinical
  specialist — regardless of whether they work at this practice or an external clinic.
  This overrides the abbreviated-name rule below.

━━━ PROPOSE FOR AUTO-REDACTION ━━━
Copy the EXACT tag name. Redact the minimum span — a name or phrase, not a whole sentence.

THIRD_PARTY_IDENTIFIER   — name or identifying detail of any private individual:
                           family member, partner, carer, neighbour, friend, employer,
                           teacher, school contact, or any unnamed member of the public.
                           This INCLUDES their date of birth, phone number, NHS number,
                           address, and any other personal data appearing in structured
                           blocks (e.g. "Perpetrator details:", "Emergency contact:",
                           "Next of kin:") — redact ALL fields in such blocks, not just
                           the name. Create a separate entry for EACH field (name, DOB,
                           phone, address) so each is individually redacted.
                           Device serial numbers for the patient's personal medical devices
                           (insulin pumps, implants, CGM sensors, home monitors) are personal
                           data — flag as THIRD_PARTY_IDENTIFIER.
                           When you flag a PERSONAL email address (firstname.lastname@,
                           initial.surname@, or similar personal format) belonging to a
                           named private individual, ALSO create a separate THIRD_PARTY_IDENTIFIER
                           entry for that person's name (e.g. if you flag
                           "anita.lobo@company.co.uk", also flag "Anita Lobo"; if you flag
                           "s.allen@sleep-centre-personal.com", also flag "Sophie Allen").
                           Do NOT apply this rule to generic role/dept addresses (support@,
                           info@, admin@, victim.support@) — and NEVER use it to flag
                           clinicians in their professional capacity.
                           Abbreviated names (e.g. "C. Murray", "Anna S.", "P. Hall") ARE
                           third-party identifiers when they refer to a NON-CLINICIAN private
                           individual — redact them exactly as written. Do NOT apply this to
                           clinicians or healthcare professionals acting in their professional
                           capacity (e.g. "Dr M. Robertson", "Dr J. Cole" are NOT redacted).
                           Police incident reference numbers, crime reference numbers, and
                           Motor Insurers' Bureau (MIB) claim references (e.g. "MV/2024/B1/04471",
                           "URN 01AZ/12345/23") are THIRD_PARTY_IDENTIFIER — they are linked
                           to a named third party in the police or insurance system and must
                           be redacted.
CONFIDENTIAL_DISCLOSURE  — information given in confidence or anonymously by a third party
                           (ICO guidance: the identity of the third party may be withheld).
                           Specific descriptions of a named or identifiable third party's
                           threatening or abusive behaviour (e.g. "sending threatening
                           messages", "verbal abuse", "threatening text messages") are
                           CONFIDENTIAL_DISCLOSURE — they characterise that private individual
                           and should not be disclosed without review.
OTHER_PATIENT_DATA       — data clearly belonging to a different patient: misfiled notes,
                           wrong-patient test results, clinic lists showing other patients.
                           Redact ALL identifying fields for the other patient including their
                           name, date of birth, NHS number, address, and any other personal
                           identifiers — create a SEPARATE entry for each field.
AGENCY_CONFIDENTIAL_INFO — (a) the name and direct contact details of any social worker,
                           police officer, prison officer, custody officer, housing officer,
                           probation officer, school staff member, university counsellor,
                           external therapist (including NHS therapists in specialist services
                           such as eating disorder, IAPT, psychological therapy, or substance
                           misuse services), support group coordinator, interpreter, solicitor
                           or legal representative, or private/employer-commissioned
                           physiotherapist or occupational health adviser
                           named individually in their professional capacity in a referral,
                           report, or correspondence — they work for a DIFFERENT data
                           controller and their personal work details are not the patient's
                           data to receive;
                           (b) the substantive content of any social work, police, probation,
                           school, or agency report that names or identifies a third party.
                           Do NOT redact the agency or organisation name itself (e.g.
                           'Kent Adult Social Care', 'Warwickshire Children's Services',
                           'Women's Refuge') — only the personal names and direct contact
                           details of named individuals working for those organisations.
                           Always create SEPARATE entries for the name and the phone
                           number — never bundle them. If you find a phone number for an
                           agency professional, you MUST also create a separate entry for
                           their name, and vice versa.
INDIRECT_IDENTIFIER      — text that would identify a private third party without naming
                           them (e.g. "your son at St Peter's Primary", "the neighbour at
                           No. 14", "your partner who works at the council").

━━━ ESCALATE FOR QUALIFIED HUMAN REVIEW — do NOT auto-redact ━━━
These require a clinical or IG professional to make an active decision before any action.

CLINICIAN_CONTEXT_AMBIGUOUS — a clinician name appearing in an ambiguous or non-professional
                              context: named as a patient in this record, named as the subject
                              of or complainant in a formal complaint or investigation, or
                              where their role is unclear (locum/agency with no stated role).
                              IMPORTANT: Documents headed 'Formal Complaint', 'Record of
                              Complaint Received', 'Patient Complaint' or similar MUST have
                              any clinician named as the SUBJECT of the complaint escalated
                              under this tag — even if their name also appears elsewhere in
                              the document in a professional capacity.
SAFEGUARDING_RISK           — safeguarding referrals, MARAC discussions, CP concerns,
                              LAC / MASH referrals. Releasing or withholding requires
                              a qualified decision; neither is automatic.
DOMESTIC_ABUSE_CONTEXT      — domestic abuse or coercive control disclosures, DASH risk
                              assessment results, MARAC referral details.
CHILD_PROTECTION            — the SUBSTANCE of CP referrals: CP plans, Section 47 or
                              Section 17 enquiry details, CP conferences, LADO referral
                              content. Do NOT use this tag for the child's name or DOB —
                              those are THIRD_PARTY_IDENTIFIER (auto-redact). Only the
                              risk assessment content and referral narrative is escalated.
SERIOUS_HARM_RISK           — content that could cause SERIOUS physical or mental harm if
                              disclosed (DPA 2018 Sch.3 para.5). Applies to ACUTE, ACTIVE
                              risk only: credible imminent suicide or self-harm risk,
                              credible current violence risk, acute psychotic risk. Routine
                              or historical mental health notes do NOT qualify.
SENSITIVE_CLINICAL_OPINION  — clinical opinion that, if disclosed, could cause serious harm
                              or identifies a third party harmfully. Specifically:
                              (a) explicit notes on factitious disorder / symptom fabrication;
                              (b) opinion recording a credible and current risk of violence
                              BY the patient; (c) opinion that would directly identify and
                              harm a named third party. Routine clinical opinion, including
                              personality disorder diagnoses, is the patient's own data and
                              must NOT be escalated under this tag.
LEGAL_PRIVILEGE             — legal advice, court reports, expert witness reports, internal
                              disciplinary or complaints investigation records (Sch.3 para.19),
                              management forecasting / planning information (Sch.3 para.6),
                              or formal negotiation records (Sch.3 para.7).
DPA_SCHEDULE3_EXEMPTION     — content that may engage a Sch.3 DPA 2018 exemption not listed
                              above: research/statistics/history data (Sch.3 para.8); exam
                              scripts before publication (Sch.3 para.9); regulatory body
                              material (Sch.3 para.10); or data whose originating data
                              controller is unclear (e.g. shared-care record, ICB-held data).

━━━ OUTPUT RULES ━━━
• "text": copy EXACTLY as it appears — character for character, minimum span only.
• "replacement": for auto-redactions only; use the format [REDACTED - reason].
  Use a plain hyphen (-), not an em-dash or any other character.
• Never include the patient's own name in any "text" field.
• For THIRD_PARTY_IDENTIFIER: if a third party's name appears in MORE THAN ONE FORM in this
  document (e.g. full name "Jane Smith" at first mention, then "Jane" alone in quoted speech),
  create a SEPARATE entry for EACH verbatim form so every occurrence is captured.
  Example: one entry with text "Jane Smith", a second with text "Jane" (if "Jane" appears alone).
• A first name used alone (e.g. "Sandra", "Brian", "Karen") IS a THIRD_PARTY_IDENTIFIER
  if it refers to a private individual — do not skip it just because a surname is absent.
• Named children appearing in safeguarding or CP referrals are THIRD_PARTY_IDENTIFIER —
  auto-redact their name and DOB as separate entries. The CP referral substance is what
  requires CHILD_PROTECTION escalation, not the child's name itself.
  Always capture just the child's name as the minimum span (e.g. text: "Lily"), and
  their approximate DOB as a second separate entry (e.g. text: "2019" or "approximately
  2019") — never bundle the name and description into one long text string.
• Escalation and auto-redaction are MUTUALLY EXCLUSIVE for the SAME span of text.
  However, a SHORTER span within an escalated passage CAN still be proposed for
  auto-redaction — e.g. if you escalate the full sentence "He mentioned his
  brother-in-law David Holmes has continued to send threatening messages" under
  DOMESTIC_ABUSE_CONTEXT, you should ALSO add a CONFIDENTIAL_DISCLOSURE entry for
  the specific phrase "threatening messages" (or similar behavioural description)
  so it is redacted automatically regardless of the human decision on the escalation.

Output this JSON and nothing else:
{{
  "proposed_redactions": [
    {{
      "text": "exact verbatim text from the document",
      "tag": "THIRD_PARTY_IDENTIFIER",
      "reason": "Brief explanation (one sentence)",
      "replacement": "[REDACTED - third-party personal data]",
      "context": "Up to 30 words of surrounding context"
    }}
  ],
  "escalations": [
    {{
      "text": "exact verbatim text",
      "tag": "SAFEGUARDING_RISK",
      "reason": "Brief explanation (one sentence)",
      "context": "Up to 30 words of surrounding context"
    }}
  ]
}}

If nothing requires redaction or escalation return exactly:
{{"proposed_redactions": [], "escalations": []}}

Document excerpt:
---
{chunk}
---"""


_CHUNK_TIMEOUT = 120   # seconds to wait for a single LLM chunk response

# =============================================================================
# Full-anonymisation prompt
# =============================================================================

_ANON_SYSTEM = (
    "You are a medical records anonymisation specialist. "
    "You respond with valid JSON only. No preamble, no explanation, no markdown."
)

_ANON_PROMPT_TMPL = """\
You are anonymising clinical records for external sharing (e.g. MDU, insurer, researcher).
ALL patient and person identifiers must be removed.

CRITICAL: Read the ENTIRE text between the --- markers from top to bottom.
Process EVERY record, EVERY patient, and EVERY person named anywhere in the text.
Do NOT stop after the first patient. Do NOT skip any record.

For each identifier found, output one JSON entry with:
  "tag"         — category (see below)
  "text"        — the EXACT string as it appears in the document
  "replacement" — the label to substitute

CATEGORIES:
  PATIENT_NAME    → [PATIENT NAME]   (patient full name, surname, first name, initials)
  PATIENT_DOB     → [DATE OF BIRTH]  (patient date of birth — use the COMPLETE date
                    string as it appears, e.g. "15 Oct 1985" or "15/10/1985";
                    NEVER flag a month name alone, a year alone, or a day number alone)
  PATIENT_NHS     → [NHS NUMBER]     (NHS number)
  PATIENT_ADDRESS → [ADDRESS]        (home address, street, town — postcodes caught separately)
  PATIENT_PHONE   → [PHONE NUMBER]
  PATIENT_EMAIL   → [EMAIL]
  PATIENT_ID      → [ID NUMBER]      (passport, NI, driving licence, case/policy reference)
  PERSON_NAME     → [NAME]           (anyone else: clinician, relative, carer, colleague, witness)
  PERSON_CONTACT  → [CONTACT DETAILS]

ALWAYS flag: every patient name, DOB, NHS number, address, every other person's full name or
identifiable initials (e.g. "Dr A. Brown", "J. Smith"), any personal phone/email.

NEVER flag: organisation names, job titles alone, clinical content, consultation dates
(appointment dates, prescription dates, referral dates — these are clinical context, NOT
identifiers), generic place names (cities, hospitals without a personal name), month names
alone (e.g. "October", "Oct"), year numbers alone (e.g. "1985"), or single day numbers.

Return ONLY valid JSON — no prose, no markdown:
{{
  "redactions": [
    {{"tag": "PATIENT_NAME", "text": "Jane Smith", "replacement": "[PATIENT NAME]"}},
    {{"tag": "PERSON_NAME",  "text": "Dr A. Brown", "replacement": "[NAME]"}}
  ]
}}

If nothing to redact: {{"redactions": []}}

---
{chunk}
---"""


def _anon_chunk(chunk: str, model: str) -> tuple:
    """Send one chunk to the LLM for full anonymisation. Returns (redactions_list, raw_string).

    Tries with format='json' first (faster, more reliable on qwen/llama).
    Falls back to plain text mode for models that don't support JSON forcing.
    """
    messages = [
        {"role": "system", "content": _ANON_SYSTEM},
        {"role": "user",   "content": _ANON_PROMPT_TMPL.format(chunk=chunk)},
    ]

    def _call(force_json: bool):
        kwargs = dict(
            model=model,
            messages=messages,
            options={"temperature": 0, "num_predict": 4096},
        )
        if force_json:
            kwargs["format"] = "json"
        return ollama.chat(**kwargs)

    raw = ""
    for use_json in (True, False):          # try JSON mode first, plain text as fallback
        try:
            ex     = ThreadPoolExecutor(max_workers=1)
            future = ex.submit(_call, use_json)
            try:
                resp = future.result(timeout=_CHUNK_TIMEOUT)
            except FuturesTimeoutError:
                ex.shutdown(wait=False)
                return [], f"[TIMEOUT] LLM did not respond within {_CHUNK_TIMEOUT}s"
            finally:
                ex.shutdown(wait=False)
            raw = resp["message"]["content"].strip()
        except Exception as exc:
            raw = f"[LLM ERROR] {exc}"
            continue

        parsed = _extract_json(raw)
        if parsed is not None:
            return parsed.get("redactions", []) or [], raw

        if not use_json:
            break   # both modes tried, give up

    return [], f"[PARSE FAILED] {raw[:300]}"


def _normalise_unicode(text: str) -> str:
    """
    Replace Unicode lookalike whitespace and punctuation with their ASCII
    equivalents so LLM output can be matched back against the source text.

    Handles non-breaking spaces (U+00A0), narrow no-break spaces (U+202F),
    non-breaking hyphens (U+2011), en/em dashes (U+2013/U+2014), and other
    common Unicode substitutes found in NHS documents.
    """
    replacements = {
        "\u00a0": " ",   # non-breaking space
        "\u202f": " ",   # narrow no-break space (used in NHS numbers)
        "\u2009": " ",   # thin space
        "\u2007": " ",   # figure space
        "\u2011": "-",   # non-breaking hyphen
        "\u2012": "-",   # figure dash
        "\u2013": "-",   # en dash
        "\u2014": "-",   # em dash
        "\u2018": "'",   # left single quotation mark
        "\u2019": "'",   # right single quotation mark
        "\u201c": '"',   # left double quotation mark
        "\u201d": '"',   # right double quotation mark
        "\u2022": "*",   # bullet
        "\u2010": "-",   # hyphen (Unicode)
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text


def _smart_chunks(text: str, max_chars: int = 2500) -> list[str]:
    """
    Split text at natural NHS record / paragraph boundaries so the LLM
    always sees complete, coherent sections rather than mid-sentence cuts.

    Priority order for split points:
      1. NHS record headers  ("Record N –", "Patient:", "NHS Number:")
      2. Double blank lines  (paragraph / section breaks)
      3. Single newlines     (line breaks)
      4. Hard character cap  (last resort — splits very long lines)

    Each returned chunk is ≤ max_chars characters.  A 50-char overlap from
    the previous chunk is prepended so names at boundaries are never missed.
    """
    # ── Step 1: split at the strongest record-header boundaries ──────────────
    _HEADER_RE = re.compile(
        r'(?m)(?=^(?:Record\s+\d|Patient\s*:|Name\s*:|NHS\s*Number\s*:|'
        r'Date\s+of\s+Birth\s*:|DOB\s*:|Dear\s+(?:Dr|Mr|Mrs|Ms|Miss|Prof)\b))',
        re.IGNORECASE,
    )
    sections = [s for s in _HEADER_RE.split(text) if s.strip()]
    if len(sections) == 1:
        # No record headers — fall back to paragraph splitting
        sections = [s for s in re.split(r'\n{2,}', text) if s.strip()]
    if len(sections) == 1:
        sections = [text]

    # ── Step 2: merge short sections / further split long ones ───────────────
    def _subsplit(blob: str) -> list[str]:
        """Split a blob that is longer than max_chars at the finest boundary."""
        if len(blob) <= max_chars:
            return [blob]
        # Try paragraph breaks first
        paras = [p for p in re.split(r'\n{2,}', blob) if p.strip()]
        if len(paras) > 1:
            out, buf = [], ""
            for p in paras:
                if len(buf) + len(p) + 2 <= max_chars:
                    buf = (buf + "\n\n" + p) if buf else p
                else:
                    if buf:
                        out.append(buf)
                    buf = p
                    if len(buf) > max_chars:
                        # Still too long — hard-split at max_chars
                        while len(buf) > max_chars:
                            out.append(buf[:max_chars])
                            buf = buf[max_chars:]
            if buf:
                out.append(buf)
            return out
        # No paragraph breaks — hard-split
        return [blob[i: i + max_chars] for i in range(0, len(blob), max_chars)]

    raw_chunks: list[str] = []
    buf = ""
    for sec in sections:
        if len(sec) > max_chars:
            if buf:
                raw_chunks.append(buf)
                buf = ""
            raw_chunks.extend(_subsplit(sec))
        elif len(buf) + len(sec) + 2 <= max_chars:
            buf = (buf + "\n\n" + sec) if buf else sec
        else:
            if buf:
                raw_chunks.append(buf)
            buf = sec
    if buf:
        raw_chunks.append(buf)

    # ── Step 3: add 50-char tail overlap so boundary names aren't missed ─────
    OVERLAP = 50
    final: list[str] = []
    for i, ch in enumerate(raw_chunks):
        prefix = raw_chunks[i - 1][-OVERLAP:] if i > 0 else ""
        final.append((prefix + ch) if prefix else ch)

    return final or [text]


# Verification prompt — run over the *already-redacted* text to catch anything missed
_ANON_VERIFY_TMPL = """\
The text below has already been anonymised but may still contain missed identifiers.

Scan every line. List ONLY items that are STILL visible and should be redacted:
  • Any person's name (patient, relative, clinician, colleague, witness)
  • Any phone number or email address

Do NOT flag: [PATIENT NAME], [NAME], [NHS NUMBER], [ADDRESS] or any other
already-redacted placeholder — only flag real, unredacted text.

Return JSON only:
{{"missed": [{{"text": "exact visible text", "replacement": "[NAME]"}}]}}
If nothing missed: {{"missed": []}}

---
{chunk}
---"""


def anonymise_document(text: str, model: str, status_cb=None) -> tuple:
    """
    Fully anonymise a document by removing all personal identifiers.

    Two-pass approach:
      Pass 1 — boundary-aware chunks sent to LLM for full redaction
      Pass 2 — verification pass over the redacted output to catch anything missed

    Returns (anonymised_text, redaction_count, raw_llm_string, llm_failures).
    """
    # Normalise Unicode lookalikes so LLM-returned strings match the source
    text = _normalise_unicode(text)

    chunks = _smart_chunks(text, max_chars=2500)

    all_redactions = []
    all_raw        = []
    llm_failures   = 0
    _p1_start      = time.time()

    for idx, chunk in enumerate(chunks, 1):
        if status_cb:
            status_cb(
                f"🔍 Pass 1 — chunk {idx}/{len(chunks)}"
                + _timing_suffix(_p1_start, idx - 1, len(chunks))
            )
        redactions, raw = _anon_chunk(chunk, model)
        all_raw.append(raw)
        all_redactions.extend(redactions)
        if raw.startswith("[PARSE FAILED]") or raw.startswith("[TIMEOUT]") or raw.startswith("[LLM ERROR]"):
            llm_failures += 1

    # Deduplicate by text value (case-insensitive)
    seen  = set()
    dedup = []
    for r in all_redactions:
        key = (r.get("text") or "").strip().lower()
        if key and key not in seen:
            seen.add(key)
            dedup.append(r)

    # ── Filter: reject overly-broad single-token proposals ───────────────────
    # A model sometimes extracts just "Oct" from "15 Oct 1985", or "1985" alone.
    # Applying those would clobber unrelated text (e.g. "doctor", "since 1985").
    _MONTH_ONLY_RE = re.compile(
        r'^(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)$',
        re.IGNORECASE,
    )
    _YEAR_ONLY_RE  = re.compile(r'^\d{4}$')
    _DAY_ONLY_RE   = re.compile(r'^\d{1,2}(?:st|nd|rd|th)?$', re.IGNORECASE)
    dedup = [
        r for r in dedup
        if not _MONTH_ONLY_RE.match((r.get("text") or "").strip())
        and not _YEAR_ONLY_RE.match((r.get("text") or "").strip())
        and not _DAY_ONLY_RE.match((r.get("text") or "").strip())
    ]

    def _word_boundary_sub(original: str, replacement: str, text: str) -> str:
        """Replace `original` in `text` using word boundaries where applicable.

        Adding \\b prevents 'Oct' matching inside 'doctor', 'Sep' inside
        'September' already matched by a longer entry, etc.
        Only adds \\b at the side where `original` starts/ends with a word char.
        """
        esc = re.escape(original)
        if original[:1:] and (original[0].isalnum() or original[0] == '_'):
            esc = r'\b' + esc
        if original[-1:] and (original[-1].isalnum() or original[-1] == '_'):
            esc = esc + r'\b'
        return re.sub(esc, replacement, text, flags=re.IGNORECASE)

    # Apply redactions — longest first to avoid partial matches
    result = text
    dedup.sort(key=lambda r: len(r.get("text") or ""), reverse=True)
    count = 0
    for r in dedup:
        original    = (r.get("text") or "").strip()
        replacement = r.get("replacement") or "[ANONYMISED]"
        if not original:
            continue
        new_result = _word_boundary_sub(original, replacement, result)
        if new_result != result:
            count += 1
        result = new_result

    # ── Regex fallbacks — catch identifiers the LLM may have missed ──────────
    # NHS numbers: 3 digits, space, 3 digits, space, 4 digits
    _before = result
    result = re.sub(r'\b(\d{3}[ \u00a0\u202f]\d{3}[ \u00a0\u202f]\d{4})\b', '[NHS NUMBER]', result)
    if result != _before:
        count += 1

    # UK postcodes (e.g. LU2 9DY, SW1A 1AA, M3 2PX)
    _before = result
    result = re.sub(
        r'\b([A-Z]{1,2}\d{1,2}[A-Z]?[ \u00a0]?\d[A-Z]{2})\b',
        '[POSTCODE]', result, flags=re.IGNORECASE,
    )
    if result != _before:
        count += 1

    # NI numbers (e.g. AB 12 34 56 C)
    _before = result
    result = re.sub(
        r'\b([A-CEGHJ-PR-TW-Z]{2}[ ]?\d{2}[ ]?\d{2}[ ]?\d{2}[ ]?[A-D])\b',
        '[NI NUMBER]', result, flags=re.IGNORECASE,
    )
    if result != _before:
        count += 1

    # DOB dates — only when explicitly labelled "DOB:" to avoid redacting
    # clinical dates (admission dates, review dates, etc.)
    _MONTHS = (
        r'Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?'
    )
    _before = result
    result = re.sub(
        rf'(?i)(DOB\s*:?\s*)(\d{{1,2}}[-/ ](?:{_MONTHS})[-/ ]\d{{2,4}}|\d{{1,2}}[/-]\d{{1,2}}[/-]\d{{2,4}})',
        r'\1[DATE OF BIRTH]',
        result,
    )
    if result != _before:
        count += 1

    # ── Pass 2: verification — scan redacted text for anything still missed ──
    verify_chunks = _smart_chunks(result, max_chars=2500)
    _p2_start = time.time()
    for idx, vchunk in enumerate(verify_chunks, 1):
        if status_cb:
            status_cb(
                f"🔎 Pass 2 — verifying chunk {idx}/{len(verify_chunks)}"
                + _timing_suffix(_p2_start, idx - 1, len(verify_chunks))
            )

        def _vcall(c=vchunk):
            return ollama.chat(
                model=model,
                messages=[
                    {"role": "system", "content": _ANON_SYSTEM},
                    {"role": "user",   "content": _ANON_VERIFY_TMPL.format(chunk=c)},
                ],
                format="json",
                options={"temperature": 0, "num_predict": 2048},
            )
        try:
            ex     = ThreadPoolExecutor(max_workers=1)
            future = ex.submit(_vcall)
            try:
                vresp = future.result(timeout=_CHUNK_TIMEOUT)
            except FuturesTimeoutError:
                ex.shutdown(wait=False)
                continue
            finally:
                ex.shutdown(wait=False)
            vraw    = vresp["message"]["content"].strip()
            vparsed = _extract_json(vraw)
            if vparsed:
                for item in vparsed.get("missed", []):
                    original    = (item.get("text") or "").strip()
                    replacement = item.get("replacement") or "[NAME]"
                    if not original or "[" in original:   # skip placeholders
                        continue
                    new_result = _word_boundary_sub(original, replacement, result)
                    if new_result != result:
                        count += 1
                    result = new_result
        except Exception:
            pass

    return result, count, "\n\n---chunk---\n\n".join(all_raw), llm_failures


def _analyse_chunk(chunk: str, model: str, patient_line: str, extra_instructions: str = "") -> tuple:
    """Send one chunk to the LLM. Returns (result_dict, raw_string)."""
    user_msg = _SAR_PROMPT_TMPL.format(patient_line=patient_line, chunk=chunk)
    if extra_instructions:
        user_msg += f"\n\nADDITIONAL INSTRUCTIONS FOR THIS SESSION ONLY:\n{extra_instructions}"

    def _call():
        return ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content": _SAR_SYSTEM},
                {"role": "user",   "content": user_msg},
            ],
            format="json",                              # forces valid JSON output for any model
            options={"temperature": 0,
                     "num_predict": 1024},              # cap output — SAR JSON rarely exceeds ~800 tokens
        )

    try:
        ex     = ThreadPoolExecutor(max_workers=1)
        future = ex.submit(_call)
        try:
            resp = future.result(timeout=_CHUNK_TIMEOUT)
        except FuturesTimeoutError:
            ex.shutdown(wait=False)   # don't block — let the stalled thread die on its own
            return (
                {"proposed_redactions": [], "escalations": [], "parse_ok": False},
                f"[TIMEOUT] LLM did not respond within {_CHUNK_TIMEOUT}s",
            )
        finally:
            ex.shutdown(wait=False)
        raw = resp["message"]["content"].strip()
    except Exception as exc:
        return {"proposed_redactions": [], "escalations": [], "parse_ok": False}, f"[LLM ERROR] {exc}"

    parsed = _extract_json(raw)
    if parsed is None:
        return {"proposed_redactions": [], "escalations": [], "parse_ok": False}, raw

    return {
        "proposed_redactions": parsed.get("proposed_redactions", []) or [],
        "escalations":         parsed.get("escalations", [])         or [],
        "parse_ok":            True,
    }, raw


def llm_analyse_document(
    text: str,
    model: str,
    patient_name: str = "",
    status_cb=None,
    extra_redactions: str = "",
    custom_instructions: str = "",
    patient_dob: str = "",
    patient_nhs: str = "",
    patient_address: str = "",
) -> tuple:
    """
    Analyse document text for SAR redactions.
    Splits long documents into overlapping chunks so the whole document is covered.
    Returns (result_dict, raw_llm_string).

    status_cb:           optional callable(message: str) for live progress updates.
    extra_redactions:    newline/comma-separated extra terms to always redact this session.
    custom_instructions: free-text extra prompt instructions appended this session.
    """
    # Build patient_line — tells the LLM the confirmed patient details so it
    # never accidentally redacts the patient's own data.
    _pdet_parts = []
    if patient_name.strip():
        _pdet_parts.append(f"name: {patient_name.strip()}")
    if patient_dob.strip():
        _pdet_parts.append(f"date of birth: {patient_dob.strip()}")
    if patient_nhs.strip():
        _pdet_parts.append(f"NHS number: {patient_nhs.strip()}")
    if patient_address.strip():
        _pdet_parts.append(f"address: {patient_address.strip()}")
    patient_line = ""
    if _pdet_parts:
        patient_line = (
            f"- The patient's own confirmed details are — {'; '.join(_pdet_parts)}.\n"
            "  NEVER flag ANY of these details, or any part of them, for redaction.\n"
            "  They are the patient's own personal data and MUST be disclosed.\n"
        )

    # Build session-specific addendum
    extra_parts = []
    if extra_redactions.strip():
        terms = [t.strip() for t in re.split(r"[,\n]+", extra_redactions) if t.strip()]
        if terms:
            quoted = ", ".join(f'"{t}"' for t in terms)
            extra_parts.append(
                f"EXTRA TERMS TO REDACT (always flag these regardless of other rules): {quoted}\n"
                "Tag each as THIRD_PARTY_IDENTIFIER unless a more specific tag clearly applies."
            )
    if custom_instructions.strip():
        extra_parts.append(custom_instructions.strip())
    extra_str = "\n\n".join(extra_parts)

    CHUNK      = 6000   # characters per chunk (~1500 words, ~2-3 GP pages)
    STRIDE     = 5500   # overlap of 500 chars catches phrases that straddle a boundary
    MAX_CHUNKS = 8      # analyse up to ~48 000 chars (≈ 12–15 pages)

    chunks = []
    pos = 0
    while pos < len(text) and len(chunks) < MAX_CHUNKS:
        chunks.append(text[pos: pos + CHUNK])
        pos += STRIDE

    all_proposed, all_escalations, all_raw = [], [], []
    parse_ok = True

    _sar_analyse_start = time.time()
    for idx, chunk in enumerate(chunks, 1):
        if status_cb:
            status_cb(
                f"🤖 Analysing chunk {idx}/{len(chunks)} "
                f"(~{len(chunk):,} chars, up to {_CHUNK_TIMEOUT}s each)"
                + _timing_suffix(_sar_analyse_start, idx - 1, len(chunks))
            )
        result, raw = _analyse_chunk(chunk, model, patient_line, extra_str)
        all_raw.append(raw)
        if not result.get("parse_ok"):
            parse_ok = False
        all_proposed.extend(result.get("proposed_redactions", []))
        all_escalations.extend(result.get("escalations", []))

    # ── Post-processing: mutual exclusivity of escalation and auto-redaction ──
    # The LLM sometimes places the same text in both proposed_redactions and
    # escalations, or uses an escalation-only tag in proposed_redactions.
    # In both cases the item should be reviewed by a human, not auto-redacted.
    _escalate_tags = {tag for tag, info in REDACTION_TAGS.items()
                      if info.get("action") == "escalate"}
    _esc_texts = {(e.get("text") or "").strip().lower() for e in all_escalations}
    all_proposed = [
        p for p in all_proposed
        if p.get("tag", "") not in _escalate_tags          # wrong tag for auto-redact
        and (p.get("text") or "").strip().lower() not in _esc_texts  # also in escalations
    ]

    # ── Post-processing: fix empty replacements ──────────────────────────────
    # The LLM sometimes returns auto-redact items (CONFIDENTIAL_DISCLOSURE etc.)
    # with an empty replacement string. Fill in the canonical default so they
    # are actually redacted in the output.
    _DEFAULT_REPLACEMENTS = {
        "THIRD_PARTY_IDENTIFIER":  "[REDACTED - third-party personal data]",
        "CONFIDENTIAL_DISCLOSURE": "[REDACTED - confidential third-party information]",
        "OTHER_PATIENT_DATA":      "[REDACTED - other patient's data]",
        "AGENCY_CONFIDENTIAL_INFO":"[REDACTED - agency confidential information]",
        "INDIRECT_IDENTIFIER":     "[REDACTED - indirect identifier]",
    }
    for item in all_proposed:
        if not (item.get("replacement") or "").strip():
            tag = item.get("tag", "")
            item["replacement"] = _DEFAULT_REPLACEMENTS.get(tag, "[REDACTED]")

    # ── Post-processing: extract concrete identifiers from escalated passages ──
    # When the LLM escalates a whole passage, concrete data items embedded within it
    # (email addresses, police/case reference numbers, phone numbers) should still be
    # auto-redacted so they are not disclosed even if the reviewer decides to release
    # the rest of the escalation context.
    # Also extract abbreviated names of agency workers (e.g. "P. Hall") from passages.
    _existing_proposed_lower = {(p.get("text") or "").strip().lower() for p in all_proposed}
    _EMAIL_RE   = re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')
    _PHONE_RE   = re.compile(r'\b(\d{5}\s\d{6}|\d{4}\s\d{3}\s\d{4}|\d{11}|\+44[\s\d]{10,13})\b')
    _REF_RE     = re.compile(r'\b[A-Z]{2,}[/\-]\d{4}[/\-][A-Z0-9]+[/\-]\d+\b')   # e.g. PC/2024/BR/004421
    # Abbreviated names: "P. Hall", "Dr Wood", "Mr Smith" etc.
    _ABBR_NAME_RE = re.compile(r'\b([A-Z]\.?\s+[A-Z][a-z]{2,})\b')
    for esc in all_escalations:
        esc_text = (esc.get("text") or "").strip()
        for pat, tag, repl in (
            (_EMAIL_RE,     "THIRD_PARTY_IDENTIFIER", "[REDACTED - third-party personal data]"),
            (_PHONE_RE,     "THIRD_PARTY_IDENTIFIER", "[REDACTED - third-party personal data]"),
            (_REF_RE,       "THIRD_PARTY_IDENTIFIER", "[REDACTED - third-party personal data]"),
            (_ABBR_NAME_RE, "AGENCY_CONFIDENTIAL_INFO", "[REDACTED - agency confidential information]"),
        ):
            for m in pat.finditer(esc_text):
                candidate = m.group(0).strip() if pat is not _ABBR_NAME_RE else m.group(1).strip()
                # For abbreviated names, skip clinician titles (Dr, Prof etc.) — they're staff
                # and skip names matching the patient's own name tokens
                if tag == "AGENCY_CONFIDENTIAL_INFO":
                    # "Dr Wood" — skip if this is the treating clinician already in the record
                    # header. Simple check: if the name appears in the document header (first 500
                    # chars) preceded by "Clinician:" or similar, skip it.
                    _header = text[:600].lower()
                    if candidate.split()[-1].lower() in _header:
                        # surname appears in header — likely the treating GP, skip
                        continue
                if candidate.lower() not in _existing_proposed_lower:
                    all_proposed.append({
                        "text":        candidate,
                        "tag":         tag,
                        "reason":      "Identifier/name extracted from escalated passage (auto-redact regardless of escalation decision)",
                        "replacement": repl,
                        "context":     esc_text[:80],
                        "approved":    True,
                    })
                    _existing_proposed_lower.add(candidate.lower())

    # ── Post-processing: family-member name extraction ────────────────────────
    # Catch children/relatives mentioned as "daughter Emily", "son James" etc.
    # The LLM sometimes escalates the surrounding context without separately
    # flagging the family member's first name as THIRD_PARTY_IDENTIFIER.
    _FAMILY_PATTERN = re.compile(
        r'\b(?:daughter|son|sister|brother|mother|father|wife|husband|partner|'
        r'fianc[eé]e?|sibling|niece|nephew|granddaughter|grandson)\s+'
        r'(?:named\s+)?([A-Z][a-z]{1,})\b'
    )
    for fm in _FAMILY_PATTERN.finditer(text):
        name = fm.group(1)
        # Skip if it's a patient name token
        _pn_toks = {t.lower() for t in patient_name.split() if len(t) >= 3}
        if name.lower() in _pn_toks:
            continue
        if name.lower() not in _existing_proposed_lower:
            all_proposed.append({
                "text":        name,
                "tag":         "THIRD_PARTY_IDENTIFIER",
                "reason":      f"Family member's first name (deterministic extraction)",
                "replacement": "[REDACTED - third-party personal data]",
                "context":     text[max(0, fm.start()-20):fm.end()+20],
                "approved":    True,
            })
            _existing_proposed_lower.add(name.lower())

    # ── Post-processing: remove clinician-only names from proposed_redactions ──
    # The LLM sometimes flags a clinician name (e.g. "Elena Morris") as a
    # THIRD_PARTY_IDENTIFIER when she appears in professional correspondence.
    # ── Post-processing: remove clinician-only names ──────────────────────────
    # Suppress THIRD_PARTY_IDENTIFIER redactions for registered health
    # professionals appearing in their professional capacity.
    # Three guards — any ONE matching all occurrences is sufficient to suppress:
    #   (a) Name itself starts with "Dr" or "Prof" (e.g. "Dr M. Robertson").
    #   (b) Leading "Dr"/"Prof" prefix within 8 chars BEFORE each occurrence.
    #   (c) Trailing ", Consultant [Specialty]" or similar title within 60
    #       chars AFTER each occurrence (e.g. "Frank Miller, Consultant Optometrist").
    _CLINICIAN_TITLE_RE = re.compile(
        r'\b(?:Dr|Prof(?:essor)?|Sister|Staff\s+Nurse|Nurse|HCA|Midwife|Paramedic|'
        r'Physio(?:therapist)?|Pharmacist|Optometrist|Radiographer|OT)\s+',
        re.IGNORECASE,
    )
    _CLINICIAN_NAME_START_RE = re.compile(r'^(?:Dr|Prof(?:essor)?)\b', re.IGNORECASE)
    _CLINICIAN_TRAILING_RE = re.compile(
        r',?\s*(?:Consultant|Senior\s+Consultant|Lead\s+Consultant|'
        r'Specialist|Principal|Registrar|Optometrist|Ophthalmologist|'
        r'Dentist|Radiographer|Pharmacist|Physiotherapist|Occupational\s+Therapist|'
        r'Speech\s+(?:and\s+Language\s+)?Therapist|Psychologist|Psychiatrist|'
        r'Podiatrist|Dietitian|Dietician|Paramedic|Midwife|'
        r'Community\s+Nurse|District\s+Nurse|Practice\s+Nurse|'
        r'Nurse\s+Practitioner|Advanced\s+(?:Nurse\s+)?Practitioner|'
        r'Clinical\s+Nurse\s+Specialist|'
        r'GP|General\s+Practitioner|Surgeon|Physician|Paediatrician|'
        r'Cardiologist|Neurologist|Oncologist|Dermatologist|Radiologist|'
        r'Gynaecologist|Obstetrician|Urologist|Gastroenterologist|'
        r'Anaesthetist|Endocrinologist|Rheumatologist|Haematologist|'
        r'(?:MB|BM|BCh|MBBS|MBChB|MRCGP|MRCP|FRCP|FRCS|FRCOphth|'
        r'FRCR|FRCPsych|MRCPsych|DCH|DRCOG|DFFP|DPH|PhD|MD))\b',
        re.IGNORECASE,
    )
    filtered_proposed = []
    for item in all_proposed:
        if item.get("tag") == "THIRD_PARTY_IDENTIFIER":
            name = (item.get("text") or "").strip()
            if " " in name:
                # Guard (a): name itself begins with Dr/Prof
                if _CLINICIAN_NAME_START_RE.match(name):
                    continue
                occurrences = list(re.finditer(
                    r'(?<!\w)' + re.escape(name) + r'(?!\w)', text, re.IGNORECASE
                ))
                # Guards (b) & (c): context around every occurrence
                if occurrences and all(
                    _CLINICIAN_TITLE_RE.search(text[max(0, m.start() - 22): m.start()])
                    or _CLINICIAN_TRAILING_RE.match(text[m.end(): m.end() + 60])
                    for m in occurrences
                ):
                    continue
        filtered_proposed.append(item)
    all_proposed = filtered_proposed

    # ── Role-title filter ─────────────────────────────────────────────────────
    # Remove THIRD_PARTY_IDENTIFIER items that look like role/job titles rather
    # than person names (e.g. "SEN coordinator", "care manager"). The DO NOT FLAG
    # section of the prompt instructs the LLM not to flag these, but it
    # occasionally does so — this filter is the code-level safety net.
    _ROLE_WORDS = {
        "coordinator", "worker", "officer", "manager", "director", "advisor",
        "adviser", "therapist", "counsellor", "nurse", "doctor", "consultant",
        "specialist", "assistant", "support", "teacher", "carer", "warden",
        "liaison", "lead", "head", "deputy", "supervisor", "practitioner",
    }
    all_proposed = [
        p for p in all_proposed
        if not (
            p.get("tag") == "THIRD_PARTY_IDENTIFIER"
            and any(
                word.lower().strip(".,;:") in _ROLE_WORDS
                for word in (p.get("text") or "").split()
                if word[:1].islower()   # only consider lowercase-starting words
            )
        )
    ]

    # ── Institutional-text filter ────────────────────────────────────────────
    # Remove proposed redactions whose text is clearly an organisation/agency
    # name rather than a person name. This catches cases where the LLM ignores
    # the DO NOT REDACT instruction for agency names (e.g. "Kent Adult Social",
    # "Suffolk County Council", "Bluebird Care Ltd").
    _INST_FILTER_WORDS = {
        "adult", "social", "care", "health", "mental", "children", "young",
        "services", "service", "authority", "council", "trust", "nhs",
        "royal", "hospital", "infirmary", "refuge", "centre", "center",
        "community", "primary", "secondary", "support", "unit",
        "foundation", "association", "police", "probation", "housing",
        # Local government / geographic body words
        "county", "borough", "district", "city", "parish", "metropolitan",
        # Org-type suffixes (even 1 institutional word + suffix = org)
        "limited", "ltd", "llp", "plc", "inc",
    }
    _ORG_SUFFIX_WORDS = {"limited", "ltd", "llp", "plc", "inc"}

    def _looks_institutional(t: str) -> bool:
        words = [w.rstrip('.,') for w in t.lower().split()]
        # Org suffix alone lowers the threshold: 1 other institutional word suffices
        if any(w in _ORG_SUFFIX_WORDS for w in words):
            return any(w in _INST_FILTER_WORDS for w in words)
        return sum(1 for w in words if w in _INST_FILTER_WORDS) >= 2

    # Also remove addresses that appear to belong to a clinic/hospital
    # (e.g. a GP surgery letterhead address rather than someone's home).
    _ADDR_WORDS = {
        "road", "street", "avenue", "lane", "drive", "close", "way", "place",
        "court", "grove", "terrace", "crescent", "square", "buildings",
        "house", "centre", "center", "walk", "gardens", "yard",
    }
    _ORG_CONTEXT_WORDS = {
        "trust", "nhs", "hospital", "surgery", "clinic", "practice",
        "infirmary", "royal", "health", "medical", "centre", "center",
        "pharmacy", "dispensary", "ward", "unit",
    }

    def _looks_clinic_address(t: str) -> bool:
        """Return True if t looks like a clinic/hospital address (not personal)."""
        words = set(re.sub(r'[^\w\s]', ' ', t.lower()).split())
        if not any(w in _ADDR_WORDS for w in words):
            return False
        # Check for postcode-like suffix (crude but effective)
        has_postcode = bool(re.search(
            r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2}\b', t, re.IGNORECASE
        ))
        if has_postcode:
            # Look for org context near this address in the document text
            addr_start = text.lower().find(t.lower()[:25])
            if addr_start != -1:
                ctx = text[max(0, addr_start - 300): addr_start + 300].lower()
                if any(w in ctx for w in _ORG_CONTEXT_WORDS):
                    return True
        return False

    all_proposed = [
        p for p in all_proposed
        if not _looks_institutional(p.get("text", ""))
        and not _looks_clinic_address(p.get("text", ""))
    ]

    # ── Guardian name filter ─────────────────────────────────────────────────
    # In paediatric records the registered parent/guardian must not be redacted.
    # Bidirectional check handles both "Mrs Laura Sanders" and "Laura Sanders"
    # (LLM may omit the title prefix when generating the proposed text).
    _guardian_name = _detect_guardian_name(text)
    if _guardian_name:
        _gn_lower = _guardian_name.strip().lower()
        all_proposed = [
            p for p in all_proposed
            if not (
                _gn_lower in (p.get("text") or "").strip().lower()
                or (p.get("text") or "").strip().lower() in _gn_lower
            )
        ]

    # ── Patient DOB filter ───────────────────────────────────────────────────
    # The patient's own DOB is their data, not a third-party identifier —
    # it must never appear in the redaction list.
    # (a) Remove any item explicitly tagged PATIENT_DOB.
    # (b) Remove any item whose text exactly matches the DOB extracted from the
    #     record header (catches cases where the LLM mislabels it).
    all_proposed = [p for p in all_proposed if p.get("tag") != "PATIENT_DOB"]
    _patient_dob = _detect_patient_dob(text)
    if _patient_dob:
        all_proposed = [
            p for p in all_proposed
            if (p.get("text") or "").strip() != _patient_dob
        ]

    # ── Police / incident reference post-processor ───────────────────────────
    # The LLM sometimes misses police incident reference numbers even when
    # explicitly prompted. This regex scans the text for reference-number
    # patterns that appear near police/legal context keywords.
    _POLICE_REF_RE = re.compile(
        r'\b([A-Z]{1,4}/\d{4}/[A-Z0-9]{1,5}/\d{3,6})\b'
    )
    _POLICE_CONTEXT_RE = re.compile(
        r'(?i)(?:police|incident|crime|MIB|motor insur|reference|URN|log\s*number)',
    )
    _existing_texts = {(p.get("text") or "").strip() for p in all_proposed}
    for m in _POLICE_REF_RE.finditer(text):
        ref = m.group(1)
        if ref in _existing_texts:
            continue
        window_start = max(0, m.start() - 100)
        window_end   = min(len(text), m.end() + 100)
        window = text[window_start:window_end]
        if _POLICE_CONTEXT_RE.search(window):
            all_proposed.append({
                "text":        ref,
                "tag":         "THIRD_PARTY_IDENTIFIER",
                "reason":      "Police/incident reference number linked to a third party.",
                "replacement": "[REDACTED - third-party personal data]",
            })
            _existing_texts.add(ref)

    # Expand multi-word name redactions to catch first-name-only mentions.
    # Pass patient_name so the expander never creates a redaction target that
    # matches the patient's own name parts (e.g. a shared family surname).
    all_proposed = _expand_name_redactions(all_proposed, text, patient_name)
    all_proposed = _expand_agency_contacts(all_proposed, text, patient_name)
    all_proposed = _expand_agency_professionals(all_proposed, text, patient_name)

    # ── Post-processing: confirmed patient details guard ─────────────────────
    # Remove any proposal whose text matches a confirmed patient detail (name
    # tokens, DOB string, NHS number, or address fragment).  This is the final
    # safety net against the LLM accidentally flagging the patient's own data.
    _protected_exact   = set()   # exact-match strings (lower)
    _protected_tokens  = set()   # individual tokens that must not be sole-redacted
    for _pv in (patient_name, patient_dob, patient_nhs, patient_address):
        _pv = (_pv or "").strip()
        if _pv:
            _protected_exact.add(_pv.lower())
            # Also protect each space-separated token of length >= 3
            for _tok in _pv.split():
                if len(_tok) >= 3:
                    _protected_tokens.add(_tok.lower().strip(".,;:"))
    if _protected_exact or _protected_tokens:
        _pat_nhs_digits = re.sub(r"\D", "", patient_nhs)  # digits only for loose match
        def _is_patient_own(item):
            _t = (item.get("text") or "").strip()
            _tl = _t.lower()
            if _tl in _protected_exact:
                return True
            # NHS number: compare digits only
            if _pat_nhs_digits and re.sub(r"\D", "", _t) == _pat_nhs_digits:
                return True
            # Short single-token match (e.g. patient first name proposed alone)
            if _tl in _protected_tokens and len(_t.split()) == 1:
                return True
            return False
        all_proposed = [p for p in all_proposed if not _is_patient_own(p)]

    # ── Post-processing: clinical vocabulary safelist ─────────────────────────
    # Remove any proposal whose ENTIRE text is a common clinical process word.
    # These describe clinical activities, not personal identifiers, and should
    # never be redacted regardless of what the LLM returns.
    _CLINICAL_NEVER_REDACT = {
        "prescription", "prescriptions", "repeat prescription",
        "repeat prescriptions", "prescription note", "prescription request",
        "medication", "medications", "medicine", "medicines",
        "diagnosis", "diagnoses", "clinical diagnosis",
        "treatment", "treatments", "treatment plan",
        "referral", "referrals", "e-referral", "onward referral",
        "consultation", "consultations",
        "appointment", "appointments",
        "examination", "examinations",
        "assessment", "assessments",
        "investigation", "investigations",
        "review", "reviews", "medication review",
        "admission", "admissions", "discharge", "discharges",
        "tablet", "tablets", "capsule", "capsules",
        "injection", "injections", "dose", "dosage",
        "inhaler", "inhalers", "cream", "ointment", "spray",
        "blood test", "blood tests",
        "scan", "x-ray", "xray", "mri", "ct scan", "ultrasound",
        "operation", "procedure", "procedures", "surgery",
        "follow up", "follow-up", "outpatient", "inpatient",
        "clinic", "ward", "unit",
        "blood pressure", "bp", "heart rate", "pulse",
        "weight", "height", "bmi",
        "cholesterol", "glucose", "hba1c",
        "test results", "results",
    }
    all_proposed = [
        p for p in all_proposed
        if (p.get("text") or "").strip().lower() not in _CLINICAL_NEVER_REDACT
    ]

    # ── Post-processing: deterministic safeguarding keyword scanner ───────────
    # The LLM can miss safeguarding content when it is incidentally mentioned
    # or buried in long text. This scanner provides a guaranteed safety net:
    # any recognised safeguarding keyword triggers a SAFEGUARDING_RISK escalation
    # regardless of what the LLM returned, so the reviewer is always alerted.
    _SG_SCAN_RE = re.compile(
        r'\b(?:'
        r'safeguarding|safe\s*guarding|'
        r'MARAC|MASH\b|IDVA|IDASA|DASH\b|LADO\b|SOVA\b|PREVENT\b|'
        r'child\s+protection(?:\s+(?:plan|conference|register|referral))?|'
        r'CP\s+(?:plan|conf(?:erence)?|register)|'
        r'section\s+(?:17|47)|s\s*\.?\s*4[57]\b|'
        r'LAC\b|CIN\b|CAF\b|early\s+help|'
        r'domestic\s+(?:abuse|violence)|'
        r'adult\s+safeguarding|MAPPA\b|PPANI\b|'
        r'FGM\b|forced\s+marriage|honour.based\s+(?:abuse|violence)|'
        r'non.accidental\s+(?:injury|harm)|NAI\b|'
        r'vulnerable\s+(?:adult|child(?:ren)?)|'
        r'child(?:ren)?\s+in\s+need|'
        r'at.risk\s+register|risk\s+of\s+significant\s+harm|'
        r'social\s+care\s+(?:plan|referral|assessment)|'
        r'multi.agency\s+(?:risk|safeguarding|meeting)|'
        r'referral\s+to\s+(?:social\s+care|children(?:\'?s)?\s+services|adult\s+services)'
        r')',
        re.IGNORECASE,
    )
    _existing_esc_lower = {(e.get("text") or "").strip().lower() for e in all_escalations}
    for _sg_m in _SG_SCAN_RE.finditer(text):
        _sg_kw = _sg_m.group(0).strip()
        if _sg_kw.lower() in _existing_esc_lower:
            continue
        _sg_ctx = text[max(0, _sg_m.start() - 150): min(len(text), _sg_m.end() + 150)]
        all_escalations.append({
            "text":    _sg_kw,
            "tag":     "SAFEGUARDING_RISK",
            "reason":  (
                f"Safeguarding keyword '{_sg_kw}' detected by automatic scan — "
                "requires clinical/IG review before release"
            ),
            "context": _sg_ctx,
        })
        _existing_esc_lower.add(_sg_kw.lower())

    return {
        "proposed_redactions": all_proposed,
        "escalations":         all_escalations,
        "parse_ok":            parse_ok,
        "chunks_analysed":     len(chunks),
        "chars_total":         len(text),
    }, f"\n\n--- CHUNK BREAK ---\n\n".join(all_raw)


# ── Variant-normalisation for dedup ──────────────────────────────────────────
_DEDUP_TITLE_RE = re.compile(
    r'\b(?:Mr|Mrs|Ms|Miss|Dr|Prof(?:essor)?|Rev|Sir|Lord|Lady|Mx|Cllr|Councillor)\b\.?',
    re.IGNORECASE,
)
_DEDUP_MONTH_MAP = {
    'january':'01','february':'02','march':'03','april':'04','may':'05','june':'06',
    'july':'07','august':'08','september':'09','october':'10','november':'11','december':'12',
    'jan':'01','feb':'02','mar':'03','apr':'04','jun':'06','jul':'07',
    'aug':'08','sep':'09','oct':'10','nov':'11','dec':'12',
}

def _normalise_date_str(text: str):
    """Try to parse text as a date; return canonical DDMMYYYY string or None."""
    t = text.lower().strip()
    m = re.match(r'^(\d{1,2})[/\-\. ](\d{1,2})[/\-\. ](\d{2,4})$', t)
    if m:
        d, mo, y = m.groups()
        if len(y) == 2:
            y = ('19' if int(y) > 30 else '20') + y
        try:
            return f"{int(d):02d}{int(mo):02d}{y}"
        except ValueError:
            return None
    m = re.match(r'^(\d{1,2})\s+([a-z]{3,9})\s+(\d{2,4})$', t)
    if not m:
        m = re.match(r'^([a-z]{3,9})\s+(\d{1,2})\s+(\d{2,4})$', t)
        if m:
            mo_s, d_s, y_s = m.group(1), m.group(2), m.group(3)
        else:
            return None
    else:
        d_s, mo_s, y_s = m.group(1), m.group(2), m.group(3)
    mo_num = _DEDUP_MONTH_MAP.get(mo_s[:3].lower())
    if not mo_num:
        return None
    if len(y_s) == 2:
        y_s = ('19' if int(y_s) > 30 else '20') + y_s
    try:
        return f"{int(d_s):02d}{mo_num}{y_s}"
    except ValueError:
        return None


def _normalise_for_dedup(text: str) -> str:
    """
    Return a canonical fingerprint for grouping variant forms of the same
    identifier together in the review deduplication.

    - Date strings → DDMMYYYY (so "15/05/1975" == "15 May 1975")
    - Names → strip titles, strip punctuation, sort tokens, lowercase
      (so "Mr John Smith" == "Smith, John" == "John Smith")
    - NHS-number-like strings → digits only
    """
    t = (text or "").strip()
    if not t:
        return ""
    # Date variant
    dt = _normalise_date_str(t)
    if dt:
        return "DATE:" + dt
    # NHS number (10 digits) → digits only
    digits = re.sub(r'\D', '', t)
    if len(digits) == 10:
        return "NHS:" + digits
    # Name / phrase — strip titles and punctuation, sort tokens
    stripped = _DEDUP_TITLE_RE.sub('', t)
    stripped = re.sub(r'[^\w\s]', ' ', stripped)
    tokens = sorted(tok for tok in stripped.lower().split() if len(tok) >= 2)
    return 'NAME:' + ' '.join(tokens) if tokens else t.lower()


def _detect_patient_details(texts: list, model: str) -> dict:
    """
    Ask the LLM to extract the patient's own demographics from a sample of
    document text.  Returns a dict with keys:
      name, dob, nhs_number, address  (all str, empty string if not found).
    Uses the first 1 500 chars of each of up to five documents.
    """
    sample = "\n\n---NEXT DOCUMENT---\n\n".join(
        t[:1500] for t in texts[:5] if (t or "").strip()
    )
    if not sample.strip():
        return {}

    prompt = (
        "You are reading NHS medical records from a Subject Access Request bundle.\n"
        "Extract the PATIENT'S OWN personal details from the document text below.\n"
        "Return ONLY a JSON object — no explanation, no markdown.\n"
        "If a detail is not present use an empty string.\n\n"
        '{"name":"patient full name","dob":"date of birth exactly as written",'
        '"nhs_number":"NHS number exactly as written (e.g. 485 777 3456)",'
        '"address":"patient home address (full address or first line)"}\n\n'
        "Rules:\n"
        "- Extract the PATIENT's details, not a clinician, not a family member.\n"
        "- Look for labels: Name:, Patient:, DOB:, Date of Birth:, NHS No:, "
        "NHS Number:, Address:, Home Address:.\n"
        "- If the same person's name appears repeatedly in correspondence headers "
        "they are most likely the patient.\n\n"
        f"Document text:\n---\n{sample}\n---\n\nReturn JSON only."
    )
    try:
        ex     = ThreadPoolExecutor(max_workers=1)
        future = ex.submit(
            ollama.chat,
            model=model,
            messages=[
                {"role": "system",
                 "content": "Extract structured data from medical documents. Return valid JSON only."},
                {"role": "user", "content": prompt},
            ],
            format="json",
            options={"temperature": 0, "num_predict": 300},
        )
        try:
            resp = future.result(timeout=60)
        except FuturesTimeoutError:
            return {}
        finally:
            ex.shutdown(wait=False)
        raw    = resp["message"]["content"].strip()
        parsed = _extract_json(raw)
        if isinstance(parsed, dict):
            return {
                "name":       str(parsed.get("name",       "") or "").strip(),
                "dob":        str(parsed.get("dob",        "") or "").strip(),
                "nhs_number": str(parsed.get("nhs_number", "") or "").strip(),
                "address":    str(parsed.get("address",    "") or "").strip(),
            }
    except Exception:
        pass
    return {}


def classify_document(text: str, model: str) -> str:
    if not text.strip():
        return "Miscellaneous"
    cats   = "\n".join(f"- {c}" for c in SECTION_ORDER)
    prompt = (
        f"Classify this NHS GP medical document into exactly ONE of these five categories:\n{cats}\n\n"
        "Definitions:\n"
        "- Clinical Records: GP consultation notes, clinical entries, SOAP notes, problem lists, "
        "medication reviews, health checks, nurse or GP encounter records, summarised care records\n"
        "- Referral Letters: Letters written BY the GP surgery and sent TO another provider or "
        "specialist — outgoing referrals, GP covering letters sent on behalf of the patient\n"
        "- Correspondence: Documents RECEIVED by the GP surgery FROM external providers — "
        "hospital discharge summaries, specialist clinic letters, letters from consultants, "
        "social care letters, letters from other agencies (exclude results/test reports)\n"
        "- Results and Investigations: Pathology results, blood tests, imaging reports (X-ray, "
        "MRI, CT, ultrasound), ECG reports, microbiology, histology, any other investigation report\n"
        "- Miscellaneous: Anything that does not clearly fit the above four categories\n\n"
        f"Document excerpt:\n---\n{text[:2000]}\n---\n\n"
        "Reply with ONLY the category name from the list, exactly as written."
    )
    try:
        ex     = ThreadPoolExecutor(max_workers=1)
        future = ex.submit(
            ollama.chat,
            model=model,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0, "num_predict": 32},
        )
        try:
            resp = future.result(timeout=30)
        except FuturesTimeoutError:
            ex.shutdown(wait=False)
            return "Miscellaneous"
        finally:
            ex.shutdown(wait=False)
        result = resp["message"]["content"].strip()
        for cat in SECTION_ORDER:
            if cat.lower() in result.lower():
                return cat
    except Exception:
        pass
    return "Miscellaneous"


# =============================================================================
# File ingest  ->  fitz.Document + extracted text
# =============================================================================

def _text_to_fitz(text: str, title: str = "") -> fitz.Document:
    doc = fitz.open()
    PAGE_W, PAGE_H = 595, 842
    MX, MY         = 50, 60
    LH, FS, CPL    = 13, 9, 105
    MAX_L          = int((PAGE_H - MY * 2) / LH)

    raw_lines = ([title, "─" * 70, ""] if title else []) + text.splitlines()
    lines = []
    for ln in raw_lines:
        while len(ln) > CPL:
            lines.append(ln[:CPL])
            ln = ln[CPL:]
        lines.append(ln)

    page, lnum = None, 0
    for line in lines:
        if page is None or lnum >= MAX_L:
            page = doc.new_page(width=PAGE_W, height=PAGE_H)
            lnum = 0
        page.insert_text((MX, MY + lnum * LH), line, fontsize=FS, fontname="cour")
        lnum += 1

    if len(doc) == 0:
        doc.new_page(width=PAGE_W, height=PAGE_H)
    return doc


def _render_context_preview(
    doc: fitz.Document,
    search_text: str,
    dpi: int = 130,
    context_px: int = 260,
) -> tuple:
    """Render the page containing search_text as PNG bytes with a yellow highlight.

    Crops to ±context_px pixels above/below the match so the reviewer sees
    only the relevant portion rather than the full page.
    Returns (png_bytes | None, 1-based page_num, found: bool).
    """
    if not search_text or doc is None or not PIL_AVAILABLE:
        return None, 0, False

    search_norm = " ".join(search_text.split())
    candidates = [search_norm, search_norm.lower(), search_norm.upper(), search_norm.title()]

    for page_num, page in enumerate(doc):
        rects = []
        for variant in candidates:
            rects = page.search_for(variant)
            if rects:
                break
        if not rects:
            continue

        scale = dpi / 72.0
        pix   = page.get_pixmap(dpi=dpi)
        img   = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples).convert("RGBA")

        # Semi-transparent yellow highlight overlay
        overlay = PILImage.new("RGBA", img.size, (0, 0, 0, 0))
        draw    = PILImageDraw.Draw(overlay)
        all_y   = []
        for rect in rects:
            x0 = max(0,          int(rect.x0 * scale) - 4)
            y0 = max(0,          int(rect.y0 * scale) - 4)
            x1 = min(img.width,  int(rect.x1 * scale) + 4)
            y1 = min(img.height, int(rect.y1 * scale) + 4)
            draw.rectangle([x0, y0, x1, y1], fill=(255, 210, 0, 170))
            draw.rectangle([x0, y0, x1, y1], outline=(200, 80, 0, 255), width=2)
            all_y += [y0, y1]

        combined = PILImage.alpha_composite(img, overlay).convert("RGB")

        # Crop to a window centred on the highlighted area
        cy      = (min(all_y) + max(all_y)) // 2
        crop_y0 = max(0,              cy - context_px)
        crop_y1 = min(combined.height, cy + context_px)
        cropped = combined.crop((0, crop_y0, combined.width, crop_y1))

        buf = io.BytesIO()
        cropped.save(buf, format="PNG")
        return buf.getvalue(), page_num + 1, True

    return None, 0, False


def ingest_file(uploaded_file) -> tuple:
    """Returns (fitz.Document | None, extracted_text, error_msg, ocr_info).

    ocr_info is a human-readable string describing how text was obtained,
    e.g. "native PDF text", "Tesseract OCR (TIFF)", "Tesseract OCR (scanned PDF)",
    "no OCR — Tesseract not available", etc.
    """
    name = uploaded_file.name
    ext  = name.rsplit(".", 1)[-1].lower()
    data = uploaded_file.read()
    try:
        if ext == "pdf":
            doc      = fitz.open(stream=data, filetype="pdf")
            text     = "".join(p.get_text() for p in doc)
            ocr_info = "native PDF text"

            # Scanned / image-only PDF — fall back to Tesseract page by page
            if not text.strip():
                if TESSERACT_AVAILABLE:
                    ocr_parts = []
                    ocr_errors = []
                    for page_num, page in enumerate(doc, 1):
                        try:
                            pix = page.get_pixmap(dpi=200)
                            img = PILImage.frombytes(
                                "RGB", [pix.width, pix.height], pix.samples
                            )
                            ocr_parts.append(pytesseract.image_to_string(img))
                        except Exception as e:
                            ocr_errors.append(f"p{page_num}: {e}")
                    text = "\n".join(ocr_parts)
                    if text.strip():
                        ocr_info = (
                            f"Tesseract OCR (scanned PDF, {len(doc)} page(s)"
                            + (f", {len(ocr_errors)} page error(s)" if ocr_errors else "")
                            + ")"
                        )
                    else:
                        ocr_info = (
                            "Tesseract OCR attempted (scanned PDF) — no text extracted"
                            + (f"; errors: {'; '.join(ocr_errors)}" if ocr_errors else "")
                        )
                else:
                    ocr_info = "scanned PDF — no text layer; Tesseract not available (install it to enable OCR)"

        elif ext in ("docx", "doc"):
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx not installed — cannot open Word files")
            d     = DocxDocument(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for r in t.rows:
                    row = " | ".join(c.text.strip() for c in r.cells if c.text.strip())
                    if row:
                        parts.append(row)
            text     = "\n".join(parts)
            doc      = _text_to_fitz(text, title=name)
            ocr_info = "Word document — native text"

        elif ext in ("tiff", "tif"):
            _tiff = fitz.open(stream=data, filetype="tiff")
            doc   = fitz.open("pdf", _tiff.convert_to_pdf())
            _tiff.close()
            text = ""
            if TESSERACT_AVAILABLE:
                try:
                    img      = PILImage.open(io.BytesIO(data))
                    text     = pytesseract.image_to_string(img)
                    ocr_info = (
                        f"Tesseract OCR (TIFF) — {len(text.split())} words extracted"
                        if text.strip() else
                        "Tesseract OCR (TIFF) — no text extracted (blank or unreadable image?)"
                    )
                except Exception as e:
                    ocr_info = f"Tesseract OCR failed: {e}"
            else:
                ocr_info = "TIFF — Tesseract not available; no text extracted"

        elif ext == "rtf":
            if RTF_AVAILABLE:
                text = parse_rtf(data.decode("utf-8", errors="ignore"))
            else:
                raw  = data.decode("utf-8", errors="ignore")
                text = re.sub(r"\\[a-z]+\d*[ ]?", " ", raw)
                text = re.sub(r"[{}\\]", "", text).strip()
            doc      = _text_to_fitz(text, title=name)
            ocr_info = "RTF — native text"

        elif ext == "txt":
            text     = data.decode("utf-8", errors="ignore")
            doc      = _text_to_fitz(text, title=name)
            ocr_info = "plain text"

        else:
            return None, "", f"Unsupported format: .{ext}", ""

        return doc, text, "", ocr_info
    except Exception as exc:
        return None, "", str(exc), ""


# =============================================================================
# Post-processing: expand name redactions to catch first-name-only occurrences
# =============================================================================

def _expand_name_redactions(proposed: list, text: str, patient_name: str = "") -> list:
    """
    For each THIRD_PARTY_IDENTIFIER redaction that looks like a full name
    (two or more words), extract each component word and add a separate
    redaction entry for any that appear STANDALONE elsewhere in the document
    (i.e. outside the immediate context of the full name).

    This catches cases like: LLM flags "Michelle Granger" but the document
    later refers to her as just "Michelle" in quoted speech.

    patient_name: the subject of the SAR — name parts matching the patient's
    own name are never added as new redaction targets.
    """
    if not text:
        return proposed

    # Build a set of the patient's own name tokens to protect from over-redaction.
    # This prevents e.g. "Sampledata" (shared surname with a family member)
    # being expanded into a redaction that would erase the patient's own header lines.
    _pn_tokens: set[str] = set()
    if patient_name.strip():
        for tok in patient_name.strip().lower().split():
            clean_tok = tok.strip(".,;:()[]'\"–—-")
            if len(clean_tok) >= 3:
                _pn_tokens.add(clean_tok)

    # Common English words that are never proper name tokens.
    _STOPWORDS = {
        "the", "a", "an", "of", "at", "on", "in", "to", "for", "from", "with",
        "who", "what", "where", "when", "how", "that", "this", "and", "or",
        "but", "not", "no", "is", "was", "are", "were", "be", "been", "has",
        "have", "had", "do", "does", "did", "will", "would", "can", "could",
        "may", "might", "she", "he", "her", "his", "their", "they", "our",
        "runs", "post", "lives", "works", "near", "next", "door", "road",
        "street", "lane", "avenue", "close", "drive", "house", "flat", "office",
        "woman", "man", "lady", "person", "child", "boy", "girl", "family",
        "local", "nearby", "down",
        # Role / occupation words — prevent expanding role titles into name tokens
        "social", "worker", "coordinator", "senior", "care", "staff", "health",
        "support", "key", "lead", "head", "deputy", "assistant", "registered",
        "qualified", "community", "liaison", "service", "services", "team",
        "manager", "director", "officer", "nurse", "doctor", "consultant",
        "specialist", "therapist", "counsellor", "advisor", "adviser",
        # Salutation / correspondence words
        "dear", "tel", "ref", "re", "via", "attn",
        # Honorifics / titles — prevent "(Mrs" being extracted as a name token
        "mrs", "miss", "prof", "sir", "rev", "lord", "dame",
    }

    existing_lower = {(r.get("text") or "").strip().lower() for r in proposed}
    extra = []

    for item in proposed:
        tag = item.get("tag", "")
        raw = (item.get("text") or "").strip()

        # For AGENCY_CONFIDENTIAL_INFO items like
        # "Claire Hughes (Warwickshire Children's Services, Tel: 01926 000055)"
        # extract the name portion before the first '(' or ',' and add it
        # as a standalone redaction if it appears elsewhere in the document.
        if tag == "AGENCY_CONFIDENTIAL_INFO":
            # Extract leading name-like segment (before first bracket or comma)
            name_part = re.split(r'[,(]', raw)[0].strip()
            # Must look like a 2-word proper name (Firstname Lastname)
            np_parts = name_part.split()
            if (2 <= len(np_parts) <= 3
                    and all(p[0].isupper() for p in np_parts if p)
                    and name_part.lower() not in existing_lower):
                # Check it appears standalone outside the full string context
                pattern = r'(?<!\w)' + re.escape(name_part) + r'(?!\w)'
                for m in re.finditer(pattern, text, re.IGNORECASE):
                    window_start = max(0, m.start() - len(raw) - 5)
                    window_end   = min(len(text), m.end() + len(raw) + 5)
                    window       = text[window_start:window_end]
                    if raw.lower() not in window.lower():
                        extra.append({
                            "text":        name_part,
                            "tag":         tag,
                            "reason":      f"Standalone name from agency contact (propagated from \"{raw}\")",
                            "replacement": item.get("replacement", "[REDACTED - agency confidential information]"),
                            "context":     item.get("context", ""),
                            "approved":    True,
                        })
                        existing_lower.add(name_part.lower())
                        # Also expand individual surname for further-downstream standalone refs
                        if len(np_parts) == 2:
                            surname = np_parts[1]
                            if (surname.lower() not in existing_lower
                                    and surname.lower() not in _STOPWORDS
                                    and len(surname) >= 3):
                                pat2 = r'(?<!\w)' + re.escape(surname) + r'(?!\w)'
                                for m2 in re.finditer(pat2, text, re.IGNORECASE):
                                    w2_start = max(0, m2.start() - len(raw) - 5)
                                    w2_end   = min(len(text), m2.end() + len(raw) + 5)
                                    w2       = text[w2_start:w2_end]
                                    if raw.lower() not in w2.lower() and name_part.lower() not in w2.lower():
                                        extra.append({
                                            "text":        surname,
                                            "tag":         tag,
                                            "reason":      f"Surname of agency contact (propagated from \"{raw}\")",
                                            "replacement": item.get("replacement", "[REDACTED - agency confidential information]"),
                                            "context":     item.get("context", ""),
                                            "approved":    True,
                                        })
                                        existing_lower.add(surname.lower())
                                        break
                        break
            continue  # do not fall through to the THIRD_PARTY_IDENTIFIER word-split logic

        if tag != "THIRD_PARTY_IDENTIFIER":
            continue
        parts = raw.split()
        if len(parts) < 2:
            continue   # already a single word — nothing to expand

        # Do NOT expand email salutation strings ("Dear Dr X", "To Whom", etc.)
        _SALUTATIONS = {"dear", "to", "re", "attn", "attention"}
        if parts[0].lower().strip(".,;:") in _SALUTATIONS:
            continue

        # Do NOT expand address strings — they contain place names (city, town) that
        # occur legitimately in institution names like "Bradford Royal Infirmary".
        _ADDRESS_KEYWORDS = {
            "road", "street", "avenue", "lane", "close", "drive", "court", "place",
            "way", "grove", "gardens", "crescent", "terrace", "walk", "parade",
        }
        if any(kw in raw.lower() for kw in _ADDRESS_KEYWORDS):
            continue

        # Only expand strings that look like proper names: every significant word
        # must start with an uppercase letter and not be a stopword.
        # This prevents "the woman who runs the post office" being split into
        # individual common words that would corrupt the document.
        for part in parts:
            # Strip common punctuation that can attach to a name in free text
            clean = part.strip(".,;:()[]'\"–—-")
            if len(clean) < 3:
                continue   # skip initials / very short tokens
            if clean.lower() in existing_lower:
                continue   # already being redacted
            if clean.lower() in _pn_tokens:
                continue   # part of patient's own name — never redact

            # Only expand parts that look like proper name tokens:
            # must start with uppercase and not be a generic English word
            if not (len(clean) >= 3 and clean[0].isupper() and clean.lower() not in _STOPWORDS):
                continue   # not a proper name token — skip this part

            # Word-boundary search for standalone occurrence
            pattern = r'(?<!\w)' + re.escape(clean) + r'(?!\w)'
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            if not matches:
                continue

            # At least one occurrence must be OUTSIDE the span of the full name
            # AND not embedded in an organisation/company name.
            _ORG_SUFFIXES_RE = re.compile(
                r'\s*(?:&|and)\s+[A-Z]|\b(?:LLP|Ltd|PLC|plc|Inc|Trust|NHS|LTD)\b',
                re.IGNORECASE,
            )
            standalone = False
            for m in matches:
                # Build the surrounding window and check full name isn't there
                window_start = max(0, m.start() - len(raw) - 5)
                window_end   = min(len(text), m.end() + len(raw) + 5)
                window       = text[window_start:window_end]
                if raw.lower() in window.lower():
                    continue  # full name also present — not a standalone occurrence
                # Check that the immediate right-hand context doesn't suggest this
                # token is the first word of an org name ("Thompson & Reed LLP")
                right_ctx = text[m.end(): m.end() + 20]
                if _ORG_SUFFIXES_RE.match(right_ctx):
                    continue  # looks like org name — skip
                standalone = True
                break

            if standalone:
                extra.append({
                    "text":        clean,
                    "tag":         "THIRD_PARTY_IDENTIFIER",
                    "reason":      f"Standalone name-part of third party (expanded from \"{raw}\")",
                    "replacement": "[REDACTED - third-party personal data]",
                    "context":     item.get("context", ""),
                    "approved":    True,
                })
                existing_lower.add(clean.lower())

    return proposed + extra


# =============================================================================
# Post-processing: expand agency contacts to catch paired name/phone
# =============================================================================

def _expand_agency_contacts(proposed: list, text: str, patient_name: str = "") -> list:
    """
    When an AGENCY_CONFIDENTIAL_INFO or THIRD_PARTY_IDENTIFIER proposed redaction
    contains a phone number but the adjacent name was missed (or vice versa), try to
    locate the missing counterpart in the surrounding text and add it.

    Handles structured blocks like:
        Social worker: Diane Okafor
        Direct line: 01925 000055
    where the LLM may catch one field but not the other.
    """
    import re as _re

    _PHONE_PAT = _re.compile(
        r'\b(\d{5}\s\d{6}|\d{4}\s\d{3}\s\d{4}|\d{11}|\+44[\s\d]{10,13})\b'
    )
    _EMAIL_PAT = _re.compile(
        r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
    )
    # Proper name: 2 words, each Title-case, each ≥ 2 chars.
    # Deliberately excludes 3+ word strings to avoid institutional names.
    _NAME_PAT = _re.compile(
        r'\b([A-Z][a-z]{1,}(?:\s+[A-Z][a-z]{1,}){1,2})\b'
    )
    _AGENCY_TAGS = {"AGENCY_CONFIDENTIAL_INFO", "THIRD_PARTY_IDENTIFIER"}

    # Words that indicate an organisation or role, not a person's name.
    # Candidates containing any of these are skipped.
    _INSTITUTIONAL_WORDS = {
        "hospital", "infirmary", "royal", "nhs", "trust", "refuge", "liaison",
        "services", "service", "clinic", "surgery", "centre", "center", "council",
        "authority", "department", "office", "association", "police", "court",
        "school", "college", "university", "academy", "foundation", "unit",
        "ward", "team", "group", "limited", "ltd", "plc", "inc", "officer",
        # Social/care sector words — prevent 'Kent Adult Social' matching as a person
        "social", "adult", "care", "health", "mental", "children", "young",
        "community", "housing", "probation", "welfare", "voluntary",
    }

    def _is_institutional(name: str) -> bool:
        """Return True if the candidate looks like an organisation name."""
        words = name.lower().split()
        return any(w in _INSTITUTIONAL_WORDS for w in words)

    def _is_plausible_person(name: str) -> bool:
        """
        Return True if the candidate looks like a real person's name:
        - 2 or 3 words (3-word names like 'P. Hall' qualify if short initial)
        - Not institutional
        """
        words = name.split()
        if not (1 < len(words) <= 3):
            return False
        return not _is_institutional(name)

    # Build patient name token set so we never add the patient as a redaction target
    _pn_tokens_agency: set = set()
    for tok in patient_name.strip().lower().split():
        clean = tok.strip(".,;:()[]'\"")
        if len(clean) >= 3:
            _pn_tokens_agency.add(clean)

    existing_lower = {(r.get("text") or "").strip().lower() for r in proposed}
    lines = text.splitlines()
    extra = []

    for item in proposed:
        if item.get("tag") not in _AGENCY_TAGS:
            continue
        item_text = (item.get("text") or "").strip()

        # Case A: item is a phone number → look for adjacent personal name (±1 line)
        if _PHONE_PAT.fullmatch(item_text.replace(" ", "")):
            for li, line in enumerate(lines):
                if item_text in line:
                    window = lines[max(0, li - 1): li + 2]
                    for wline in window:
                        for m in _NAME_PAT.finditer(wline):
                            candidate = m.group(1)
                            if (candidate.lower() not in existing_lower
                                    and _is_plausible_person(candidate)):
                                extra.append({
                                    "text":        candidate,
                                    "tag":         item.get("tag"),
                                    "reason":      f"Name associated with agency contact (paired with {item_text})",
                                    "replacement": "[REDACTED - agency contact]",
                                    "context":     wline.strip(),
                                    "approved":    True,
                                })
                                existing_lower.add(candidate.lower())
                    break

        # Case B: item is a name → look for adjacent phone numbers (±1 line)
        elif _NAME_PAT.fullmatch(item_text) and _is_plausible_person(item_text):
            for li, line in enumerate(lines):
                if item_text in line:
                    window = lines[max(0, li - 1): li + 2]
                    for wline in window:
                        for m in _PHONE_PAT.finditer(wline):
                            candidate = m.group(0)
                            if candidate.lower() not in existing_lower:
                                extra.append({
                                    "text":        candidate,
                                    "tag":         item.get("tag"),
                                    "reason":      f"Phone associated with agency contact (paired with {item_text})",
                                    "replacement": "[REDACTED - agency contact]",
                                    "context":     wline.strip(),
                                    "approved":    True,
                                })
                                existing_lower.add(candidate.lower())
                    break

        # Case C: item is an email address → look for the owner's name on the same line
        elif _EMAIL_PAT.fullmatch(item_text):
            for li, line in enumerate(lines):
                if item_text in line:
                    # Search the same line and ±1 lines for a proper name
                    window = lines[max(0, li - 1): li + 2]
                    for wline in window:
                        for m in _NAME_PAT.finditer(wline):
                            candidate = m.group(1)
                            # Skip if any part of the candidate matches a patient name token
                            candidate_toks = {w.lower() for w in candidate.split()}
                            if candidate_toks & _pn_tokens_agency:
                                continue
                            if (candidate.lower() not in existing_lower
                                    and _is_plausible_person(candidate)):
                                extra.append({
                                    "text":        candidate,
                                    "tag":         "THIRD_PARTY_IDENTIFIER",
                                    "reason":      f"Named owner of email address {item_text}",
                                    "replacement": "[REDACTED - third-party personal data]",
                                    "context":     wline.strip(),
                                    "approved":    True,
                                })
                                existing_lower.add(candidate.lower())
                    break

    return proposed + extra


def _expand_agency_professionals(proposed: list, text: str, patient_name: str = "") -> list:
    """
    Code-level fallback to catch named agency professionals that the LLM misses
    even when prompted.  Targets three high-miss patterns:

      1. "by [Name] (private/independent/employer-commissioned [role])"
         e.g. "by Lisa Torn (private physiotherapist, commissioned by employer's insurer)"

      2. "[Role] [Name] ([org/service]...)"
         e.g. "Therapist Claire Inder (NHS Eating Disorder Service, Peterborough)"

      3. "solicitor [Name] (" — solicitors named in correspondence
         e.g. "solicitor James Hazeldine (Hazeldine & Partners LLP)"

    Adds new AGENCY_CONFIDENTIAL_INFO entries for matched names not already proposed.
    """
    _NAME = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'

    _PATTERNS = [
        # "by [Name] (private / independent / employer-commissioned [clinical role])"
        re.compile(
            r'\bby\s+' + _NAME +
            r'\s*\((?:private|independent|employer[- ]commissioned)\s*'
            r'(?:physiotherapist|physio|therapist|counsellor|psychologist|'
            r'occupational health)',
            re.IGNORECASE,
        ),
        # "[Clinical role title] [Name] ([org"
        re.compile(
            r'\b(?:Therapist|Physiotherapist|Counsellor|Psychologist|'
            r'Occupational Health Adviser|Occupational Health Advisor)\s+' +
            _NAME + r'\s*[\(,]',
            re.IGNORECASE,
        ),
        # "solicitor [Name]" or "from solicitor [Name]"
        re.compile(
            r'\bsolicitor\s+' + _NAME + r'\b',
            re.IGNORECASE,
        ),
    ]

    _pn_lower = (patient_name or "").strip().lower()
    _existing = {(p.get("text") or "").strip().lower() for p in proposed}

    additions = []
    for pat in _PATTERNS:
        for m in pat.finditer(text):
            name = m.group(1).strip()
            if not name:
                continue
            if name.lower() in _existing:
                continue
            if _pn_lower and name.lower() == _pn_lower:
                continue
            additions.append({
                "text":        name,
                "tag":         "AGENCY_CONFIDENTIAL_INFO",
                "reason":      "Named agency professional (pattern-based extraction).",
                "replacement": "[REDACTED - agency confidential information]",
            })
            _existing.add(name.lower())

    return proposed + additions


# =============================================================================
# Apply approved redactions
# =============================================================================

def _find_text_on_page(page, needle: str) -> list:
    """
    Case-insensitive, whitespace-tolerant text search.
    Handles line-breaks that split words across PDF lines.
    Returns list of Rect.
    """
    needle = " ".join(needle.split())   # normalise all whitespace to single spaces
    if not needle:
        return []

    # 1. Exact match
    rects = page.search_for(needle)
    if rects:
        return rects

    # 2. Common case variants
    for variant in (needle.lower(), needle.upper(), needle.title()):
        if variant != needle:
            rects = page.search_for(variant)
            if rects:
                return rects

    # 3. Case-insensitive via raw page text (handles different casing)
    page_text = page.get_text("text")
    pos = page_text.lower().find(needle.lower())
    if pos != -1:
        actual = page_text[pos: pos + len(needle)]
        rects = page.search_for(actual)
        if rects:
            return rects

    # 4. Whitespace-normalised: collapse all whitespace in page text too.
    #    Catches cases where a line-break sits between two words of the needle.
    page_flat = " ".join(page_text.split())
    pos = page_flat.lower().find(needle.lower())
    if pos != -1:
        actual_flat = page_flat[pos: pos + len(needle)]
        rects = page.search_for(actual_flat)
        if rects:
            return rects
        # Try case variants of the flat version
        for variant in (actual_flat.lower(), actual_flat.upper(), actual_flat.title()):
            rects = page.search_for(variant)
            if rects:
                return rects

    return []


def _rects_overlap(r1: fitz.Rect, r2: fitz.Rect, threshold: float = 0.4) -> bool:
    """Return True if r1 and r2 overlap by more than threshold of the smaller rect's area."""
    ix0 = max(r1.x0, r2.x0)
    iy0 = max(r1.y0, r2.y0)
    ix1 = min(r1.x1, r2.x1)
    iy1 = min(r1.y1, r2.y1)
    if ix0 >= ix1 or iy0 >= iy1:
        return False
    inter = (ix1 - ix0) * (iy1 - iy0)
    a1    = max((r1.x1 - r1.x0) * (r1.y1 - r1.y0), 1e-6)
    a2    = max((r2.x1 - r2.x0) * (r2.y1 - r2.y0), 1e-6)
    return inter / min(a1, a2) > threshold


def _sanitise_replacement(text: str) -> str:
    """Replace characters unsupported by PDF built-in Helvetica with safe equivalents."""
    return (
        text
        .replace("\u2014", "-")   # em-dash → hyphen
        .replace("\u2013", "-")   # en-dash → hyphen
        .replace("\u2018", "'")   # left single quote
        .replace("\u2019", "'")   # right single quote
        .replace("\u201c", '"')   # left double quote
        .replace("\u201d", '"')   # right double quote
    )


def apply_approved_redactions(doc: fitz.Document, approved_items: list) -> tuple:
    """
    Black-box all approved strings. Returns (modified_doc, redaction_count).

    Improvements over naïve approach:
    - Sanitises replacement text (em-dashes etc.) for Helvetica compatibility.
    - Deduplicates overlapping rects before adding annotations so two redactions
      covering the same span don't interleave their replacement text labels.
    """
    # apply_redactions() requires a PDF; convert if necessary (e.g. TIFF opened directly)
    if not doc.is_pdf:
        doc = fitz.open("pdf", doc.convert_to_pdf())

    count  = 0
    unique = {}
    for item in approved_items:
        t = (item.get("text") or "").strip()
        if t and len(t) >= 2:
            raw_repl = item.get("replacement", "[REDACTED]")
            unique[t] = _sanitise_replacement(raw_repl)

    for page in doc:
        # Collect all (rect, replacement) pairs for this page first so we can
        # deduplicate overlapping rects before calling add_redact_annot.
        pending: list[tuple[fitz.Rect, str]] = []
        for s, replacement in unique.items():
            for rect in _find_text_on_page(page, s):
                # Skip if this rect significantly overlaps one already queued
                overlaps = any(_rects_overlap(rect, existing_r) for existing_r, _ in pending)
                if not overlaps:
                    pending.append((rect, replacement))

        for rect, replacement in pending:
            try:
                page.add_redact_annot(
                    rect,
                    text=replacement,
                    fontname="helv",
                    fontsize=5,
                    fill=(0.85, 0.85, 0.85),
                )
            except Exception:
                page.add_redact_annot(rect, fill=(0, 0, 0))
            count += 1

        page.apply_redactions()

    return doc, count


# =============================================================================
# Bundle assembly — cover page + section dividers + documents
# =============================================================================

def _cover_page(sar_ref, operator, date_str, total_docs) -> fitz.Document:
    doc  = fitz.open()
    page = doc.new_page(width=595, height=842)

    page.draw_rect(fitz.Rect(0, 0, 595, 150), color=NHS_BLUE, fill=NHS_BLUE)
    page.insert_text((40, 75),  "NHS",                      fontsize=40, color=WHITE, fontname="helv")
    page.insert_text((40, 108), "Subject Access Request",   fontsize=17, color=WHITE, fontname="helv")
    page.insert_text((40, 133), "REDACTED DOCUMENT BUNDLE", fontsize=13, color=WHITE, fontname="helv")

    def row(lbl, val, y):
        page.insert_text((40,  y), lbl,        fontsize=11, color=GREY,  fontname="helv")
        page.insert_text((210, y), val or "—", fontsize=11, color=BLACK, fontname="helv")
        return y + 26

    y = 210
    y = row("SAR Reference / Subject:", sar_ref,         y)
    y = row("Processed by:",            operator,        y)
    y = row("Date processed:",          date_str,        y)
    y = row("Total documents:",         str(total_docs), y)

    page.draw_line((40, y + 10), (555, y + 10), color=LT_GREY, width=1)

    notice = (
        "Processed under UK GDPR / DPA 2018 / ICO SAR guidance. "
        "Third-party and safeguarding information has been reviewed by the named operator. "
        "All redaction decisions have been individually approved before this bundle was produced."
    )
    yn = y + 30
    for chunk in [notice[i:i + 90] for i in range(0, len(notice), 90)]:
        page.insert_text((40, yn), chunk, fontsize=9, color=GREY, fontname="helv")
        yn += 14

    page.draw_rect(fitz.Rect(0, 830, 595, 842), color=NHS_BLUE, fill=NHS_BLUE)
    page.insert_text((40, 839), "CONFIDENTIAL — FOR AUTHORISED VIEWING ONLY",
                     fontsize=8, color=WHITE, fontname="helv")
    return doc


def _divider_page(section, count, idx, total, date_range="") -> fitz.Document:
    doc  = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.draw_rect(fitz.Rect(0, 0, 12, 842), color=NHS_BLUE, fill=NHS_BLUE)
    page.insert_text((30, 60),  f"SECTION {idx} OF {total}",
                     fontsize=9,  color=GREY,  fontname="helv")
    page.insert_text((30, 400), section,
                     fontsize=30, color=BLACK, fontname="helv")
    subtitle = f"{count} document{'s' if count != 1 else ''}"
    if date_range:
        subtitle += f"  ·  {date_range}"
    page.insert_text((30, 435), subtitle,
                     fontsize=13, color=GREY,  fontname="helv")
    page.insert_text((30, 455), "Ordered most recent first",
                     fontsize=9,  color=GREY,  fontname="helv")
    page.draw_rect(fitz.Rect(0, 830, 595, 842), color=LT_GREY, fill=LT_GREY)
    return doc


def build_bundle(proc_docs, sar_ref="", operator="", date_str="") -> fitz.Document:
    groups = {}
    for item in proc_docs:
        groups.setdefault(item["section"], []).append(item)

    # Sort each section: most recent document first; undated docs go to the end
    for sec in groups:
        groups[sec].sort(
            key=lambda x: x.get("doc_date", datetime.date.min),
            reverse=True,
        )

    ordered = [(s, groups[s]) for s in SECTION_ORDER if s in groups]
    out     = fitz.open()
    out.insert_pdf(_cover_page(sar_ref, operator, date_str, len(proc_docs)))
    for idx, (sec, docs) in enumerate(ordered, 1):
        # Build a date-range subtitle for the divider (e.g. "Jan 2019 – Mar 2024")
        real_dates = [
            d["doc_date"] for d in docs
            if d.get("doc_date") and d["doc_date"] != datetime.date.min
        ]
        if real_dates:
            oldest  = min(real_dates).strftime("%b %Y")
            newest  = max(real_dates).strftime("%b %Y")
            date_range = newest if oldest == newest else f"{oldest} – {newest}"
        else:
            date_range = ""
        out.insert_pdf(_divider_page(sec, len(docs), idx, len(ordered), date_range))
        for item in docs:
            out.insert_pdf(item["doc"])
    return out


# =============================================================================
# INSURANCE FORM FILLER — helpers
# =============================================================================

_FF_SYSTEM = (
    "You are a medical form-filling assistant for an NHS GP practice. "
    "You respond with valid JSON only. No preamble, no explanation, no markdown."
)

_FF_EXTRACT_PROMPT = """\
You are analysing text extracted by OCR from a scanned medical insurance or GP report form.
Identify EVERY field or question that needs to be completed on this form.
Look for: field labels followed by blank lines/spaces, numbered questions, table rows with labels.

For EACH field return one JSON object with these keys:
- "label"             : exact text of the field label as it appears (trim trailing colons/underscores/spaces)
- "field_type"        : "text" | "date" | "yes_no" | "number" | "checkboxes" | "signature" | "textarea"
- "needs_manual_input": true if this info is unlikely to be in a patient medical record
  (e.g. policy number, insurer name/address, claim reference, witness, authorisation signature)
- "manual_hint"       : (only if needs_manual_input=true) short guidance for the user

Return ONLY this JSON — nothing else:
{{"fields": [...]}}

Form text (OCR extracted):
---
{form_text}
---"""

_FF_ANSWER_PROMPT = """\
You are completing a medical form for patient: {patient_name}.

Using ONLY the patient record text provided below, answer each form field as concisely as a form requires.
- If the record contains the answer, provide it.
- For dates use DD/MM/YYYY format.
- If the record does NOT clearly contain the answer, set "answer" to null.
- In "evidence" copy the exact sentence/phrase from the record that supports your answer.
- Set confidence: "high" (clearly stated), "medium" (inferred), "low" (uncertain), "none" (not found).

Fields to complete (JSON):
{fields_json}

Patient record:
---
{epr_text}
---

Return ONLY this JSON:
{{"answers": [
  {{"label": "...", "answer": "..." or null, "evidence": "...", "confidence": "high|medium|low|none"}}
]}}"""


def extract_form_fields_llm(form_text: str, model: str) -> list:
    """Ask LLM to identify all fillable fields in the form. Returns list of field dicts."""
    prompt = _FF_EXTRACT_PROMPT.format(form_text=form_text[:8000])
    try:
        ex     = ThreadPoolExecutor(max_workers=1)
        future = ex.submit(
            ollama.chat,
            model=model,
            messages=[
                {"role": "system", "content": _FF_SYSTEM},
                {"role": "user",   "content": prompt},
            ],
            format="json",
            options={"temperature": 0.05, "num_predict": 2048},
        )
        try:
            resp = future.result(timeout=120)
        except FuturesTimeoutError:
            ex.shutdown(wait=False)
            return []
        finally:
            ex.shutdown(wait=False)
        parsed = _extract_json(resp["message"]["content"].strip())
        if parsed and "fields" in parsed:
            return parsed["fields"] or []
    except Exception:
        pass
    return []


def answer_fields_from_epr(
    fields: list,
    epr_text: str,
    patient_name: str,
    model: str,
    status_cb=None,
) -> list:
    """Ask LLM to answer all form fields using EPR records. Returns enriched field dicts."""
    BATCH      = 8
    epr_sample = epr_text[:14000]
    answered   = []

    for batch_start in range(0, len(fields), BATCH):
        batch         = fields[batch_start : batch_start + BATCH]
        auto_fields   = [f for f in batch if not f.get("needs_manual_input")]
        manual_fields = [f for f in batch if f.get("needs_manual_input")]

        if status_cb:
            status_cb(
                f"Answering fields {batch_start + 1}–"
                f"{min(batch_start + BATCH, len(fields))} of {len(fields)}…"
            )

        # Manual fields — mark for user input immediately
        for f in manual_fields:
            f = dict(f)
            f.update({"answer": None, "evidence": "", "confidence": "none",
                      "approved": False, "final_answer": ""})
            answered.append(f)

        if not auto_fields:
            continue

        fields_json = json.dumps(
            [{"label": f.get("label", ""), "field_type": f.get("field_type", "text")}
             for f in auto_fields],
            indent=2,
        )
        prompt = _FF_ANSWER_PROMPT.format(
            patient_name=patient_name or "the patient",
            fields_json=fields_json,
            epr_text=epr_sample,
        )
        try:
            ex     = ThreadPoolExecutor(max_workers=1)
            future = ex.submit(
                ollama.chat,
                model=model,
                messages=[
                    {"role": "system", "content": _FF_SYSTEM},
                    {"role": "user",   "content": prompt},
                ],
                format="json",
                options={"temperature": 0.05, "num_predict": 2048},
            )
            try:
                resp = future.result(timeout=120)
            except FuturesTimeoutError:
                ex.shutdown(wait=False)
                resp = None
            finally:
                ex.shutdown(wait=False)

            ans_map = {}
            if resp:
                parsed = _extract_json(resp["message"]["content"].strip())
                if parsed and "answers" in parsed:
                    ans_map = {a.get("label", ""): a for a in parsed["answers"]}

            for f in auto_fields:
                f   = dict(f)
                ans = ans_map.get(f.get("label", ""), {})
                f["answer"]       = ans.get("answer")
                f["evidence"]     = ans.get("evidence", "")
                f["confidence"]   = ans.get("confidence", "none")
                f["approved"]     = bool(f["answer"])
                f["final_answer"] = f["answer"] or ""
                answered.append(f)

        except Exception:
            for f in auto_fields:
                f = dict(f)
                f.update({"answer": None, "evidence": "", "confidence": "none",
                          "approved": False, "final_answer": ""})
                answered.append(f)

    # Restore original field order
    label_to_ans = {f.get("label", ""): f for f in answered}
    return [label_to_ans.get(f.get("label", ""), f) for f in fields]


def _find_label_in_ocr(ocr_data: dict, label: str):
    """Return (x, y, w, h) pixel bbox of the label in pytesseract data, or None."""
    if not ocr_data:
        return None
    label_words = [w.lower() for w in label.split() if w.strip()]
    if not label_words:
        return None
    texts = [ocr_data["text"][i].lower().strip() for i in range(len(ocr_data["text"]))]
    confs = [int(ocr_data["conf"][i]) for i in range(len(ocr_data["conf"]))]

    # Try sequences of decreasing length
    for attempt_len in range(min(len(label_words), 4), 0, -1):
        seq = label_words[:attempt_len]
        for i in range(len(texts) - attempt_len + 1):
            if confs[i] < 20:
                continue
            if all(texts[i + j] == seq[j] for j in range(attempt_len)):
                last_i = i + attempt_len - 1
                x  = ocr_data["left"][i]
                y  = ocr_data["top"][i]
                x2 = ocr_data["left"][last_i] + ocr_data["width"][last_i]
                h  = max(ocr_data["height"][i : last_i + 1])
                return x, y, x2 - x, h
    return None


def _ingest_form(uploaded_file) -> tuple:
    """
    Ingest a scanned insurance form (PDF / image).
    Returns (fitz.Document | None, extracted_text, error_msg, has_text_layer).
    """
    name = uploaded_file.name
    ext  = name.rsplit(".", 1)[-1].lower()
    data = uploaded_file.read()

    try:
        if ext == "pdf":
            doc        = fitz.open(stream=data, filetype="pdf")
            text       = "".join(p.get_text() for p in doc)
            has_native = bool(text.strip())
            if not has_native and TESSERACT_AVAILABLE and PIL_AVAILABLE:
                parts = []
                for page in doc:
                    pix = page.get_pixmap(dpi=200)
                    img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    parts.append(pytesseract.image_to_string(img))
                text = "\n".join(parts)
            return doc, text, "", has_native

        elif ext in ("jpg", "jpeg", "png", "tiff", "tif", "bmp"):
            if not PIL_AVAILABLE:
                return None, "", "Pillow not installed — cannot open image files.", False
            img    = PILImage.open(io.BytesIO(data)).convert("RGB")
            buf    = io.BytesIO()
            img.save(buf, format="PDF", resolution=150)
            buf.seek(0)
            doc    = fitz.open(stream=buf.read(), filetype="pdf")
            text   = pytesseract.image_to_string(img) if TESSERACT_AVAILABLE else ""
            return doc, text, "", False

        else:
            return None, "", f"Unsupported form format: .{ext}", False

    except Exception as exc:
        return None, "", str(exc), False


def _append_summary_page(doc: fitz.Document, answered_fields: list):
    """Append a formatted Q&A summary page to doc (in-place)."""
    PAGE_W, PAGE_H = 595, 842
    MX, MY         = 40, 70
    LH             = 17

    def _new_pg(title="COMPLETED FORM — ANSWER SUMMARY"):
        p = doc.new_page(width=PAGE_W, height=PAGE_H)
        p.draw_rect(fitz.Rect(0, 0, PAGE_W, 48), color=NHS_BLUE, fill=NHS_BLUE)
        p.insert_text((MX, 32), title, fontsize=12, color=WHITE, fontname="helv")
        return p, MY + 8

    page, y = _new_pg()

    for field in answered_fields:
        label  = (field.get("label") or "—")[:90]
        answer = (field.get("final_answer") or "").strip() or "[NOT COMPLETED]"
        conf   = field.get("confidence", "none")
        ev     = (field.get("evidence") or "")[:240]

        need_h = LH * 2 + (10 * (1 + len(ev) // 90)) + 12
        if y + need_h > PAGE_H - 40:
            page, y = _new_pg("COMPLETED FORM — ANSWER SUMMARY (cont.)")

        page.insert_text((MX, y), label, fontsize=8, color=GREY, fontname="helv")
        y += LH - 3

        marker = {"high": "✓", "medium": "~", "low": "?", "none": "—"}.get(conf, "")
        page.insert_text(
            (MX + 8, y),
            f"{marker}  {answer}"[:95],
            fontsize=10, color=BLACK, fontname="helv",
        )
        y += LH

        if ev:
            for i in range(0, min(len(ev), 240), 90):
                page.insert_text(
                    (MX + 8, y), ev[i : i + 90],
                    fontsize=7, color=GREY, fontname="helv",
                )
                y += 10

        page.draw_line((MX, y + 2), (PAGE_W - MX, y + 2), color=LT_GREY, width=0.4)
        y += 9


def build_filled_form_pdf(
    form_doc: fitz.Document,
    answered_fields: list,
    has_text_layer: bool = False,
) -> fitz.Document:
    """
    Return a new fitz.Document with answers overlaid on the form pages,
    plus a Q&A summary page appended.
    Uses PIL image overlay for scanned forms; text insertion for native PDFs.
    """
    DPI   = 150
    SCALE = DPI / 72.0
    BLUE  = (0, 60, 180)          # PIL colour (RGB)
    BFITZ = (0.0, 0.25, 0.75)    # fitz colour (0-1)
    FS    = 8.5

    out = fitz.open()

    # Try to load a font for PIL rendering (Windows: Arial; fallback: default)
    _pil_font = None
    if PIL_AVAILABLE:
        try:
            from PIL import ImageFont as _PILFont
            _pil_font = _PILFont.truetype("arial.ttf", 14)
        except Exception:
            try:
                from PIL import ImageFont as _PILFont
                _pil_font = _PILFont.load_default()
            except Exception:
                _pil_font = None

    for page_num, page in enumerate(form_doc):
        if has_text_layer:
            # ── Native PDF: use search_for + insert_text ────────────────────
            new_pg = out.new_page(width=page.rect.width, height=page.rect.height)
            new_pg.show_pdf_page(new_pg.rect, form_doc, page_num)
            pw = page.rect.width

            for field in answered_fields:
                final = (field.get("final_answer") or "").strip()
                label = (field.get("label") or "").strip()
                if not final or not label:
                    continue

                rect = None
                for attempt in [label, label[:50], label[:30],
                                " ".join(label.split()[:3]),
                                " ".join(label.split()[:2])]:
                    attempt = attempt.strip(": _")
                    if len(attempt) < 3:
                        continue
                    rects = page.search_for(attempt)
                    if rects:
                        rect = rects[0]
                        break

                if not rect:
                    continue

                space_r  = pw - rect.x1 - 8
                max_ch   = int((pw - rect.x1 - 14) / 4.5)
                if space_r >= max(50, len(final) * 4.2):
                    pt = fitz.Point(rect.x1 + 6, rect.y1 - 1)
                else:
                    max_ch = int((pw - rect.x0 - 14) / 4.5)
                    pt = fitz.Point(rect.x0, rect.y1 + 11)

                try:
                    new_pg.insert_text(
                        pt, final[: max(10, max_ch)],
                        fontsize=FS, fontname="helv", color=BFITZ,
                    )
                except Exception:
                    pass

        elif PIL_AVAILABLE and TESSERACT_AVAILABLE:
            # ── Scanned form: PIL image overlay ─────────────────────────────
            pix = page.get_pixmap(dpi=DPI)
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            drw = PILImageDraw.Draw(img)

            try:
                ocr_data = pytesseract.image_to_data(
                    img, output_type=pytesseract.Output.DICT
                )
            except Exception:
                ocr_data = None

            for field in answered_fields:
                final = (field.get("final_answer") or "").strip()
                label = (field.get("label") or "").strip()
                if not final or not label:
                    continue

                pos = _find_label_in_ocr(ocr_data, label)
                if not pos:
                    continue

                fx, fy, fw, fh = pos
                ans_x = fx + fw + 6
                ans_y = fy
                if ans_x + len(final) * 8 > img.width - 10:
                    ans_x = fx
                    ans_y = fy + fh + 4

                try:
                    drw.text((ans_x, ans_y), final, fill=BLUE, font=_pil_font)
                except Exception:
                    pass

            # Convert annotated image back to a single-page PDF
            img_buf = io.BytesIO()
            img.save(img_buf, format="PDF", resolution=DPI)
            img_buf.seek(0)
            img_pdf = fitz.open(stream=img_buf.read(), filetype="pdf")
            out.insert_pdf(img_pdf)
            img_pdf.close()
            continue  # already appended — skip the out.new_page() path

        else:
            # No OCR / PIL — just copy the page unchanged
            new_pg = out.new_page(width=page.rect.width, height=page.rect.height)
            new_pg.show_pdf_page(new_pg.rect, form_doc, page_num)

    _append_summary_page(out, answered_fields)
    return out


# =============================================================================
# Session state
# =============================================================================

for _k, _v in [
    ("stage",         "upload"),
    ("analyses",      []),
    ("bundle_bytes",  None),
    ("bundle_fname",  "SAR_REDACTED_BUNDLE.pdf"),
    ("proc_summary",  []),
    ("play_sound",    None),
    ("sar_input_ver",    0),   # incremented by _reset() to wipe sidebar form fields
    # ── Detected patient demographics (auto-filled or manually entered) ───────
    ("pat_det_name",    ""),
    ("pat_det_dob",     ""),
    ("pat_det_nhs",     ""),
    ("pat_det_address", ""),
    ("pat_det_done",    False),  # True once detection has been run
    # ── Form filler ──────────────────────────────────────────────────────────
    ("tool_mode",         "sar"),        # "sar" | "form_filler" | "anon"
    ("ff_stage",          "ff_upload"),  # "ff_upload" | "ff_review" | "ff_export"
    ("ff_epr_text",       ""),
    ("ff_epr_docs",       []),           # list of fitz.Document (for context preview)
    ("ff_form_doc",       None),
    ("ff_form_text",      ""),
    ("ff_has_text_layer", False),
    ("ff_fields",         []),
    ("ff_filled_bytes",   None),
    ("ff_patient_name",   ""),
]:
    if _k not in st.session_state:
        st.session_state[_k] = _v


def _reset():
    """Reset SAR redaction session state and wipe all sidebar form fields."""
    keys_to_clear = [
        k for k in list(st.session_state.keys())
        if k in ("stage", "analyses", "bundle_bytes", "bundle_fname", "proc_summary")
        or k.startswith("editor_")
        or k.startswith("sec_")
        or k.startswith("esc_add_")
        or k.startswith("app_all_")
        or k.startswith("rej_all_")
        or k.startswith("approve_")
        or k.startswith("repl_")
        or k.startswith("esc_dec_")
        or k.startswith("page_excl_")
    ]
    for k in keys_to_clear:
        del st.session_state[k]
    # Increment version counter → all sidebar inputs with key=f"..._{sar_input_ver}"
    # become new widgets and default back to empty/None on the next rerun.
    st.session_state.sar_input_ver = st.session_state.get("sar_input_ver", 0) + 1
    # Clear detected patient demographics
    for _dk in ("pat_det_name", "pat_det_dob", "pat_det_nhs", "pat_det_address", "pat_det_done"):
        st.session_state[_dk] = "" if _dk != "pat_det_done" else False


def _reset_ff():
    """Reset form filler session state."""
    keys_to_clear = [
        k for k in list(st.session_state.keys())
        if k in ("ff_stage", "ff_epr_text", "ff_epr_docs", "ff_form_doc",
                 "ff_form_text", "ff_has_text_layer", "ff_fields",
                 "ff_filled_bytes", "ff_patient_name")
        or k.startswith("ff_ans_")
        or k.startswith("ff_appr_")
        or k.startswith("ff_man_")
    ]
    for k in keys_to_clear:
        del st.session_state[k]


# =============================================================================
# Page config
# =============================================================================

st.set_page_config(
    page_title="SAR Redaction Tool",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="expanded",
)
_inject_css()

# Consume pending sound (set before st.rerun() on stage transitions)
if st.session_state.get("play_sound"):
    _play_sound(st.session_state.play_sound)
    st.session_state.play_sound = None


# =============================================================================
# Sidebar — always visible
# =============================================================================

with st.sidebar:
    # Logo + title
    if _LOGO_B64:
        st.markdown(
            f'<div style="text-align:center;padding:12px 0 6px">'
            f'<img src="{_LOGO_B64}" style="max-height:64px;max-width:100%;border-radius:8px;'
            f'box-shadow:0 4px 16px rgba(0,94,184,.3)"></div>',
            unsafe_allow_html=True,
        )
    st.markdown(
        '<div style="text-align:center;font-size:1.1rem;font-weight:700;'
        'color:#0a2040;margin:6px 0 2px">NHS Clinical Tools</div>'
        '<div style="text-align:center;font-size:.72rem;color:#5a7090;margin-bottom:8px">'
        'SAR Redaction · Forms</div>',
        unsafe_allow_html=True,
    )

    # Badges
    st.markdown(
        '<div style="display:flex;gap:6px;justify-content:center;margin-bottom:10px;flex-wrap:wrap">'
        '<span class="badge-local">🔒 100% Local — No data leaves this PC</span>'
        '<span class="badge-test">⚠ Beta</span>'
        '</div>',
        unsafe_allow_html=True,
    )

    st.divider()

    # ── Mode selector ─────────────────────────────────────────────────────────
    _mode_choice = st.radio(
        "Select tool",
        ["🔒 SAR Redaction", "🕵️ Anonymise", "📋 Forms"],
        key="tool_mode_radio",
        label_visibility="collapsed",
    )
    if "Forms" in _mode_choice:
        tool_mode = "form_filler"
    elif "Anonymise" in _mode_choice:
        tool_mode = "anon"
    else:
        tool_mode = "sar"
    st.session_state.tool_mode = tool_mode

    st.divider()

    # ── Ollama — shared by both modes ─────────────────────────────────────────
    connected, available_models = check_ollama_connection()
    if connected and available_models:
        st.success(f"Ollama ✓ — {len(available_models)} model(s) available")
        selected_model = st.selectbox("LLM Model", available_models)
    elif connected:
        st.warning("Ollama running — no models loaded")
        selected_model = st.text_input("Model name", value="llama3")
    else:
        st.error("Ollama not running — start via run.bat")
        selected_model = st.text_input("Model name", value="llama3")

    st.divider()

    if tool_mode == "sar":
        # ── SAR-specific stage indicator ──────────────────────────────────────
        _stage_labels = {
            "upload": "① Upload & Analyse",
            "review": "② Review & Approve",
            "export": "③ Export",
        }
        _stage_colours = {"upload": "#4a9eff", "review": "#f0a030", "export": "#3cb86a"}
        _sc = _stage_colours.get(st.session_state.stage, "#888")
        st.markdown(
            f'<div style="background:#ffffff;border:1px solid #dde5f0;'
            f'border-radius:8px;padding:8px 12px;font-size:.82rem;font-weight:600;'
            f'color:{_sc}">▶ {_stage_labels.get(st.session_state.stage, "")}</div>',
            unsafe_allow_html=True,
        )
        st.divider()

        st.subheader("Settings")
        auto_classify = True   # always on
        show_debug = st.toggle(
            "Show LLM debug output",
            value=False,
            help="Show the raw LLM response for each document. Useful when no redactions appear.",
        )

        st.divider()
        st.subheader("SAR Details")
        _sv = st.session_state.sar_input_ver   # version suffix — incremented by _reset()
        sar_ref       = st.text_input("SAR reference / case ID",  placeholder="e.g. SAR-2024-001", key=f"sar_ref_{_sv}")
        patient_name  = st.text_input("Patient full name",         placeholder="e.g. John Smith",   key=f"patient_name_{_sv}")
        operator_name = st.text_input("Operator name",             placeholder="Your name / initials", key=f"operator_name_{_sv}")
        sar_date_input = st.date_input("SAR received date", value=None, format="DD/MM/YYYY", key=f"sar_date_{_sv}")
        if sar_date_input:
            _deadline  = sar_date_input + datetime.timedelta(days=30)
            _days_left = (_deadline - datetime.date.today()).days
            _colour    = "green" if _days_left > 10 else ("orange" if _days_left > 3 else "red")
            st.markdown(
                f"Deadline: **{_deadline.strftime('%d/%m/%Y')}**  \n"
                f":{_colour}[{_days_left} days remaining]"
            )

        st.divider()
        with st.expander("⚙ Custom redaction (this session only)", expanded=False):
            st.caption(
                "These settings apply only until the page is refreshed or the app is restarted. "
                "They do not affect the default behaviour."
            )
            extra_terms = st.text_area(
                "Extra terms to always redact",
                placeholder="e.g. Jane Smith\nAcme Care Ltd, Reference XYZ-99",
                height=90,
                key=f"extra_terms_{_sv}",
                help="Names, organisations or phrases that should always be redacted in this session. "
                     "Separate with commas or new lines.",
            )
            custom_instructions = st.text_area(
                "Custom LLM instructions",
                placeholder=(
                    "e.g. Also flag any medication names.\n"
                    "Treat all street addresses as third-party identifiers."
                ),
                height=110,
                key=f"custom_inst_{_sv}",
                help="Free-text instructions appended to the LLM redaction prompt for every document "
                     "in this session. Use this to fine-tune what the model flags.",
            )

        if st.session_state.stage != "upload":
            st.divider()
            if st.button("🔄 Start New SAR", use_container_width=True):
                _reset()
                st.rerun()

    elif tool_mode == "anon":
        # ── Anonymise sidebar ─────────────────────────────────────────────────
        st.markdown(
            '<div style="background:#ffffff;border:1px solid #dde5f0;'
            'border-radius:8px;padding:8px 12px;font-size:.82rem;color:#1a2a4a">'
            '<b>🕵️ Full Anonymisation</b><br>'
            'Removes <em>all</em> patient and person identifiers — suitable for sharing with '
            'MDU, insurers, researchers, or any external body where the patient must not be '
            'identifiable. Clinical content is preserved.'
            '</div>',
            unsafe_allow_html=True,
        )
        st.divider()
        if st.button("🔄 Start New", use_container_width=True):
            for _k in ("anon_results", "anon_zip_bytes"):
                st.session_state.pop(_k, None)
            st.rerun()

        # Set defaults for SAR-mode variables that won't be defined
        sar_ref = operator_name = patient_name = ""
        auto_classify = True
        show_debug    = False
        extra_terms   = ""
        custom_instructions = ""

    else:
        # ── Form filler sidebar ───────────────────────────────────────────────
        _ff_stage_labels = {
            "ff_upload": "① Load Records & Form",
            "ff_review": "② Review Answers",
            "ff_export": "③ Download Filled Form",
        }
        _ff_colours = {"ff_upload": "#4a9eff", "ff_review": "#f0a030", "ff_export": "#3cb86a"}
        _fsc = _ff_colours.get(st.session_state.ff_stage, "#888")
        st.markdown(
            f'<div style="background:#ffffff;border:1px solid #dde5f0;'
            f'border-radius:8px;padding:8px 12px;font-size:.82rem;font-weight:600;'
            f'color:{_fsc}">▶ {_ff_stage_labels.get(st.session_state.ff_stage, "")}</div>',
            unsafe_allow_html=True,
        )
        st.divider()

        # Patient name used by both modes — set via sidebar
        patient_name = st.text_input(
            "Patient full name",
            placeholder="e.g. John Smith",
            key="ff_patient_name_input",
        )

        if st.session_state.ff_stage != "ff_upload":
            st.divider()
            if st.button("🔄 Start New Form", use_container_width=True):
                _reset_ff()
                st.rerun()

        # Set defaults for SAR-mode variables that won't be defined
        sar_ref = operator_name = ""
        auto_classify = True
        show_debug    = False
        extra_terms   = ""
        custom_instructions = ""

    # Disclaimer — always shown
    st.divider()
    st.markdown(
        '<div class="sar-disclaimer">'
        '<b>⚠ Beta Software — No Warranty</b><br>'
        'This tool is in active development and provided for evaluation only. '
        'All AI suggestions must be independently reviewed by a qualified clinician '
        'or Information Governance professional before use. '
        'The authors accept no liability for errors, omissions, or misuse.<br><br>'
        '<b>🔒 Data Privacy</b><br>'
        'All processing uses a local LLM running on this computer. '
        '<b>No document content, patient data, or metadata is transmitted over the internet</b> '
        'or shared with any third party.'
        '</div>',
        unsafe_allow_html=True,
    )


# =============================================================================
# Page header
# =============================================================================

_logo_tag = f'<img src="{_LOGO_B64}" alt="Logo">' if _LOGO_B64 else \
    '<div style="width:54px;height:54px;background:rgba(0,94,184,.5);border-radius:8px;display:flex;align-items:center;justify-content:center;font-size:1.6rem">🔒</div>'

if tool_mode == "sar":
    _header_title = "SAR Redaction Tool"
    _header_sub   = "NHS Subject Access Request · Multi-format bundle processor · UK GDPR / DPA 2018 / ICO compliant"
elif tool_mode == "anon":
    _header_title = "Full Anonymisation"
    _header_sub   = "Remove all patient & person identifiers · Suitable for MDU, insurers, researchers · 100% local AI"
else:
    _header_title = "Forms"
    _header_sub   = "Complete insurance & GP report forms automatically from patient records · 100% local AI"

st.markdown(f"""
<div class="sar-header">
  {_logo_tag}
  <div class="sar-header-text">
    <h1>{_header_title}</h1>
    <p>{_header_sub}</p>
  </div>
  <div style="margin-left:auto;display:flex;flex-direction:column;gap:6px;align-items:flex-end">
    <span class="badge-local">🔒 Fully Local — Zero data egress</span>
    <span class="badge-test">⚠ Beta — Not for live use without review</span>
  </div>
</div>
""", unsafe_allow_html=True)

if tool_mode == "sar":
    if not DOCX_AVAILABLE:
        st.warning("python-docx not installed — Word files unsupported. Run: `pip install python-docx`")
    if not TESSERACT_AVAILABLE:
        st.info("Tesseract not available — TIFF/image files included in bundle but text redaction is not applied to image-only pages.")


# =============================================================================
# SAR REDACTION — stages 1-3
# =============================================================================

if tool_mode == "sar" and st.session_state.stage == "upload":
    st.divider()
    st.subheader("① Upload Documents")
    st.caption(
        "Drag-and-drop files below, or paste a folder path to load an entire patient record folder.  "
        "Supported: **PDF · Word (.docx) · TIFF · RTF · TXT · ZIP**"
    )

    uploaded_files = st.file_uploader(
        "Browse or drop files (or a ZIP archive)",
        type=ACCEPTED_FORMATS,
        accept_multiple_files=True,
        label_visibility="collapsed",
        key=f"sar_uploader_{st.session_state.sar_input_ver}",
    )

    folder_path_input = st.text_input(
        "Or load all documents from a folder path:",
        placeholder=r"e.g. C:\Patient Records\John Smith",
        key=f"folder_path_{st.session_state.sar_input_ver}",
        help="Paste the full path to a folder on this computer. "
             "All supported files (PDF, DOCX, TIFF, RTF, TXT, ZIP) inside it will be included.",
    ).strip()

    # Preview folder contents
    _folder_files = []
    if folder_path_input:
        _fp = Path(folder_path_input)
        if _fp.is_dir():
            _folder_files = [
                f for f in sorted(_fp.iterdir())
                if f.is_file() and f.suffix.lower().lstrip(".") in _SUPPORTED_EXTS | {"zip"}
            ]
            if _folder_files:
                st.info(
                    f"📁 **{len(_folder_files)} file(s) found in folder:**  "
                    + "  ·  ".join(f.name for f in _folder_files[:12])
                    + ("…" if len(_folder_files) > 12 else "")
                )
            else:
                st.warning("No supported files found in that folder.")
        else:
            st.error(f"Folder not found: `{folder_path_input}`")

    _any_input = bool(uploaded_files) or bool(_folder_files)

    if _any_input:
        _n_uploaded = len(uploaded_files) if uploaded_files else 0
        _n_folder   = len(_folder_files)
        _n_total    = _n_uploaded + _n_folder
        _zip_count  = sum(1 for f in (uploaded_files or []) if f.name.lower().endswith(".zip"))
        _zip_count += sum(1 for f in _folder_files if f.suffix.lower() == ".zip")
        _summary    = []
        if _n_uploaded:
            _summary.append(f"{_n_uploaded} uploaded file(s)")
        if _n_folder:
            _summary.append(f"{_n_folder} folder file(s)")
        if _zip_count:
            _summary.append(f"{_zip_count} ZIP(s) will be extracted")
        st.info("**Ready:** " + "  ·  ".join(_summary))

        # ── Patient details detection panel ───────────────────────────────────
        st.divider()
        st.markdown("#### 👤 Step 1: Confirm Patient Details")
        st.caption(
            "These details will be **protected from redaction** across every document in the bundle. "
            "Click **Auto-detect** to extract them from the documents, then confirm or correct."
        )

        _sv2 = st.session_state.sar_input_ver   # version key for versioned inputs
        _det_c1, _det_c2 = st.columns([3, 1])
        with _det_c2:
            st.markdown("&nbsp;")   # vertical alignment spacer
            if st.button("🔍 Auto-detect", use_container_width=True,
                         help="Scan the first few pages of the uploaded documents to detect patient details automatically"):
                _sample_files = _collect_all_files(
                    st.session_state.get(f"sar_uploader_{_sv2}") or [],
                    folder_path_input,
                )[:5]
                _det_texts = []
                for _sf in _sample_files:
                    if getattr(_sf, "_zip_error", None):
                        continue
                    _, _dt, _, _ = ingest_file(_sf)
                    if (_dt or "").strip():
                        _det_texts.append(_dt)
                if _det_texts:
                    with st.spinner("🔍 Detecting patient details from documents…"):
                        _det = _detect_patient_details(_det_texts, selected_model)
                    if _det.get("name"):
                        st.session_state.pat_det_name    = _det.get("name",       "")
                        st.session_state.pat_det_dob     = _det.get("dob",        "")
                        st.session_state.pat_det_nhs     = _det.get("nhs_number", "")
                        st.session_state.pat_det_address = _det.get("address",    "")
                        st.session_state.pat_det_done    = True
                        st.rerun()
                    else:
                        st.warning("Could not detect patient details automatically. "
                                   "Please fill in the fields manually.")
                else:
                    st.warning("No readable document text found. Upload files first.")

        with _det_c1:
            _pdn = st.text_input(
                "Patient name",
                value=st.session_state.pat_det_name,
                key=f"pat_det_name_input_{_sv2}",
                placeholder="e.g. John Smith",
            )
            st.session_state.pat_det_name = _pdn

        _dc1, _dc2 = st.columns(2)
        with _dc1:
            _pdd = st.text_input(
                "Date of birth",
                value=st.session_state.pat_det_dob,
                key=f"pat_det_dob_input_{_sv2}",
                placeholder="e.g. 12/05/1975",
            )
            st.session_state.pat_det_dob = _pdd

            _pdn2 = st.text_input(
                "NHS number",
                value=st.session_state.pat_det_nhs,
                key=f"pat_det_nhs_input_{_sv2}",
                placeholder="e.g. 485 777 3456",
            )
            st.session_state.pat_det_nhs = _pdn2

        with _dc2:
            _pda = st.text_area(
                "Home address",
                value=st.session_state.pat_det_address,
                key=f"pat_det_addr_input_{_sv2}",
                height=100,
                placeholder="e.g. 12 High Street\nManchester\nM1 2AB",
            )
            st.session_state.pat_det_address = _pda

        # Confirmation banner — shown once any detail is filled
        _eff_name = (st.session_state.pat_det_name or patient_name or "").strip()
        _eff_dob  = st.session_state.pat_det_dob.strip()
        _eff_nhs  = st.session_state.pat_det_nhs.strip()
        _eff_addr = st.session_state.pat_det_address.strip()

        if _eff_name or _eff_dob or _eff_nhs or _eff_addr:
            _shield_parts = []
            if _eff_name: _shield_parts.append(f"**Name:** {_eff_name}")
            if _eff_dob:  _shield_parts.append(f"**DOB:** {_eff_dob}")
            if _eff_nhs:  _shield_parts.append(f"**NHS No:** {_eff_nhs}")
            if _eff_addr: _shield_parts.append(f"**Address:** {_eff_addr.splitlines()[0]}")
            st.success(
                "🛡️ **Protected from redaction across all documents:**  \n"
                + "  ·  ".join(_shield_parts)
            )

        st.divider()

        # Require at least a name before allowing analysis to proceed
        _pname_missing = not _eff_name
        if _pname_missing:
            st.error(
                "⚠️ **Patient name is required before analysis can begin.** "
                "Click **Auto-detect** above, or enter the name manually in the Patient name field."
            )

        if st.button(
            "▶ Analyse Documents",
            type="primary",
            use_container_width=True,
            disabled=_pname_missing,
        ):
            # Carry effective name forward: detected field takes precedence over sidebar
            if _eff_name and not patient_name.strip():
                patient_name = _eff_name
            # Show UI immediately so the user knows something is happening
            prog   = st.progress(0.0, text="⏳ Preparing files…")
            status = st.empty()
            status.info("📂 Collecting and unpacking files — please wait…")

            # Collect + expand ZIPs
            all_files = _collect_all_files(uploaded_files, folder_path_input)
            if not all_files:
                st.error("No files to process after expansion.")
                st.stop()

            analyses          = []
            _model_warm_shown = False   # track whether we've warned about model warm-up
            _sar_outer_start  = time.time()

            for i, ufile in enumerate(all_files):
                prog.progress(
                    i / len(all_files),
                    text=(
                        f"File {i + 1}/{len(all_files)}: {ufile.name}"
                        + _timing_suffix(_sar_outer_start, i, len(all_files))
                    ),
                )

                # Surface ZIP extraction errors
                zip_err = getattr(ufile, "_zip_error", None)
                if zip_err:
                    analyses.append({
                        "filename":            ufile.name,
                        "section":             "Miscellaneous",
                        "doc":                 None,
                        "text":                "",
                        "has_text":            False,
                        "error":               f"ZIP extraction failed: {zip_err}",
                        "proposed_redactions": [],
                        "escalations":         [],
                        "llm_raw":             "",
                        "llm_parse_ok":        False,
                        "doc_date":            datetime.date.min,
                    })
                    continue

                # Ingest
                status.markdown(f"📥 **Ingesting** `{ufile.name}`…")
                fitz_doc, text, err, ocr_info = ingest_file(ufile)

                if err or fitz_doc is None:
                    analyses.append({
                        "filename":            ufile.name,
                        "section":             "Miscellaneous",
                        "doc":                 None,
                        "text":                "",
                        "has_text":            False,
                        "error":               err,
                        "ocr_info":            ocr_info,
                        "proposed_redactions": [],
                        "escalations":         [],
                        "llm_raw":             "",
                        "llm_parse_ok":        False,
                        "doc_date":            datetime.date.min,
                    })
                    continue

                has_text = bool(text.strip())

                # Extract document date for bundle ordering
                doc_date = _extract_document_date(text) if has_text else datetime.date.min

                # Classify
                section = "Miscellaneous"
                if auto_classify and has_text:
                    if not _model_warm_shown:
                        status.info(
                            f"🔍 **Classifying** `{ufile.name}`…  \n"
                            "⏳ *First call loads the AI model into memory — "
                            "this can take up to 60 seconds. The app is working.*"
                        )
                        _model_warm_shown = True
                    else:
                        status.markdown(f"🔍 **Classifying** `{ufile.name}`…")
                    section = classify_document(text, selected_model)

                # LLM analysis
                llm_raw      = ""
                llm_parse_ok = True
                proposed     = []
                escalations  = []

                if has_text:
                    # ── Resolve effective patient name ────────────────────────
                    # Use the operator-entered name if available; otherwise try to
                    # auto-detect it from the filename or document header so the LLM
                    # and name-expansion filter know whose data must NOT be redacted.
                    # Priority: sidebar name → detected name → filename/header heuristic
                    effective_pname = (
                        patient_name.strip()
                        or st.session_state.get("pat_det_name", "").strip()
                    )
                    if not effective_pname:
                        effective_pname = _detect_patient_name(ufile.name, text)
                        if effective_pname:
                            status.info(
                                f"🔍 **Patient name auto-detected** from filename/header: "
                                f"**{effective_pname}** — will be protected from redaction. "
                                f"Enter the name in the sidebar to suppress this message."
                            )

                    if not _model_warm_shown:
                        status.info(
                            f"🤖 **Analysing** `{ufile.name}`…  \n"
                            "⏳ *First call loads the AI model into memory — "
                            "this can take up to 60 seconds. The app is working.*"
                        )
                        _model_warm_shown = True
                    else:
                        status.markdown(f"🤖 **Analysing** `{ufile.name}` for SAR redactions…")
                    result, llm_raw = llm_analyse_document(
                        text, selected_model, effective_pname,
                        status_cb=lambda msg: status.markdown(f"🤖 **`{ufile.name}`** — {msg}"),
                        extra_redactions=extra_terms,
                        custom_instructions=custom_instructions,
                        patient_dob     = st.session_state.get("pat_det_dob",     ""),
                        patient_nhs     = st.session_state.get("pat_det_nhs",     ""),
                        patient_address = st.session_state.get("pat_det_address", ""),
                    )
                    llm_parse_ok    = result.get("parse_ok", False)
                    raw_prop        = result.get("proposed_redactions", [])
                    escalations     = result.get("escalations", [])

                    # Deduplicate proposed redactions by text
                    seen = set()
                    for item in raw_prop:
                        t = (item.get("text") or "").strip()
                        if t and t not in seen:
                            item["text"]     = t
                            item["approved"] = True
                            seen.add(t)
                            proposed.append(item)

                    # Remove any item that matches the patient's own name or any
                    # individual token within it (prevents shared surnames being redacted).
                    if effective_pname:
                        pn_lower   = effective_pname.lower()
                        pn_tokens  = {
                            tok.lower() for tok in pn_lower.split() if len(tok) >= 3
                        }
                        proposed = [
                            p for p in proposed
                            if pn_lower not in p["text"].lower()
                            and p["text"].lower() not in pn_lower
                            and p["text"].lower() not in pn_tokens
                        ]
                        escalations = [
                            e for e in escalations
                            if pn_lower not in (e.get("text") or "").lower()
                        ]

                analyses.append({
                    "filename":            ufile.name,
                    "section":             section,
                    "doc":                 fitz_doc,
                    "text":                text,
                    "has_text":            has_text,
                    "error":               "",
                    "ocr_info":            ocr_info,
                    "proposed_redactions": proposed,
                    "escalations":         escalations,
                    "llm_raw":             llm_raw,
                    "llm_parse_ok":        llm_parse_ok,
                    "doc_date":            doc_date,
                })

            prog.progress(1.0, text="Analysis complete")
            status.empty()
            st.session_state.analyses   = analyses
            # Clear stale approval/replacement/escalation-decision keys
            for _k in list(st.session_state.keys()):
                if _k.startswith(("approve_", "repl_", "esc_dec_", "page_excl_")):
                    del st.session_state[_k]
            st.session_state.stage      = "review"
            st.session_state.play_sound = "chime"
            st.rerun()

    else:
        st.info("Upload one or more documents above, or enter a folder path to begin.")
        with st.expander("Accepted formats"):
            st.markdown("""
| Format | Extensions | Notes |
|--------|-----------|-------|
| PDF | `.pdf` | Full text redaction; image-only pages included without OCR unless Tesseract is installed |
| Word | `.docx` `.doc` | Requires `python-docx` |
| TIFF | `.tiff` `.tif` | Requires Tesseract for OCR redaction |
| RTF | `.rtf` | Requires `striprtf`; falls back to basic text stripping |
| Plain text | `.txt` | Full redaction support |
| ZIP archive | `.zip` | All supported files inside will be extracted and processed automatically |
| **Folder** | — | Paste a folder path in the text box above to load all files in one go |
            """)


# =============================================================================
# STAGE 2 — REVIEW & APPROVE
# =============================================================================

elif tool_mode == "sar" and st.session_state.stage == "review":
    st.divider()

    _analyses   = st.session_state.analyses
    _total_prop = sum(len(a["proposed_redactions"]) for a in _analyses)
    _total_esc  = sum(len(a["escalations"])         for a in _analyses)

    m1, m2, m3 = st.columns(3)
    m1.metric("Documents",           len(_analyses))
    m2.metric("Proposed redactions", _total_prop)
    m3.metric("Escalations",         _total_esc)

    if _total_esc > 0:
        st.error(
            f"⚠️ **{_total_esc} escalation(s)** require clinical or IG review before release. "
            "See the highlighted sections below — these are NOT automatically redacted."
        )

    if _total_prop == 0 and _total_esc == 0:
        st.warning(
            "The LLM did not propose any redactions or escalations across any document.  \n\n"
            "**If this seems wrong**, enable **Show LLM debug output** in the sidebar to see "
            "the raw response from the model. Common causes:  \n"
            "• Model not outputting JSON — try **llama3.1** or **qwen2.5**  \n"
            "• Documents genuinely contain no third-party or sensitive content  \n\n"
            "You can still build and download the bundle using the button below."
        )

    # ── Escalation decision options (defined here so progress counter can use them) ──
    _DEC_OPTS = [
        "⚠️ Awaiting decision",
        "🔴 Redact this passage",
        "✅ Release as-is",
    ]

    _ESC_TAG_DOT = {
        "CLINICIAN_CONTEXT_AMBIGUOUS": "🟡",
        "SAFEGUARDING_RISK":           "🔴",
        "DOMESTIC_ABUSE_CONTEXT":      "🔴",
        "CHILD_PROTECTION":            "🔴",
        "SERIOUS_HARM_RISK":           "🔴",
        "SENSITIVE_CLINICAL_OPINION":  "🟠",
        "LEGAL_PRIVILEGE":             "🟠",
        "DPA_SCHEDULE3_EXEMPTION":     "🟠",
    }

    # ── Phase 1 — Build canonical maps (runs every rerender, before any UI) ──
    # Uses _normalise_for_dedup() so that variant forms of the same identifier
    # (e.g. "John Smith" / "Mr John Smith" / "Smith, John", or "15/05/1975" /
    # "15 May 1975") are treated as a single item and only shown once.

    _unique_items_ordered = []        # [(text_lower, canon_si, canon_ri, rr_dict)]
    _vk_to_canonical = {}             # variant-key  → (canon_si, canon_ri)
    _tl_to_canonical = {}             # text_lower   → (canon_si, canon_ri)  [exact fallback]
    _vk_to_docs = {}                  # variant-key  → [filename, ...]

    for _si, _sa in enumerate(_analyses):
        for _ri, _rr in enumerate(_sa["proposed_redactions"]):
            _raw = (_rr.get("text") or "").strip()
            _tl  = _raw.lower()
            _vk  = _normalise_for_dedup(_raw)
            if not _vk:
                continue
            if _vk not in _vk_to_canonical:
                _unique_items_ordered.append((_tl, _si, _ri, _rr))
                _vk_to_canonical[_vk] = (_si, _ri)
                _tl_to_canonical[_tl] = (_si, _ri)
                _vk_to_docs[_vk] = [_sa["filename"]]
            else:
                # Variant of an already-seen identifier → point to same canonical
                _tl_to_canonical[_tl] = _vk_to_canonical[_vk]
                if _sa["filename"] not in _vk_to_docs[_vk]:
                    _vk_to_docs[_vk].append(_sa["filename"])

    # Init session_state for CANONICAL instances only
    for _tl, _csi, _cri, _crr in _unique_items_ordered:
        _ak = f"approve_{_csi}_{_cri}"
        _rk = f"repl_{_csi}_{_cri}"
        if _ak not in st.session_state:
            st.session_state[_ak] = _crr.get("approved", True)
        if _rk not in st.session_state:
            st.session_state[_rk] = _crr.get("replacement", "[REDACTED]") or "[REDACTED]"

    # Escalations canonical map
    _esc_unique_ordered = []       # [(etl, canon_si, canon_ei, esc_dict, [filenames])]
    _etl_to_canonical = {}
    _etl_to_docs = {}
    for _si, _sa in enumerate(_analyses):
        for _ei, _esc in enumerate(_sa["escalations"]):
            _etl = (_esc.get("text") or "").strip().lower()
            if not _etl:
                continue
            if _etl not in _etl_to_canonical:
                _esc_unique_ordered.append((_etl, _si, _ei, _esc, [_sa["filename"]]))
                _etl_to_canonical[_etl] = (_si, _ei)
                _etl_to_docs[_etl] = [_sa["filename"]]
            elif _sa["filename"] not in _etl_to_docs[_etl]:
                _etl_to_docs[_etl].append(_sa["filename"])

    for _etl, _esi, _eei, _esc, _edocs in _esc_unique_ordered:
        _dk = f"esc_dec_{_esi}_{_eei}"
        if _dk not in st.session_state:
            st.session_state[_dk] = _DEC_OPTS[0]

    # ── Phase 2 — Propagate decisions ──────────────────────────────────────────

    def _propagate_decisions():
        # Sync canonical session_state → every rr in _analyses
        for _sa in _analyses:
            for _rr in _sa["proposed_redactions"]:
                _raw = (_rr.get("text") or "").strip()
                _tl  = _raw.lower()
                _vk  = _normalise_for_dedup(_raw)
                # Try exact match first, fall back to variant key
                _canon = _tl_to_canonical.get(_tl) or _vk_to_canonical.get(_vk)
                if _canon:
                    _csi, _cri = _canon
                    _rr["approved"]    = st.session_state.get(f"approve_{_csi}_{_cri}", True)
                    _rr["replacement"] = st.session_state.get(f"repl_{_csi}_{_cri}", "[REDACTED]") or "[REDACTED]"

        # Apply escalation decisions → proposed_redactions for every doc
        for _si, _sa in enumerate(_analyses):
            for _esc in _sa["escalations"]:
                _etl   = (_esc.get("text") or "").strip().lower()
                _canon = _etl_to_canonical.get(_etl)
                if not _canon:
                    continue
                _csi, _cei = _canon
                _dec       = st.session_state.get(f"esc_dec_{_csi}_{_cei}", _DEC_OPTS[0])
                _flag      = _esc.get("text", "")
                _etag      = _esc.get("tag", "")
                _ereason   = _esc.get("reason", "")
                _econtext  = _esc.get("context", "")
                _existing  = {r["text"] for r in _sa["proposed_redactions"]}
                _tlbl      = REDACTION_TAGS.get(_etag, {}).get("label", _etag)
                if _dec == _DEC_OPTS[1] and _flag and _flag not in _existing:
                    _sa["proposed_redactions"].append({
                        "text": _flag, "tag": _etag, "reason": _ereason,
                        "replacement": f"[REDACTED – {_ereason or _tlbl}]",
                        "context": _econtext, "approved": True,
                    })
                elif _dec == _DEC_OPTS[2]:
                    _sa["proposed_redactions"] = [
                        r for r in _sa["proposed_redactions"]
                        if not (r.get("text") == _flag and r.get("tag") == _etag)
                    ]

        # Re-sync after escalation additions
        for _sa in _analyses:
            for _rr in _sa["proposed_redactions"]:
                _raw = (_rr.get("text") or "").strip()
                _tl  = _raw.lower()
                _vk  = _normalise_for_dedup(_raw)
                _canon = _tl_to_canonical.get(_tl) or _vk_to_canonical.get(_vk)
                if _canon:
                    _csi, _cri = _canon
                    _rr["approved"]    = st.session_state.get(f"approve_{_csi}_{_cri}", True)
                    _rr["replacement"] = st.session_state.get(f"repl_{_csi}_{_cri}", "[REDACTED]") or "[REDACTED]"
                else:
                    _rr.setdefault("approved", True)

    _propagate_decisions()

    # ── Phase 3 — Progress banner (counts based on unique items) ───────────────

    _n_unique_total    = len(_unique_items_ordered)
    _n_unique_approved = sum(1 for _, _csi, _cri, _ in _unique_items_ordered
                             if st.session_state.get(f"approve_{_csi}_{_cri}", True))
    _n_esc_unique      = len(_esc_unique_ordered)
    _n_esc_decided     = sum(1 for _, _esi, _eei, _, _ in _esc_unique_ordered
                             if st.session_state.get(f"esc_dec_{_esi}_{_eei}", _DEC_OPTS[0]) != _DEC_OPTS[0])
    _n_esc_pending     = _n_esc_unique - _n_esc_decided
    _all_ready         = _n_esc_pending == 0

    # Progress bar HTML banner
    _esc_pct  = int(_n_esc_decided / _n_esc_unique * 100) if _n_esc_unique > 0 else 100
    _bar_col  = "#3cb86a" if _all_ready else "#f0a030"
    st.markdown(
        f"""
        <div style="background:#ffffff;border:1px solid #dde5f0;
                    border-radius:14px;padding:16px 22px;margin-bottom:16px;">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">
            <span style="color:#1a3060;font-weight:600;font-size:1rem">
              📋 Review Progress
            </span>
            <span style="color:#5a7090;font-size:.85rem">
              Scroll down to review all suggestions
            </span>
          </div>
          <div style="display:flex;gap:32px;margin-bottom:12px">
            <div>
              <div style="color:#6a80a0;font-size:.75rem;text-transform:uppercase;letter-spacing:.5px">Redactions approved</div>
              <div style="color:#0a2040;font-size:1.3rem;font-weight:700">{_n_unique_approved} <span style="color:#7a90b0;font-size:.9rem">/ {_n_unique_total}</span></div>
            </div>
            <div>
              <div style="color:#6a80a0;font-size:.75rem;text-transform:uppercase;letter-spacing:.5px">Escalations decided</div>
              <div style="color:#0a2040;font-size:1.3rem;font-weight:700">{_n_esc_decided} <span style="color:#7a90b0;font-size:.9rem">/ {_n_esc_unique}</span></div>
            </div>
            <div style="margin-left:auto;display:flex;align-items:center">
              {"<span style='background:#e6f9ec;border:1px solid #52c271;border-radius:20px;padding:4px 14px;color:#1a7a36;font-weight:600'>✅ Ready to build</span>" if _all_ready else f"<span style='background:#fff8e6;border:1px solid #f0a030;border-radius:20px;padding:4px 14px;color:#8a5a00;font-weight:600'>⚠️ {_n_esc_pending} escalation{'s' if _n_esc_pending != 1 else ''} pending</span>"}
            </div>
          </div>
          <div style="background:#e8edf5;border-radius:6px;height:14px;overflow:hidden">
            <div style="width:{_esc_pct}%;height:100%;background:linear-gradient(90deg,{_bar_col},{_bar_col}cc);
                        border-radius:6px;transition:width .4s ease"></div>
          </div>
          <div style="color:#5a7090;font-size:.73rem;margin-top:5px;text-align:right">
            {"All escalations resolved — click Apply button below to build the bundle" if _all_ready else f"Scroll down · find sections marked 🔴 · make a decision for each escalation"}
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ── Phase 4 — STEP 1: Global unique identifiers table ─────────────────────

    st.markdown("### Step 1 — Bundle-wide Redaction Decisions")
    st.caption(
        "Each identifier appears once regardless of how many documents contain it. "
        "Your ✓/✗ decision applies to every document in the bundle automatically."
    )

    _ga1, _ga2, _ = st.columns([1, 1, 5])
    if _ga1.button("✅ Approve All", key="app_all_global"):
        for _tl, _csi, _cri, _ in _unique_items_ordered:
            st.session_state[f"approve_{_csi}_{_cri}"] = True
        st.rerun()
    if _ga2.button("❌ Reject All", key="rej_all_global"):
        for _tl, _csi, _cri, _ in _unique_items_ordered:
            st.session_state[f"approve_{_csi}_{_cri}"] = False
        st.rerun()

    # Group by tag
    _items_by_tag: dict[str, list] = {}
    for _tl, _csi, _cri, _crr in _unique_items_ordered:
        _tag = _crr.get("tag", "")
        _items_by_tag.setdefault(_tag, []).append((_tl, _csi, _cri, _crr))

    for _tag, _tag_items in _items_by_tag.items():
        _tag_info  = REDACTION_TAGS.get(_tag, {})
        _tag_label = _tag_info.get("label", _tag) if _tag else "Untagged"
        _n_items   = len(_tag_items)
        _n_approved_tag = sum(
            1 for _, _csi, _cri, _ in _tag_items
            if st.session_state.get(f"approve_{_csi}_{_cri}", True)
        )
        with st.expander(
            f"**{_tag_label}** — {_n_items} unique · {_n_approved_tag} approved",
            expanded=True,
        ):
            _ca1, _ca2, _ = st.columns([1, 1, 5])
            if _ca1.button("✅ All", key=f"app_cat_{_tag}"):
                for _, _csi, _cri, _ in _tag_items:
                    st.session_state[f"approve_{_csi}_{_cri}"] = True
                st.rerun()
            if _ca2.button("❌ None", key=f"rej_cat_{_tag}"):
                for _, _csi, _cri, _ in _tag_items:
                    st.session_state[f"approve_{_csi}_{_cri}"] = False
                st.rerun()

            _hdr1, _hdr2, _hdr3, _hdr4 = st.columns([1, 4, 2, 2])
            _hdr1.markdown("<span style='color:#6a80a0;font-size:.78rem'>APPROVE</span>", unsafe_allow_html=True)
            _hdr2.markdown("<span style='color:#6a80a0;font-size:.78rem'>TEXT / REASON</span>", unsafe_allow_html=True)
            _hdr3.markdown("<span style='color:#6a80a0;font-size:.78rem'>REPLACEMENT LABEL</span>", unsafe_allow_html=True)
            _hdr4.markdown("<span style='color:#6a80a0;font-size:.78rem'>FOUND IN</span>", unsafe_allow_html=True)
            st.markdown("<hr style='margin:4px 0 8px;border-color:#dde5f0'>", unsafe_allow_html=True)

            for _tl, _csi, _cri, _crr in _tag_items:
                _ak       = f"approve_{_csi}_{_cri}"
                _rk       = f"repl_{_csi}_{_cri}"
                _text     = _crr.get("text", "")
                _reason   = (_crr.get("reason") or "")[:140]
                _approved = st.session_state.get(_ak, True)
                _docs_for_tl = _tl_to_docs.get(_tl, [])

                _txt_style = (
                    "background:rgba(220,50,50,.10);border-radius:4px;"
                    "padding:2px 6px;color:#c0001a;font-family:monospace;font-size:.85rem"
                    if _approved else
                    "background:#f0f4f8;border-radius:4px;"
                    "padding:2px 6px;color:#6a80a0;font-family:monospace;font-size:.85rem;"
                    "text-decoration:line-through"
                )

                c_chk, c_txt, c_repl, c_found = st.columns([1, 4, 2, 2])
                with c_chk:
                    _cb_val = st.checkbox(
                        "✓",
                        key=_ak,
                        help="Tick to redact · untick to keep",
                        label_visibility="collapsed",
                    )
                    st.caption("✓ Redact" if _cb_val else "✗ Keep")
                with c_txt:
                    st.markdown(
                        f"<span style='{_txt_style}'>{_text}</span>"
                        + (f"<br><span style='color:#3a6090;font-size:.78rem'>{_reason}</span>" if _reason else ""),
                        unsafe_allow_html=True,
                    )
                with c_repl:
                    st.text_input(
                        "Replacement",
                        key=_rk,
                        label_visibility="collapsed",
                        placeholder="[REDACTED]",
                    )
                with c_found:
                    if len(_docs_for_tl) == 1:
                        st.caption(f"📄 {_docs_for_tl[0]}")
                    else:
                        _doc_list_md = "\n".join(f"- {_d}" for _d in _docs_for_tl)
                        st.caption(f"📄 {len(_docs_for_tl)} documents:\n{_doc_list_md}")
                st.markdown(
                    "<hr style='margin:2px 0 6px;border-color:#e8edf5'>",
                    unsafe_allow_html=True,
                )

    # ── Phase 5 — STEP 2: Escalations (unique, deduplicated) ──────────────────

    st.markdown("### Step 2 — Escalations · Require Clinical / IG Review")
    st.info(
        "Each unique passage is shown once; your decision applies to every document where it appears."
    )

    if _esc_unique_ordered:
        _n_esc_total_unique = len(_esc_unique_ordered)
        for _esc_idx, (_etl, _esi, _eei, _esc_item, _esc_docs) in enumerate(_esc_unique_ordered):
            _tag      = _esc_item.get("tag", "")
            _tag_info = REDACTION_TAGS.get(_tag, {})
            _label    = _tag_info.get("label", _tag)
            _desc     = _tag_info.get("desc", "")
            _dot      = _ESC_TAG_DOT.get(_tag, "🟡")
            _flagged  = _esc_item.get("text", "")
            _reason   = _esc_item.get("reason", "")
            _context  = _esc_item.get("context", "")
            _dec_key  = f"esc_dec_{_esi}_{_eei}"

            if len(_esc_docs) > 1:
                st.warning(
                    f"Found in **{len(_esc_docs)} documents** — decision applies to all: "
                    + ", ".join(_esc_docs)
                )

            with st.container(border=True):
                # ── Card header ──────────────────────────────────────────────
                hcol1, hcol2 = st.columns([5, 1])
                with hcol1:
                    st.markdown(f"##### {_dot} {_label}")
                with hcol2:
                    st.caption(f"*{_esc_idx + 1} / {_n_esc_total_unique}*")

                if _desc:
                    st.caption(f"📋 **Category:** {_desc}")

                st.markdown("---")

                # ── Flagged text + AI reasoning ──────────────────────────────
                fcol, rcol = st.columns([3, 2])
                with fcol:
                    st.markdown("**🚩 Flagged passage**")
                    st.code(_flagged or "*(no text captured)*", language=None)
                with rcol:
                    st.markdown("**🤖 AI reasoning**")
                    if _reason:
                        st.markdown(
                            f"<div style='background:#e8f0fc;"
                            f"border-left:3px solid #005EB8;"
                            f"border-radius:6px;padding:10px 14px;"
                            f"color:#1a2a4a;font-size:.88rem'>"
                            f"{_reason}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.caption("*(no reasoning recorded)*")

                # ── Inline document context preview ──────────────────────────
                # Use first document that contains this escalation for context
                _ctx_found = False
                for _sa in _analyses:
                    if _sa["filename"] in _esc_docs:
                        _doc_text = _sa.get("text", "")
                        if _doc_text and _flagged:
                            _idx = _doc_text.find(_flagged)
                            if _idx == -1:
                                _probe = _flagged[:50].lower()
                                _idx   = _doc_text.lower().find(_probe)
                            if _idx != -1:
                                _start   = max(0, _idx - 350)
                                _end     = min(len(_doc_text), _idx + len(_flagged) + 350)
                                _snippet = _doc_text[_start:_end]
                                if _start > 0:
                                    _snippet = "…" + _snippet
                                if _end < len(_doc_text):
                                    _snippet = _snippet + "…"
                                with st.expander("📄 Show surrounding document context", expanded=False):
                                    st.caption(
                                        "The passage below is extracted from the original document "
                                        "around the flagged text, to help you judge the full context."
                                    )
                                    st.code(_snippet, language=None)
                                _ctx_found = True
                                break

                if not _ctx_found and _context:
                    with st.expander("📄 Context (from AI)", expanded=False):
                        st.code(_context, language=None)

                st.markdown("---")

                # ── Decision control ─────────────────────────────────────────
                _decision = st.radio(
                    "**Your decision:**",
                    options=_DEC_OPTS,
                    key=_dec_key,
                    horizontal=True,
                    help=(
                        "• Awaiting decision — no action yet  \n"
                        "• Redact this passage — adds to the redaction list below  \n"
                        "• Release as-is — text will appear in the final bundle"
                    ),
                )
                if _decision == _DEC_OPTS[1]:
                    _n_aff = len(_esc_docs)
                    st.error(
                        f"This passage will be **redacted** in all {_n_aff} document(s)."
                        if _n_aff > 1 else
                        "This passage will be **redacted** in the final bundle."
                    )
                elif _decision == _DEC_OPTS[2]:
                    st.success("This passage will be **released** as-is.")

        _propagate_decisions()
    else:
        st.success("No escalations.")

    # ── Phase 6 — STEP 3: Document Details ────────────────────────────────────

    st.markdown("### Step 3 — Document Details")
    st.caption(
        "Verify section assignments and optionally remove whole pages. "
        "All redaction decisions from Steps 1 & 2 are already applied."
    )

    for i, analysis in enumerate(_analyses):
        fname   = analysis["filename"]
        has_err = bool(analysis.get("error"))
        n_esc   = len(analysis["escalations"])

        _n_red  = sum(1 for r in analysis["proposed_redactions"] if r.get("approved", True))

        _doc_date = analysis.get("doc_date")
        _date_str = (
            _doc_date.strftime("%d/%m/%Y")
            if _doc_date and _doc_date != datetime.date.min
            else "date unknown"
        )

        # Determine pending escalations for this doc (any not decided)
        _doc_esc_pending = sum(
            1 for _etl, _esi, _eei, _esc_item, _esc_docs in _esc_unique_ordered
            if fname in _esc_docs
            and st.session_state.get(f"esc_dec_{_esi}_{_eei}", _DEC_OPTS[0]) == _DEC_OPTS[0]
        )

        icon = "❌" if has_err else ("🔴" if _doc_esc_pending > 0 else ("✏️" if _n_red > 0 else "✅"))

        with st.expander(
            f"{icon} {fname} — {_n_red} redactions · {_date_str} · section: {analysis['section']}",
            expanded=has_err,
        ):
            if has_err:
                st.error(f"Failed to process: {analysis['error']}")
                continue

            # LLM debug output
            if show_debug:
                n_red_full   = len(analysis["proposed_redactions"])
                llm_raw      = analysis.get("llm_raw", "")
                llm_parse_ok = analysis.get("llm_parse_ok", True)
                n_chunks     = analysis.get("chunks_analysed", 1)
                chars_total  = analysis.get("chars_total", 0)
                with st.expander(
                    "🔧 Raw LLM response",
                    expanded=(not llm_parse_ok or (n_red_full == 0 and n_esc == 0)),
                ):
                    if chars_total:
                        covered = min(n_chunks * 6000, chars_total)
                        pct = int(covered / chars_total * 100)
                        st.caption(
                            f"Document: {chars_total:,} chars · "
                            f"Analysed: {n_chunks} chunk(s) · "
                            f"Coverage: ~{pct}%"
                        )
                    if not llm_raw:
                        st.info("No LLM response recorded (file had no extractable text).")
                    elif not llm_parse_ok:
                        st.warning(
                            "⚠️ The LLM response could not be parsed as JSON.  \n"
                            "The model may not be following the output format instruction.  \n"
                            "Try **qwen2.5:14b** for reliable JSON output."
                        )
                        st.code(llm_raw, language=None)
                    else:
                        if n_red_full == 0 and n_esc == 0:
                            st.success(
                                "✅ JSON parsed successfully. "
                                "The LLM found no third-party or sensitive content "
                                "requiring redaction in this document."
                            )
                        st.code(llm_raw, language="json")

            # OCR / text extraction status badge
            _ocr = analysis.get("ocr_info", "")
            if _ocr:
                _ocr_lower = _ocr.lower()
                if "tesseract ocr" in _ocr_lower and "failed" not in _ocr_lower and "no text" not in _ocr_lower and "not available" not in _ocr_lower:
                    st.success(f"🔎 **Text extraction:** {_ocr}")
                elif "not available" in _ocr_lower or "failed" in _ocr_lower or "no text extracted" in _ocr_lower:
                    st.error(f"⚠️ **Text extraction:** {_ocr}")
                elif "scanned" in _ocr_lower and "not available" not in _ocr_lower:
                    st.info(f"🔎 **Text extraction:** {_ocr}")
                else:
                    st.caption(f"📄 **Text extraction:** {_ocr}")

            if not analysis["has_text"]:
                st.warning(
                    "No text layer detected. This file will be included in the bundle "
                    "but automated redaction cannot be applied. Manual review recommended."
                )

            # Section override
            new_sec = st.selectbox(
                "Bundle section",
                SECTION_ORDER,
                index=SECTION_ORDER.index(analysis["section"])
                      if analysis["section"] in SECTION_ORDER else 0,
                key=f"sec_{i}",
            )
            analysis["section"] = new_sec
            st.divider()

            # ── Redaction summary ─────────────────────────────────────────────
            _approved_texts = [
                r.get("text", "") for r in analysis["proposed_redactions"]
                if r.get("approved", True)
            ]
            if _approved_texts:
                _shown = _approved_texts[:12]
                _summary_parts = " ".join(
                    f"<code style='background:#fdecea;padding:1px 5px;"
                    f"border-radius:3px;color:#c0001a;font-size:.8rem'>{t}</code>"
                    for t in _shown
                )
                if len(_approved_texts) > 12:
                    _summary_parts += f" <span style='color:#7a90b0;font-size:.8rem'>+{len(_approved_texts) - 12} more</span>"
                st.markdown(
                    f"<div style='margin-bottom:8px'><b>Approved redactions:</b> {_summary_parts}</div>",
                    unsafe_allow_html=True,
                )

            # ── Context preview panel ─────────────────────────────────────────
            _all_items = analysis["proposed_redactions"] + [
                {**e, "_is_esc": True} for e in analysis.get("escalations", [])
            ]
            if _all_items and analysis.get("doc") and PIL_AVAILABLE:
                st.markdown("#### 🔍 Context Preview")
                _prev_options = ["— select a redaction to preview —"] + [
                    f"{'⚠ ' if r.get('_is_esc') else ''}[{REDACTION_TAGS.get(r.get('tag',''),{}).get('label', r.get('tag',''))}]  {r.get('text','')[:70]}"
                    for r in _all_items
                ]
                _prev_sel = st.selectbox(
                    "Preview term in document:",
                    _prev_options,
                    key=f"prev_sel_{i}",
                    label_visibility="collapsed",
                )
                if _prev_sel and _prev_sel != _prev_options[0]:
                    _sel_idx  = _prev_options.index(_prev_sel) - 1
                    _sel_item = _all_items[_sel_idx]
                    _sel_text = _sel_item.get("text", "")
                    _tag_info = REDACTION_TAGS.get(_sel_item.get("tag", ""), {})

                    _pcol1, _pcol2 = st.columns([2, 3])
                    with _pcol1:
                        st.markdown(f"**{_tag_info.get('label', _sel_item.get('tag', ''))}**")
                        st.caption(_sel_item.get("reason", ""))
                        st.code(_sel_text, language=None)
                        _ctx = _sel_item.get("context", "")
                        if _ctx:
                            st.markdown("**Surrounding context (from LLM):**")
                            _ctx_hl = _ctx.replace(
                                _sel_text,
                                f"**:orange[{_sel_text}]**"
                            )
                            st.markdown(f"> {_ctx_hl}")

                    with _pcol2:
                        _png, _pnum, _found = _render_context_preview(
                            analysis["doc"], _sel_text
                        )
                        if _found:
                            st.caption(f"📄 Page {_pnum} — yellow = matched text")
                            st.image(_png, use_container_width=True)
                        else:
                            st.caption(
                                "Term not found on a rendered page — "
                                "may be in an image layer or OCR text only."
                            )

            elif analysis["has_text"] and not analysis["escalations"]:
                st.success("No redactions proposed. Verify the document manually if needed.")

            # ── Page exclusion ────────────────────────────────────────────────
            _doc_for_pages = analysis.get("doc")
            if _doc_for_pages is not None and _doc_for_pages.page_count > 0:
                st.markdown("#### 🗑️ Remove Whole Pages")
                st.caption(
                    "Select pages to delete entirely from the final bundle — e.g. misfiled "
                    "records, letters about another patient, blank pages."
                )
                _page_labels = []
                for _pn in range(_doc_for_pages.page_count):
                    _pg_txt = _doc_for_pages[_pn].get_text("text")[:90].replace("\n", " ").strip()
                    _pg_lbl = f"Page {_pn + 1}"
                    if _pg_txt:
                        _pg_lbl += f" — {_pg_txt}…"
                    _page_labels.append(_pg_lbl)

                _excl_sel = st.multiselect(
                    "Pages to remove:",
                    options=_page_labels,
                    key=f"page_excl_{i}",
                    placeholder="None — all pages will be included",
                )
                if _excl_sel:
                    if len(_excl_sel) == _doc_for_pages.page_count:
                        st.error(
                            "⚠️ All pages selected for removal — this document will be "
                            "excluded entirely from the bundle."
                        )
                    else:
                        st.warning(
                            f"🗑️ {len(_excl_sel)} page(s) will be removed: "
                            + ", ".join(
                                lbl.split(" — ")[0] for lbl in _excl_sel
                            )
                        )

    # ── Apply button ──────────────────────────────────────────────────────────
    st.divider()
    _approved_final = sum(
        sum(1 for r in a["proposed_redactions"] if r.get("approved", True))
        for a in _analyses
    )
    st.markdown(f"**{_approved_final} redaction(s) approved** and ready to apply.")

    if st.button("Apply Approved Redactions & Build Bundle", type="primary", use_container_width=True):
        prog   = st.progress(0.0)
        status = st.empty()
        proc   = []

        _apply_start = time.time()
        for i, analysis in enumerate(_analyses):
            if analysis.get("error") or analysis.get("doc") is None:
                continue
            prog.progress(
                i / max(len(_analyses), 1),
                text=(
                    f"Redacting {analysis['filename']} ({i + 1}/{len(_analyses)})"
                    + _timing_suffix(_apply_start, i, len(_analyses))
                ),
            )
            status.markdown(f"✏️ Applying redactions to **{analysis['filename']}**…")
            approved = [r for r in analysis["proposed_redactions"] if r.get("approved", True)]
            doc, cnt = apply_approved_redactions(analysis["doc"], approved)

            # ── Remove any pages the reviewer marked for deletion ─────────────
            _excl_labels = st.session_state.get(f"page_excl_{i}", [])
            if _excl_labels:
                # Parse "Page N — ..." → zero-based index; delete in reverse to preserve indices
                _excl_nums = sorted(
                    [int(lbl.split(" — ")[0].replace("Page ", "").strip()) - 1
                     for lbl in _excl_labels],
                    reverse=True,
                )
                for _pn in _excl_nums:
                    if 0 <= _pn < doc.page_count:
                        doc.delete_page(_pn)
                cnt += len(_excl_labels)   # count deleted pages as redaction events

            # Skip document if all pages were removed
            if doc.page_count == 0:
                continue

            proc.append({
                "filename":        analysis["filename"],
                "section":         analysis["section"],
                "doc":             doc,
                "redaction_count": cnt,
                "doc_date":        analysis.get("doc_date", datetime.date.min),
            })

        prog.progress(0.9, text="Building PDF bundle…")
        status.markdown("📎 Building bundle PDF…")

        bundle = build_bundle(
            proc,
            sar_ref=sar_ref,
            operator=operator_name,
            date_str=datetime.date.today().strftime("%d/%m/%Y"),
        )
        buf = io.BytesIO()
        bundle.save(buf)
        buf.seek(0)

        today    = datetime.date.today().strftime("%Y%m%d")
        safe_ref = re.sub(r"[^\w\-]", "_", sar_ref or "SAR")
        # Include patient name in filename so the output is clearly identified.
        # Use the entered patient name; fall back to auto-detection from the first file.
        _fname_patient = patient_name.strip()
        if not _fname_patient and _analyses:
            _first_fname = analyses[0].get("filename", "")
            _first_text  = analyses[0].get("text", "")
            _fname_patient = _detect_patient_name(_first_fname, _first_text)
        safe_pname = re.sub(r"[^\w\-]", "_", _fname_patient) if _fname_patient else ""
        if safe_pname:
            _fname_core = f"{today}_{safe_pname}_REDACTED_BUNDLE.pdf"
        else:
            _fname_core = f"{today}_{safe_ref}_REDACTED_BUNDLE.pdf"
        st.session_state.bundle_bytes  = buf.getvalue()
        st.session_state.bundle_fname  = _fname_core
        st.session_state.proc_summary  = [
            {
                "File":       p["filename"],
                "Section":    p["section"],
                "Redactions": p["redaction_count"],
                "Status":     "✅",
            }
            for p in proc
        ]
        prog.progress(1.0)
        status.empty()
        st.session_state.stage      = "export"
        st.session_state.play_sound = "fanfare"
        st.rerun()


# =============================================================================
# STAGE 3 — EXPORT
# =============================================================================

elif tool_mode == "sar" and st.session_state.stage == "export":
    st.divider()
    st.success("✅ Redacted bundle is ready for download — review all decisions before releasing to the data subject.")
    st.markdown(
        '<div class="sar-disclaimer">'
        '<b>⚠ Important — Human review required before release</b><br>'
        'This output was generated by an AI system currently in beta testing. '
        'It must be checked by a qualified Information Governance or clinical professional '
        'before being sent to the data subject. The tool\'s authors accept no liability for '
        'incorrect, incomplete, or excessive redactions. Use of this tool in a live SAR '
        'process is entirely at the discretion and responsibility of the operating organisation.'
        '</div>',
        unsafe_allow_html=True,
    )

    summary = st.session_state.get("proc_summary", [])
    if summary:
        total_r = sum(d["Redactions"] for d in summary)
        c1, c2, c3 = st.columns(3)
        c1.metric("Documents in bundle",      len(summary))
        c2.metric("Total redactions applied", total_r)
        c3.metric("Reviewed by",              operator_name or "—")

        if PANDAS_AVAILABLE:
            st.dataframe(
                pd.DataFrame(summary)[["File", "Section", "Redactions", "Status"]],
                use_container_width=True,
                hide_index=True,
            )

    st.download_button(
        label="⬇  Download Redacted Bundle PDF",
        data=st.session_state.bundle_bytes,
        file_name=st.session_state.bundle_fname,
        mime="application/pdf",
        use_container_width=True,
        type="primary",
    )

    st.info(
        "**Governance reminder (ICO / BMA / NHS England):**  "
        "All SAR redaction decisions should be reviewed by an appropriate clinician or IG lead. "
        "Decisions must be defensible under UK GDPR and DPA 2018. "
        "Patients should be informed when exemptions have been applied and of their right "
        "to complain to the ICO."
    )

    if st.button("🔄 Process Another SAR", use_container_width=True):
        _reset()
        st.rerun()


# =============================================================================
# FULL ANONYMISATION MODE
# =============================================================================

elif tool_mode == "anon":
    st.divider()
    st.subheader("① Upload Documents to Anonymise")
    st.caption(
        "Upload one or more clinical documents. All patient and person identifiers will be "
        "automatically removed and replaced with labelled placeholders — e.g. **[PATIENT NAME]**, "
        "**[DATE OF BIRTH]**, **[NAME]**. Clinical content is preserved.  \n"
        "Supported: **PDF · Word (.docx) · RTF · TXT · ZIP**"
    )

    anon_files = st.file_uploader(
        "Browse or drop files",
        type=ACCEPTED_FORMATS,
        accept_multiple_files=True,
        label_visibility="collapsed",
        key="anon_uploader",
    )

    anon_folder = st.text_input(
        "Or load all documents from a folder path:",
        placeholder=r"e.g. C:\Patient Records\John Smith",
        key="anon_folder_input",
    ).strip()

    _anon_folder_files = []
    if anon_folder:
        _afp = Path(anon_folder)
        if _afp.is_dir():
            _anon_folder_files = [
                f for f in sorted(_afp.iterdir())
                if f.is_file() and f.suffix.lower().lstrip(".") in _SUPPORTED_EXTS | {"zip"}
            ]
            if _anon_folder_files:
                st.info(
                    f"📁 **{len(_anon_folder_files)} file(s) found:**  "
                    + "  ·  ".join(f.name for f in _anon_folder_files[:12])
                    + ("…" if len(_anon_folder_files) > 12 else "")
                )
            else:
                st.warning("No supported files found in that folder.")
        else:
            st.error(f"Folder not found: `{anon_folder}`")

    _anon_any = bool(anon_files) or bool(_anon_folder_files)

    if _anon_any and st.button("🕵️ Anonymise Documents", type="primary", use_container_width=True):
        all_anon_files = _collect_all_files(anon_files, anon_folder)
        results = []
        prog   = st.progress(0.0)
        status = st.empty()

        _anon_outer_start = time.time()
        for i, uf in enumerate(all_anon_files):
            status.info(
                f"Processing file {i + 1}/{len(all_anon_files)}: {uf.name}"
                + _timing_suffix(_anon_outer_start, i, len(all_anon_files))
            )
            _doc, text, _err, _ocr = ingest_file(uf)
            if not text.strip():
                results.append({"name": uf.name, "text": "", "count": 0, "error": _err or "Could not extract text"})
                prog.progress((i + 1) / len(all_anon_files))
                continue

            def _cb(msg, _name=uf.name):
                status.info(f"**{_name}** — {msg}")

            anon_text, count, _raw, llm_fails = anonymise_document(text, selected_model, status_cb=_cb)
            warn = (f"⚠️ {llm_fails} chunk(s) failed LLM parsing — "
                    "only regex-based redactions (NHS numbers, postcodes) were applied. "
                    "Try switching to qwen2.5:14b or qwen2.5:32b for full name/DOB redaction."
                    ) if llm_fails else None
            results.append({"name": uf.name, "text": anon_text, "count": count, "error": warn})
            prog.progress((i + 1) / len(all_anon_files))

        # Build ZIP of anonymised .txt files
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for r in results:
                stem = Path(r["name"]).stem
                zf.writestr(f"{stem}_ANONYMISED.txt", r["text"] or f"[Error: {r['error']}]")
        zip_buf.seek(0)

        st.session_state["anon_results"]   = results
        st.session_state["anon_zip_bytes"] = zip_buf.getvalue()
        prog.progress(1.0)
        status.empty()
        st.rerun()

    # ── Results ───────────────────────────────────────────────────────────────
    if st.session_state.get("anon_results"):
        results   = st.session_state["anon_results"]
        zip_bytes = st.session_state.get("anon_zip_bytes")

        st.success(f"✅ {len(results)} document(s) anonymised.")
        total_redactions = sum(r["count"] for r in results)
        c1, c2 = st.columns(2)
        c1.metric("Documents processed", len(results))
        c2.metric("Identifiers removed", total_redactions)

        st.markdown("---")

        today = datetime.date.today().strftime("%Y%m%d")
        st.download_button(
            label="⬇  Download Anonymised Documents (ZIP)",
            data=zip_bytes,
            file_name=f"{today}_ANONYMISED_BUNDLE.zip",
            mime="application/zip",
            use_container_width=True,
            type="primary",
        )

        st.markdown("---")
        st.subheader("Preview")
        for r in results:
            with st.expander(f"📄 {r['name']}  —  {r['count']} identifier(s) removed"):
                if r["error"]:
                    st.error(r["error"])
                else:
                    st.text_area(
                        "Anonymised text",
                        value=r["text"],
                        height=300,
                        key=f"anon_preview_{r['name']}",
                        label_visibility="collapsed",
                    )

        st.info(
            "**Reminder:** Review the anonymised output before sharing externally. "
            "The AI may occasionally miss an identifier or leave an indirect re-identification risk. "
            "All processing is 100% local — no data leaves this machine."
        )


# =============================================================================
# INSURANCE FORM FILLER — STAGE 1: UPLOAD & EXTRACT
# =============================================================================

elif tool_mode == "form_filler" and st.session_state.ff_stage == "ff_upload":
    st.divider()
    st.subheader("① Load Patient Records & Upload Insurance Form")
    st.caption(
        "Upload the **patient record ZIP** (exported from your EPR system) and the "
        "**insurance / GP report form** you need to complete. "
        "The AI will read the records and suggest answers for every field on the form."
    )

    col_epr, col_form = st.columns(2)

    with col_epr:
        st.markdown("**Patient Records (EPR export)**")
        epr_upload = st.file_uploader(
            "EPR ZIP or individual files",
            type=["zip", "pdf", "docx", "doc", "txt", "rtf", "tiff", "tif"],
            accept_multiple_files=True,
            key="ff_epr_uploader",
            label_visibility="collapsed",
        )
        epr_folder = st.text_input(
            "Or load from folder path:",
            placeholder=r"e.g. C:\Patient Records\John Smith",
            key="ff_epr_folder",
        ).strip()

    with col_form:
        st.markdown("**Insurance / GP Report Form**")
        form_upload = st.file_uploader(
            "Scanned form (PDF, JPEG, PNG, TIFF)",
            type=["pdf", "jpg", "jpeg", "png", "tiff", "tif"],
            accept_multiple_files=False,
            key="ff_form_uploader",
            label_visibility="collapsed",
        )
        if form_upload:
            st.caption(f"Form: **{form_upload.name}**")
            if not TESSERACT_AVAILABLE and form_upload.name.lower().endswith(
                (".jpg", ".jpeg", ".png", ".tiff", ".tif")
            ):
                st.warning("Tesseract OCR not available — text cannot be extracted from image files.")

    _epr_ready  = bool(epr_upload) or bool(epr_folder)
    _form_ready = bool(form_upload)

    if _epr_ready and _form_ready:
        if st.button(
            "Extract Form Fields & Answer from Patient Records",
            type="primary",
            use_container_width=True,
        ):
            prog   = st.progress(0.0, text="⏳ Ingesting patient records…")
            status = st.empty()

            # ── Ingest EPR records ────────────────────────────────────────────
            status.info("📂 Reading patient record files…")
            all_epr_files = _collect_all_files(epr_upload, epr_folder)
            epr_text_parts = []
            epr_docs       = []

            _epr_start = time.time()
            for fi, epr_file in enumerate(all_epr_files):
                prog.progress(
                    0.05 + 0.3 * fi / max(len(all_epr_files), 1),
                    text=(
                        f"Reading {epr_file.name} ({fi + 1}/{len(all_epr_files)})"
                        + _timing_suffix(_epr_start, fi, len(all_epr_files))
                    ),
                )
                if getattr(epr_file, "_zip_error", None):
                    continue
                fdoc, ftext, ferr, _ = ingest_file(epr_file)
                if ftext.strip():
                    epr_text_parts.append(ftext)
                if fdoc:
                    epr_docs.append(fdoc)

            epr_combined = "\n\n---\n\n".join(epr_text_parts)
            if not epr_combined.strip():
                st.error("Could not extract text from the patient record files. "
                         "Check the files are readable and Tesseract is installed for image-only PDFs.")
                st.stop()

            # ── Ingest the form ───────────────────────────────────────────────
            prog.progress(0.38, text="Reading insurance form…")
            status.info("📋 Reading insurance form…")
            form_doc, form_text, form_err, has_text_layer = _ingest_form(form_upload)

            if form_err or form_doc is None:
                st.error(f"Could not open the form: {form_err}")
                st.stop()
            if not form_text.strip():
                st.error(
                    "No text could be extracted from the form.  \n"
                    "Ensure Tesseract is installed for scanned images/PDFs."
                )
                st.stop()

            # ── LLM: extract form fields ──────────────────────────────────────
            prog.progress(0.45, text="Identifying form fields…")
            status.info("🤖 Identifying form fields (this may take a moment)…")
            fields = extract_form_fields_llm(form_text, selected_model)

            if not fields:
                st.error(
                    "The AI could not identify any fields on this form.  \n"
                    "Try a different model or check the form text was extracted correctly."
                )
                st.stop()

            # ── LLM: answer fields from EPR ───────────────────────────────────
            prog.progress(0.6, text="Answering form fields from patient record…")
            _pname = patient_name.strip() if patient_name.strip() else st.session_state.get("ff_patient_name", "")

            answered = answer_fields_from_epr(
                fields,
                epr_combined,
                _pname,
                selected_model,
                status_cb=lambda msg: status.info(f"🤖 {msg}"),
            )

            prog.progress(1.0)
            status.empty()

            st.session_state.ff_epr_text       = epr_combined
            st.session_state.ff_epr_docs       = epr_docs
            st.session_state.ff_form_doc       = form_doc
            st.session_state.ff_form_text      = form_text
            st.session_state.ff_has_text_layer = has_text_layer
            st.session_state.ff_fields         = answered
            st.session_state.ff_patient_name   = _pname
            st.session_state.ff_stage          = "ff_review"
            st.session_state.play_sound        = "chime"
            st.rerun()

    elif not _epr_ready:
        st.info("Upload the patient record ZIP or files (left column) to begin.")
    elif not _form_ready:
        st.info("Upload the insurance form (right column) to begin.")


# =============================================================================
# INSURANCE FORM FILLER — STAGE 2: REVIEW ANSWERS
# =============================================================================

elif tool_mode == "form_filler" and st.session_state.ff_stage == "ff_review":
    st.divider()

    _ff_fields   = st.session_state.ff_fields
    _form_doc    = st.session_state.ff_form_doc
    _epr_text    = st.session_state.ff_epr_text
    _epr_docs    = st.session_state.ff_epr_docs
    _has_tl      = st.session_state.ff_has_text_layer
    _pname       = st.session_state.ff_patient_name

    _n_auto      = sum(1 for f in _ff_fields if not f.get("needs_manual_input"))
    _n_manual    = sum(1 for f in _ff_fields if f.get("needs_manual_input"))
    _n_answered  = sum(1 for f in _ff_fields if f.get("answer"))

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Fields found",         len(_ff_fields))
    m2.metric("Answered from record", _n_answered)
    m3.metric("Need your input",      _n_manual)
    m4.metric("Patient",              _pname or "—")

    st.markdown(
        "Review each suggested answer below. Edit any answer, then approve it. "
        "Fields highlighted in **orange** require information that isn't in the patient record — "
        "please fill these in manually."
    )

    # ── Form preview ─────────────────────────────────────────────────────────
    if _form_doc and PIL_AVAILABLE:
        with st.expander("🖼 View original form", expanded=False):
            for pg_num, pg in enumerate(_form_doc):
                pix = pg.get_pixmap(dpi=110)
                img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                st.image(buf.getvalue(), caption=f"Form page {pg_num + 1}", use_container_width=True)

    st.divider()

    # ── Fields answered from patient record ───────────────────────────────────
    _auto_fields = [f for f in _ff_fields if not f.get("needs_manual_input")]
    if _auto_fields:
        st.markdown("### Answers from Patient Record")
        st.caption(
            "These answers were suggested by the AI from the patient record. "
            "Edit if needed, then tick **Approved** to include in the output."
        )

    for fi, field in enumerate(_ff_fields):
        if field.get("needs_manual_input"):
            continue

        label      = field.get("label", f"Field {fi + 1}")
        conf       = field.get("confidence", "none")
        evidence   = field.get("evidence", "")
        key_ans    = f"ff_ans_{fi}"
        key_appr   = f"ff_appr_{fi}"

        _conf_col = {
            "high":   "rgba(28,155,60,.18)",
            "medium": "rgba(0,94,184,.14)",
            "low":    "rgba(255,140,0,.14)",
            "none":   "rgba(180,0,0,.10)",
        }.get(conf, "rgba(255,255,255,.04)")

        with st.container(border=True):
            c_label, c_conf = st.columns([4, 1])
            with c_label:
                st.markdown(f"**{label}**")
            with c_conf:
                _cmap = {"high": "🟢 High", "medium": "🔵 Medium", "low": "🟡 Low", "none": "🔴 None"}
                st.caption(_cmap.get(conf, conf))

            # Editable answer
            current_ans = st.session_state.get(key_ans, field.get("final_answer", ""))
            new_ans = st.text_input(
                "Answer",
                value=current_ans,
                key=key_ans,
                label_visibility="collapsed",
                placeholder="Type answer here…",
            )
            field["final_answer"] = new_ans

            # Approval checkbox
            default_approved = bool(field.get("answer")) and conf in ("high", "medium")
            field["approved"] = st.checkbox(
                "Approved — include this answer in the output",
                value=st.session_state.get(key_appr, default_approved),
                key=key_appr,
            )

            # Evidence panel
            if evidence:
                with st.expander("📄 Evidence from patient record"):
                    st.markdown(f"> {evidence}")
                    # Try to find and highlight the text in the EPR document
                    if _epr_docs and PIL_AVAILABLE:
                        for epr_doc in _epr_docs[:3]:   # check up to 3 docs
                            _png, _pnum, _found = _render_context_preview(epr_doc, evidence[:60])
                            if _found:
                                st.caption(f"Page {_pnum} of patient record — highlighted in yellow")
                                st.image(_png, use_container_width=True)
                                break

    # ── Manual input fields ───────────────────────────────────────────────────
    _manual_fields = [f for f in _ff_fields if f.get("needs_manual_input")]
    if _manual_fields:
        st.divider()
        st.markdown("### Additional Information Required from You")
        st.caption(
            "The following fields could not be answered from the patient record. "
            "Please provide the information manually."
        )

        for fi, field in enumerate(_ff_fields):
            if not field.get("needs_manual_input"):
                continue

            label = field.get("label", f"Field {fi + 1}")
            hint  = field.get("manual_hint", "")
            key_m = f"ff_man_{fi}"

            with st.container(border=True):
                st.markdown(f"**{label}**")
                if hint:
                    st.caption(hint)
                man_val = st.text_input(
                    "Your answer",
                    value=st.session_state.get(key_m, ""),
                    key=key_m,
                    label_visibility="collapsed",
                    placeholder="Enter answer…",
                )
                field["final_answer"] = man_val
                field["approved"]     = bool(man_val.strip())
                field["confidence"]   = "none"
                field["evidence"]     = ""

    # ── Build output button ───────────────────────────────────────────────────
    st.divider()
    _n_approved = sum(1 for f in _ff_fields if f.get("approved") and f.get("final_answer", "").strip())
    st.markdown(f"**{_n_approved} of {len(_ff_fields)} fields** will be included in the output.")

    if st.button("Build Filled Form PDF", type="primary", use_container_width=True):
        with st.spinner("Building PDF…"):
            filled_doc = build_filled_form_pdf(
                _form_doc, _ff_fields, has_text_layer=_has_tl
            )
            buf = io.BytesIO()
            filled_doc.save(buf)
            buf.seek(0)
            st.session_state.ff_filled_bytes = buf.getvalue()
            st.session_state.ff_stage        = "ff_export"
            st.session_state.play_sound      = "fanfare"
            st.rerun()


# =============================================================================
# INSURANCE FORM FILLER — STAGE 3: EXPORT
# =============================================================================

elif tool_mode == "form_filler" and st.session_state.ff_stage == "ff_export":
    st.divider()
    st.success("✅ Filled form PDF is ready — review all answers before sending to the insurer.")

    st.markdown(
        '<div class="sar-disclaimer">'
        '<b>⚠ Important — Human review required</b><br>'
        'All answers were suggested by an AI and must be verified by the responsible clinician '
        'or practice manager before the form is submitted. The authors accept no liability for '
        'incorrect or incomplete answers.'
        '</div>',
        unsafe_allow_html=True,
    )

    _ff_fields = st.session_state.ff_fields
    _n_filled  = sum(1 for f in _ff_fields if f.get("final_answer", "").strip())
    _n_manual  = sum(1 for f in _ff_fields if f.get("needs_manual_input") and f.get("final_answer", "").strip())
    _n_epr     = _n_filled - _n_manual

    c1, c2, c3 = st.columns(3)
    c1.metric("Total fields",         len(_ff_fields))
    c2.metric("Answered from record", _n_epr)
    c3.metric("Manually provided",    _n_manual)

    # Summary table
    if PANDAS_AVAILABLE:
        summary_rows = []
        for f in _ff_fields:
            source = "Manual" if f.get("needs_manual_input") else f"AI ({f.get('confidence','?')} confidence)"
            summary_rows.append({
                "Field":   f.get("label", ""),
                "Answer":  f.get("final_answer", "") or "—",
                "Source":  source,
                "Status":  "✅" if f.get("final_answer", "").strip() else "⚠ Empty",
            })
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

    today    = datetime.date.today().strftime("%Y%m%d")
    pname    = st.session_state.get("ff_patient_name", "Patient") or "Patient"
    safe_p   = re.sub(r"[^\w\-]", "_", pname)

    st.download_button(
        label="⬇  Download Filled Form PDF",
        data=st.session_state.ff_filled_bytes,
        file_name=f"{today}_{safe_p}_COMPLETED_FORM.pdf",
        mime="application/pdf",
        use_container_width=True,
        type="primary",
    )

    if st.button("🔄 Complete Another Form", use_container_width=True):
        _reset_ff()
        st.rerun()
